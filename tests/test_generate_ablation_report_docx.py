import csv
import json
import sys
import unittest
import tempfile
from pathlib import Path

try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    Document = None
    HAS_PYTHON_DOCX = False


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from generate_ablation_report_docx import (  # noqa: E402
    CORE_TABLES,
    build_document,
    read_csv_rows,
    summarize_findings,
)


def write_csv(path: Path, rows):
    if not rows:
        path.write_text("", encoding="utf-8")
        return
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


class GenerateAblationReportDocxTest(unittest.TestCase):
    @unittest.skipUnless(HAS_PYTHON_DOCX, "python-docx is not installed")
    def test_read_csv_rows_returns_empty_for_missing_file(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            rows = read_csv_rows(Path(tmpdir) / "missing.csv")

        self.assertEqual(rows, [])

    @unittest.skipUnless(HAS_PYTHON_DOCX, "python-docx is not installed")
    def test_summarize_findings_uses_table_values(self):
        tables = {
            "ablation": [
                {
                    "variant": "Full BEA-Judge",
                    "head": "pairwise",
                    "accuracy": "0.75",
                    "macro_f1": "0.67",
                    "tie_recall": "0.52",
                },
                {
                    "variant": "Full BEA-Judge",
                    "head": "factuality",
                    "accuracy": "0.76",
                    "macro_f1": "0.74",
                },
                {
                    "variant": "w/o Evidence Module",
                    "head": "factuality",
                    "macro_f1": "0.65",
                },
            ],
            "significance": [
                {
                    "variant": "w/o Evidence Module",
                    "head": "factuality",
                    "delta_macro_f1_full_minus_variant": "0.09",
                    "mcnemar_p": "0.001",
                },
                {
                    "variant": "w/o Calibration",
                    "head": "pairwise",
                    "delta_macro_f1_full_minus_variant": "0.03",
                    "mcnemar_p": "0.04",
                },
            ],
            "bias_utility": [
                {
                    "setting": "bias_as_decision_features",
                    "review_capture_rate": "0.75",
                },
                {
                    "setting": "no_bias_decision_features",
                    "macro_f1": "0.69",
                },
            ],
        }

        findings = summarize_findings(tables)

        joined = "\n".join(findings)
        self.assertIn("0.75", joined)
        self.assertIn("0.09", joined)
        self.assertIn("0.75", joined)

    @unittest.skipUnless(HAS_PYTHON_DOCX, "python-docx is not installed")
    def test_build_document_creates_docx_with_core_sections(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp = Path(tmpdir)
            tables_dir = tmp / "tables"
            tables_dir.mkdir()
            payload = {
                "ablation": [
                    {
                        "variant": "Full BEA-Judge",
                        "head": "pairwise",
                        "n": "2",
                        "accuracy": "1.0",
                        "macro_f1": "1.0",
                        "ece": "0.0",
                        "brier": "0.0",
                        "tie_recall": "1.0",
                        "review_rate": "0.5",
                    },
                    {
                        "variant": "Full BEA-Judge",
                        "head": "factuality",
                        "n": "1",
                        "accuracy": "0.9",
                        "macro_f1": "0.8",
                        "ece": "0.1",
                        "brier": "0.2",
                        "tie_recall": "",
                        "review_rate": "0.5",
                    },
                    {
                        "variant": "w/o Evidence Module",
                        "head": "factuality",
                        "n": "1",
                        "accuracy": "0.7",
                        "macro_f1": "0.6",
                        "ece": "0.2",
                        "brier": "0.3",
                        "tie_recall": "",
                        "review_rate": "0.5",
                    },
                ],
                "significance": [
                    {
                        "variant": "w/o Evidence Module",
                        "head": "factuality",
                        "paired_n": "1",
                        "delta_accuracy_full_minus_variant": "0.2",
                        "delta_macro_f1_full_minus_variant": "0.2",
                        "mcnemar_full_only_correct": "1",
                        "mcnemar_variant_only_correct": "0",
                        "mcnemar_p": "0.5",
                    },
                    {
                        "variant": "w/o Calibration",
                        "head": "pairwise",
                        "paired_n": "2",
                        "delta_accuracy_full_minus_variant": "0.1",
                        "delta_macro_f1_full_minus_variant": "0.1",
                        "mcnemar_full_only_correct": "1",
                        "mcnemar_variant_only_correct": "0",
                        "mcnemar_p": "0.5",
                    },
                ],
                "evidence_groups": [
                    {
                        "feature_group": "full",
                        "weighted_calibration": "True",
                        "feature_count": "4",
                        "accuracy": "0.9",
                        "macro_f1": "0.8",
                        "ece": "0.1",
                        "brier": "0.2",
                    }
                ],
                "bias_utility": [
                    {
                        "setting": "bias_as_decision_features",
                        "head": "pairwise",
                        "n": "2",
                        "accuracy": "1.0",
                        "macro_f1": "1.0",
                        "ece": "0.0",
                        "review_rate": "0.5",
                        "review_capture_rate": "1.0",
                    },
                    {
                        "setting": "no_bias_decision_features",
                        "head": "pairwise",
                        "n": "2",
                        "accuracy": "1.0",
                        "macro_f1": "1.0",
                        "ece": "0.0",
                        "review_rate": "0.0",
                        "review_capture_rate": "0.0",
                    },
                ],
                "calibration": [
                    {
                        "method": "temperature",
                        "split": "test",
                        "accuracy": "1.0",
                        "ece": "0.0",
                        "mce": "0.0",
                        "brier": "0.0",
                        "nll": "0.0",
                        "coverage": "",
                        "set_size_avg": "",
                    }
                ],
                "risk_coverage": [
                    {
                        "head": "pairwise",
                        "split": "test",
                        "review_rate": "0.2000",
                        "review_count": "1",
                        "error_capture_rate": "1.0",
                        "auto_accept_count": "1",
                        "auto_accept_accuracy": "1.0",
                        "risk_threshold": "0.5",
                    }
                ],
                "ragtruth": [
                    {
                        "split": "test",
                        "n": "1",
                        "accuracy": "1.0",
                        "macro_f1": "1.0",
                        "ece": "0.0",
                        "brier": "0.0",
                        "review_rate": "0.0",
                        "supported_to_unsupported": "0",
                        "unsupported_to_supported": "0",
                    }
                ],
            }
            for key, filename in CORE_TABLES.items():
                write_csv(tables_dir / filename, payload.get(key, []))
            ablation_report = tmp / "ablation.json"
            ablation_report.write_text(
                json.dumps({"created_at": "2026-05-23T00:00:00+00:00", "input_dataset": "unit"}),
                encoding="utf-8",
            )
            output = tmp / "out.docx"

            build_document(
                title="Unit Ablation Report",
                tables_dir=tables_dir,
                ablation_report_path=ablation_report,
                output_path=output,
            )

            self.assertTrue(output.exists())
            document = Document(output)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("Unit Ablation Report", text)
            self.assertIn("执行摘要", text)
            self.assertGreaterEqual(len(document.tables), 8)


if __name__ == "__main__":
    unittest.main()
