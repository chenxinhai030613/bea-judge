from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "数据集构建章节_SCI规范稿.md"
OUTPUT = ROOT / "数据集构建章节_SCI规范稿.docx"


ACCENT = RGBColor(31, 78, 121)
MUTED = RGBColor(90, 90, 90)
GRID = "D9E2F3"
HEADER_FILL = "EAF2F8"
CODE_FILL = "F5F7FA"


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_style_font(style, latin: str, east_asia: str, size: int, bold: bool = False, color=None) -> None:
    font = style.font
    font.name = latin
    font.size = Pt(size)
    font.bold = bold
    if color is not None:
        font.color.rgb = color
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_bottom_border(paragraph, color: str = "B7C9DA", size: str = "6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths_dxa: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def format_sha(text: str) -> str:
    if re.fullmatch(r"[0-9a-f]{64}", text):
        return "\n".join(text[i : i + 16] for i in range(0, 64, 16))
    return text


def parse_markdown(md: str):
    lines = md.splitlines()
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("```"):
            lang = line.strip("`").strip()
            code = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            blocks.append(("code", lang, "\n".join(code)))
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            blocks.append(("heading", level, line[level:].strip()))
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = []
            for idx, row in enumerate(table_lines):
                cells = [c.strip() for c in row.strip().strip("|").split("|")]
                if idx == 1 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
                    continue
                rows.append(cells)
            blocks.append(("table", rows))
            continue
        if line.startswith("**") and line.endswith("**"):
            blocks.append(("caption", line.strip("*")))
            i += 1
            continue
        paragraph = [line.strip()]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if not nxt.strip() or nxt.startswith("#") or nxt.startswith("|") or nxt.startswith("```"):
                break
            if nxt.startswith("**") and nxt.endswith("**"):
                break
            paragraph.append(nxt.strip())
            i += 1
        blocks.append(("paragraph", " ".join(paragraph)))
    return blocks


def add_markdown_text(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        set_east_asia_font(run, "宋体")


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.keep_together = True
    add_bottom_border(p, color="E1E6EF", size="4")
    for idx, line in enumerate(code.splitlines()):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CODE_FILL)
    p_pr.append(shd)


def column_widths(num_cols: int) -> list[int]:
    width = 9360
    presets = {
        2: [3600, 5760],
        3: [2800, 2480, 4080],
        4: [2460, 2300, 1200, 3400],
        5: [1700, 1700, 1100, 2250, 2610],
        6: [1420, 1420, 1510, 700, 2460, 1850],
    }
    return presets.get(num_cols, [width // num_cols] * num_cols)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    widths = column_widths(len(rows[0]))
    set_table_width(table, widths)

    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                shade_cell(cell, HEADER_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 18 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(format_sha(text))
            set_east_asia_font(run, "宋体")
            run.font.size = Pt(8 if len(rows[0]) >= 5 else 9)
            run.bold = r_idx == 0
        prevent_row_split(table.rows[r_idx])
    mark_header_row(table.rows[0])
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    set_style_font(styles["Normal"], "Times New Roman", "宋体", 11)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    styles["Normal"].paragraph_format.space_after = Pt(6)

    set_style_font(styles["Title"], "Times New Roman", "黑体", 20, True, ACCENT)
    styles["Title"].paragraph_format.space_after = Pt(12)
    set_style_font(styles["Heading 1"], "Times New Roman", "黑体", 16, True, ACCENT)
    styles["Heading 1"].paragraph_format.space_before = Pt(12)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    set_style_font(styles["Heading 2"], "Times New Roman", "黑体", 13, True, ACCENT)
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(5)

    if "Code Block" not in [s.name for s in styles]:
        style = styles.add_style("Code Block", 1)
    else:
        style = styles["Code Block"]
    set_style_font(style, "Consolas", "Consolas", 9)
    style.paragraph_format.left_indent = Cm(0.25)
    style.paragraph_format.right_indent = Cm(0.25)
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.0

    header_p = section.header.paragraphs[0]
    header_p.text = "BEA-Judge 数据集构建章节"
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header_p.runs[0]
    set_east_asia_font(header_run, "宋体")
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = MUTED
    add_bottom_border(header_p)

    footer_p = section.footer.paragraphs[0]
    add_page_number(footer_p)
    for run in footer_p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def build_docx() -> None:
    md = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)

    for block in parse_markdown(md):
        kind = block[0]
        if kind == "heading":
            _, level, text = block
            if level == 1:
                p = doc.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_markdown_text(p, text)
            else:
                doc.add_paragraph(text, style=f"Heading {min(level - 1, 2)}")
        elif kind == "caption":
            if block[1].startswith("表 1 ") or block[1].startswith("表 3 "):
                doc.add_page_break()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(block[1])
            set_east_asia_font(run, "黑体")
            run.font.size = Pt(10)
            run.bold = True
            run.font.color.rgb = ACCENT
        elif kind == "paragraph":
            p = doc.add_paragraph()
            add_markdown_text(p, block[1])
        elif kind == "code":
            add_code_block(doc, block[2])
        elif kind == "table":
            add_table(doc, block[1])

    props = doc.core_properties
    props.title = "数据集构建章节 SCI 规范稿"
    props.subject = "BEA-Judge dataset construction"
    props.author = "Codex"
    props.keywords = "BEA-Judge, dataset construction, SCI manuscript"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
