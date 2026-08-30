"""
Update the dataset-construction section of the BEA-Judge project plan DOCX.

The update keeps the original document structure, but aligns the data-source
table and preprocessing rules with the implemented Chinese 1000-sample build.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


DEFAULT_DOCX = Path("BEA-Judge模型构建全流程方案_期刊课题版.docx")


STATUS_PARAGRAPH = (
    "当前实现状态：截至 2026-05-13，中文专业标注集已扩充至 1000 条，"
    "其中 open_qa 400 条、pairwise_bias 400 条、factuality_rag 200 条；"
    "输出文件包括 datasets/processed/chinese_professional_annotated_1000.json、"
    "datasets/processed/chinese_professional_annotated_latest.json、datasets/splits_zh/"
    "train|dev|test.json、datasets/chinese_dataset_statistics.json 和 "
    "datasets/chinese_annotation_report.json。"
)


IMPLEMENTATION_PARAGRAPHS = [
    STATUS_PARAGRAPH,
    (
        "数据源整合规则：公开数据继续保留 MT-Bench、PandaLM、JudgeBench、"
        "WikiEval 与 ARES 的来源字段、原始记录标识和标签映射；自建中文数据统一写入 "
        "self_built_chinese_annotation 来源，并通过 task_type、language、split、"
        "human_score、metadata.field_contract 与 metadata.missing_reason 对齐到同一结构标准。"
    ),
    (
        "缺失值处理规则：必填文本字段 prompt、answer_a、answer_b 在相应任务契约下必须为非空；"
        "非必需的 context、reference 不再使用空字符串占位，统一使用 null，并在 metadata.missing_reason "
        "中记录 not_required_for_task。factuality_rag 样本必须保留 context，reference 建议保留。"
    ),
    (
        "质量门槛：生成后必须通过样本数、任务分布、语言分布、重复样本、split 泄漏、"
        "score_format/scoring_system 完整性、空字符串扫描和中文标注一致性检查；中文标注报告需记录 "
        "Cohen's kappa、仲裁数量、标签分布和领域分布。"
    ),
]


def set_matching_paragraph(document: Document, prefix: str, replacement: str) -> bool:
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            paragraph.text = replacement
            return True
    return False


def update_data_source_table(document: Document) -> bool:
    updated = False
    for table in document.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if headers[:3] != ["数据类型", "推荐数据源", "建议样本规模"]:
            continue
        for row in table.rows[1:]:
            cells = row.cells
            label = cells[0].text.strip()
            if label == "开放式回答质量":
                cells[2].text = "300-800；核心公开集按实验目标抽样，当前可复现构建支持每任务 400/800 条。"
                cells[4].text = "明确是否使用官方划分；统一标签为 A>B、B>A、Tie；保留 source_url 与原始记录标识。"
                updated = True
            elif label == "Judge 偏差样本":
                cells[2].text = "300-800；需覆盖 position、length、format、rubric_sensitivity 扰动。"
                cells[4].text = "swap 后必须映射回原始实际答案；同一 parent_id 的扰动样本不得跨 split 泄漏。"
                updated = True
            elif label == "RAG/事实性评价":
                cells[2].text = "300-800；优先保留带 context 的样本。"
                cells[4].text = "context 为 factuality_rag 必填字段；claim/evidence 标签和截断策略需可追溯。"
                updated = True
            elif label == "中文专业场景":
                cells[1].text = "管理、财经、政策、科研方法问答自建样本；含开放问答、偏差扰动和事实性/RAG。"
                cells[2].text = "1000（已生成：open_qa 400、pairwise_bias 400、factuality_rag 200）。"
                cells[3].text = "增强中文与专业场景适用性，并用于验证偏差、证据和校准模块在中文任务上的稳定性。"
                cells[4].text = "需提供标注规范、匿名化样例、字段契约；可选字段缺失使用 null，避免空字符串。"
                updated = True
        break
    return updated


def insert_status_block(document: Document) -> bool:
    if any(paragraph.text.strip() == STATUS_PARAGRAPH for paragraph in document.paragraphs):
        return False
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == "4.2 预处理流程":
            for text in IMPLEMENTATION_PARAGRAPHS:
                paragraph.insert_paragraph_before(text)
            return True
    return False


def update_document(path: Path) -> None:
    document = Document(path)

    set_matching_paragraph(
        document,
        "数据设计需覆盖",
        (
            "数据设计需覆盖开放式回答质量、评价偏差、事实性/RAG 和中文专业场景四类任务，"
            "以保证模型结论不是单一数据集上的偶然表现。公开数据优先保证可复现性，自建中文数据用于"
            "体现应用价值和场景适配；当前中文专业数据已按统一结构标准扩充至 1000 条。"
        ),
    )
    set_matching_paragraph(
        document,
        "完整性检查：",
        (
            "完整性检查：删除必填字段缺失、乱码、重复样本、明显无效回答和过短回答，并输出剔除日志；"
            "非必填字段缺失使用 null 表示，不使用空字符串或空格占位。"
        ),
    )
    set_matching_paragraph(
        document,
        "字段统一：",
        (
            "字段统一：将 instruction/query/question 统一为 prompt，将 response/answer/completion "
            "统一为 answer_a/answer_b，并用 field_contract 明确各任务的 context、reference、answer_b 要求。"
        ),
    )
    set_matching_paragraph(
        document,
        "数据统计：",
        (
            "数据统计：报告样本数、任务类型、语言、平均 prompt/answer/context 长度、标签分布、Tie 比例、"
            "claim 数量、空字符串扫描结果、字段契约覆盖率和标注一致性。"
        ),
    )
    set_matching_paragraph(
        document,
        "每条中文自建样本建议",
        (
            "中文自建样本按 1000 条规模组织，每条建议由 2 名标注者独立评分；分歧超过 2 分或偏好标签冲突时"
            "由第 3 人仲裁，并保留 annotator_votes、arbiter_label 与 agreement 字段。"
        ),
    )
    set_matching_paragraph(
        document,
        "中文专业数据样本规模有限",
        (
            "中文专业数据已由初始小规模样本扩充至 1000 条，但其外部有效性仍需在更多行业文本和真实应用场景中验证。"
        ),
    )
    update_data_source_table(document)
    insert_status_block(document)

    document.save(path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Update the dataset section in the BEA-Judge DOCX plan.")
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    update_document(args.docx)
    print(args.docx)


if __name__ == "__main__":
    main()
