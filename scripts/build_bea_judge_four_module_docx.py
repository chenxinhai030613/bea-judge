"""Build an academic DOCX describing the BEA-Judge four-module framework."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "论文撰写"
FIG_DIR = OUT_DIR / "figures"
DOCX_PATH = OUT_DIR / "BEA-Judge四模块模型框架学术说明.docx"
FIG_PATH = FIG_DIR / "bea_judge_four_module_framework.png"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2EC") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)


def add_formula(doc: Document, label: str, formula: str, explanation: str) -> None:
    p = doc.add_paragraph()
    p.style = "Formula"
    run = p.add_run(f"{label}  {formula}")
    run.bold = True
    doc.add_paragraph(explanation, style="FormulaExplanation")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in [
        ("Title", 20, "17365D"),
        ("Heading 1", 15, "17365D"),
        ("Heading 2", 12.5, "255F85"),
        ("Heading 3", 11.5, "255F85"),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)

    for style_name in ("Formula", "FormulaExplanation"):
        if style_name not in styles:
            styles.add_style(style_name, 1)
        style = styles[style_name]
        style.font.name = "Consolas" if style_name == "Formula" else "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(9.5)
        style.paragraph_format.left_indent = Cm(0.45)
        style.paragraph_format.right_indent = Cm(0.2)
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(3)
    styles["Formula"].font.color.rgb = RGBColor.from_string("1F4E79")
    styles["FormulaExplanation"].font.color.rgb = RGBColor.from_string("4A5568")


def build_diagram() -> None:
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

    FIG_DIR.mkdir(parents=True, exist_ok=True)
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    fig, ax = plt.subplots(figsize=(12, 6), dpi=220)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6)
    ax.axis("off")

    boxes = [
        (0.5, 3.7, 2.3, 1.0, "输入样本", "prompt/context\nanswer A/B\nrubric/reference"),
        (3.2, 4.1, 2.4, 0.9, "模块1：基础 Judge", "base scores\nA>B / B>A / Tie"),
        (3.2, 2.8, 2.4, 0.9, "模块2：偏差感知", "position/length/format\nrubric/source risk"),
        (3.2, 1.5, 2.4, 0.9, "模块3：证据增强事实性", "support/contradiction\nambiguity evidence"),
        (6.5, 2.8, 2.7, 1.25, "模块4：融合校准", "feature fusion\ncalibration\nconfidence"),
        (9.7, 3.15, 1.9, 0.9, "输出", "label\nconfidence\nreview flag"),
    ]
    colors = ["#E8F1FA", "#DFF3EA", "#FFF3D6", "#FCE7E7", "#E7E6F7", "#EDF7ED"]
    for (x, y, w, h, title, body), color in zip(boxes, colors):
        patch = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.04,rounding_size=0.08",
            linewidth=1.4,
            edgecolor="#375A7F",
            facecolor=color,
        )
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h * 0.66, title, ha="center", va="center", fontsize=10.5, weight="bold", color="#17365D")
        ax.text(x + w / 2, y + h * 0.30, body, ha="center", va="center", fontsize=8.5, color="#263238")

    arrows = [
        ((2.8, 4.2), (3.2, 4.55)),
        ((2.8, 4.0), (3.2, 3.25)),
        ((2.8, 3.8), (3.2, 1.95)),
        ((5.6, 4.55), (6.5, 3.65)),
        ((5.6, 3.25), (6.5, 3.35)),
        ((5.6, 1.95), (6.5, 3.05)),
        ((9.2, 3.42), (9.7, 3.60)),
    ]
    for start, end in arrows:
        ax.add_patch(FancyArrowPatch(start, end, arrowstyle="-|>", mutation_scale=13, linewidth=1.25, color="#375A7F"))

    ax.text(
        6,
        0.45,
        "注：偏差模块主要用于风险识别与人工复核优先级；证据模块主要服务事实性判断；融合校准层输出最终类别和置信度。",
        ha="center",
        va="center",
        fontsize=9,
        color="#4A5568",
    )
    fig.tight_layout(pad=0.4)
    fig.savefig(FIG_PATH, bbox_inches="tight")
    plt.close(fig)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].width = Cm(widths[i])
        set_cell_shading(hdr[i], "DDEBF7")
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].width = Cm(widths[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(8.8)
    doc.add_paragraph()


def build_docx() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_diagram()
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("BEA-Judge 四模块模型框架学术说明")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("用于论文方法部分撰写的真实、可复核技术描述")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("4A5568")

    doc.add_paragraph(
        "本文档基于当前项目中已经实现和已观察到的实验结果，系统说明 BEA-Judge 的四个组成模块：基础 Judge 评分模块、偏差感知模块、证据增强事实性模块，以及融合校准与置信度输出模块。为保证学术规范，本文仅陈述可由项目脚本、配置或已完成 seed42 pilot 结果支持的内容；未完成的 3-seed 复现实验和外部 holdout 不作为已验证结论。",
    )

    doc.add_heading("1. 框架总览", level=1)
    doc.add_paragraph(
        "BEA-Judge 采用模块化设计，将轻量级基础 Judge 的原始判断、偏差风险特征、证据支撑特征和校准后的融合预测结合起来。该设计的目标不是训练一个新的大语言模型，而是在 3B 级 Judge backbone 之上进行参数高效微调，并通过后处理与融合校准提升评估可靠性。"
    )
    doc.add_picture(str(FIG_PATH), width=Inches(6.8))
    add_caption(doc, "图 1. BEA-Judge 四模块框架结构图。")
    doc.add_paragraph(
        "图 1 显示了系统的数据流。输入样本首先进入基础 Judge 模块，得到 A/B/Tie 等原始评分信号；偏差感知模块并行估计位置、长度、格式和 rubric 等风险；证据增强事实性模块提取回答与上下文之间的支持关系；最后，融合校准模块整合多源特征，输出最终标签、校准置信度和人工复核信号。"
    )

    doc.add_heading("2. 符号定义与总体公式", level=1)
    add_table(
        doc,
        ["符号", "含义"],
        [
            ["x", "待评估样本，包括 prompt、context、rubric、reference、answer A 与 answer B。"],
            ["z_base", "基础 Judge 评分模块输出的结构化特征，如 A/B 分数、预测类别和 Tie 信号。"],
            ["z_bias", "偏差感知模块输出的风险特征，如 position risk、length risk、format risk。"],
            ["z_evid", "证据增强事实性模块输出的证据支撑特征。"],
            ["p(y|x)", "融合校准模块输出的最终类别概率分布。"],
            ["c(x)", "最终预测的置信度，即最大类别概率。"],
            ["r(x)", "人工复核风险信号，定义为 1 - c(x)。"],
        ],
        [2.2, 13.6],
    )
    add_formula(
        doc,
        "公式 (1)",
        "z(x) = [z_base(x); z_bias(x); z_evid(x); z_text(x)]",
        "该公式表示融合层的输入特征向量。方括号表示特征拼接；z_text 为任务相关的浅层文本特征。该表达式对应项目中 pairwise_feature_dict 与 factuality_feature_dict 对不同特征源的组合逻辑。",
    )
    add_formula(
        doc,
        "公式 (2)",
        "p(y = k | x) = softmax_k(W z(x) + b)",
        "该公式表示融合分类器的概率输出。W 与 b 为训练得到的分类参数；k 表示候选类别。pairwise 任务的类别为 A>B、B>A、Tie；factuality 任务的类别为 supported、unsupported、ambiguous。",
    )
    add_formula(
        doc,
        "公式 (3)",
        "ŷ = argmax_k p(y = k | x),     c(x) = max_k p(y = k | x)",
        "该公式定义最终预测类别和置信度。ŷ 是概率最高的类别；c(x) 是系统对该预测的置信度。后续复核策略基于该置信度计算风险。",
    )
    add_formula(
        doc,
        "公式 (4)",
        "r(x) = 1 - c(x),     review(x) = I[r(x) ≥ τ]",
        "该公式定义人工复核触发规则。τ 为开发集上选择的复核阈值；I[·] 为指示函数。该机制使系统能够把低置信度或高风险样本交给人工复核，而不是无条件接受自动判断。",
    )

    doc.add_heading("3. 模块一：基础 Judge 评分模块", level=1)
    doc.add_paragraph(
        "基础 Judge 评分模块提供系统的第一层判别信号。对于 pairwise 样本，该模块读取 prompt、context、rubric 或 reference 以及两个候选回答，调用 M-Prometheus-3B 或其 QLoRA adapter，输出 A>B、B>A 或 Tie。输出随后被解析为 base_scores.json，其中包含样本 ID、gold label、pred label、A/B 分数、原始文本输出、解析状态和后端信息。"
    )
    add_formula(
        doc,
        "公式 (5)",
        "z_base(x) = [s_A, s_B, s_A - s_B, I(ŷ_base = Tie), I(ŷ_base = A>B), I(ŷ_base = B>A)]",
        "该公式抽象表示基础 Judge 特征。s_A 与 s_B 是解析后的 A/B 侧得分；指示变量编码基础 Judge 的离散预测。实际实现中还保留 judge_backend、parse_status 等元信息用于覆盖率检查和审计。",
    )
    doc.add_paragraph(
        "学术表述上，该模块应被称为 lightweight base judge scoring module。当前 seed42 pilot 显示，QLoRA 后的基础 Judge 在 test pairwise 上 parse_success_rate = 1.0，pairwise accuracy = 0.8338，macro-F1 = 0.7370；该结果支持“参数高效微调显著改善基础 Judge 判断能力”的主张，但仍应标注为 seed42 epoch1 checkpoint pilot，而不是 3-seed 正式结论。"
    )

    doc.add_heading("4. 模块二：偏差感知模块", level=1)
    doc.add_paragraph(
        "偏差感知模块用于识别评估过程中的非语义风险因素。项目实现中，该模块生成 position、length、format、rubric sensitivity 和 source bias 等风险特征，并进一步给出 overall_bias_risk 与 review_required。该模块的功能重点是审计与复核分流，而不是直接宣称提升整体准确率。"
    )
    add_formula(
        doc,
        "公式 (6)",
        "z_bias(x) = [ρ_pos, ρ_len, ρ_fmt, ρ_rubric, ρ_src, ρ_overall, I(review_required)]",
        "该公式表示偏差风险特征。ρ_pos 表示位置偏差风险，ρ_len 表示长度偏差风险，ρ_fmt 表示格式偏差风险，ρ_rubric 表示评分标准敏感性风险，ρ_src 表示来源偏差风险。I(review_required) 表示该样本是否建议进入人工复核。"
    )
    doc.add_paragraph(
        "在论文中，该模块宜表述为 bias-aware audit module 或 review triage overlay。这样更符合当前实验事实：偏差模块的主要贡献是风险识别、偏差诊断和可解释性增强，而不是被过度表述为准确率提升模块。"
    )

    doc.add_heading("5. 模块三：证据增强事实性模块", level=1)
    doc.add_paragraph(
        "证据增强事实性模块用于 factuality assessment。该模块围绕回答与上下文证据之间的支持关系构造特征，帮助系统区分 supported、unsupported 与 ambiguous。对于 pairwise factuality 样本，该模块也可为 pairwise 决策提供证据相关特征。"
    )
    add_formula(
        doc,
        "公式 (7)",
        "z_evid(x) = [η_support, η_low_support, η_gap, η_hallucination, η_ambiguity]",
        "该公式概括证据增强特征。其中 η_support 表示证据支持强度，η_low_support 表示低支持比例，η_gap 表示证据锚定缺口，η_hallucination 表示潜在幻觉严重度，η_ambiguity 表示证据不足或语义不确定性。不同任务会启用不同子集，部分特征仅用于审计而不进入默认决策。"
    )
    add_formula(
        doc,
        "公式 (8)",
        "y_fact ∈ {supported, unsupported, ambiguous}",
        "该公式给出事实性任务的目标标签空间。supported 表示回答被证据支持；unsupported 表示回答与证据冲突或缺乏支持；ambiguous 表示证据不足、语义不明确或无法可靠判定。"
    )
    doc.add_paragraph(
        "该模块的论文贡献应强调 evidence-grounded evaluation。它使系统不完全依赖基础 Judge 的隐式推理，而是显式考虑回答与证据之间的支持关系，从而增强事实性判断的可解释性和可审计性。"
    )

    doc.add_heading("6. 模块四：融合校准与置信度输出模块", level=1)
    doc.add_paragraph(
        "融合校准模块整合基础 Judge、偏差风险、证据特征和文本特征，训练轻量分类头，并输出校准后的预测概率。该模块同时服务 pairwise 和 factuality 两个 head。除最终类别外，系统还输出 confidence、risk score、review threshold、review rate、error capture rate 和 ECE 等指标。"
    )
    add_formula(
        doc,
        "公式 (9)",
        "ECE = Σ_{m=1}^{M} (|B_m| / n) · |acc(B_m) - conf(B_m)|",
        "该公式定义 Expected Calibration Error。B_m 表示第 m 个置信度分箱；acc(B_m) 为该分箱内样本准确率；conf(B_m) 为平均置信度。ECE 越低，说明模型置信度与真实正确率越一致。"
    )
    add_formula(
        doc,
        "公式 (10)",
        "Macro-F1 = (1 / K) Σ_{k=1}^{K} F1_k",
        "该公式定义宏平均 F1。K 为类别数，F1_k 为第 k 类的 F1 值。相较 accuracy，Macro-F1 对类别不平衡更敏感，适合包含 Tie 或 ambiguous 等少数类的评估任务。"
    )
    doc.add_paragraph(
        "当前已完成的 seed42 pilot 中，QLoRA-BEA-Judge 相比 Current BEA-Judge 的 pairwise test 指标从 accuracy = 0.7512、macro-F1 = 0.6730、ECE = 0.0558、Tie recall = 0.5231，提升至 accuracy = 0.8395、macro-F1 = 0.7644、ECE = 0.0301、Tie recall = 0.5385。该结果支持融合校准层在性能和校准方面的作用，但论文中应明确这是 seed42 pilot 结果，正式均值和标准差需等待 3-seed 复现实验完成。"
    )

    doc.add_heading("7. 四模块贡献与论文表述边界", level=1)
    add_table(
        doc,
        ["模块", "主要输入", "主要输出", "论文中可主张的贡献"],
        [
            ["基础 Judge", "prompt、context、rubric、answer A/B", "base scores、A/B/Tie", "轻量 3B Judge backbone 经 QLoRA 后可显著提升原始 pairwise 判断能力。"],
            ["偏差感知", "样本元数据与回答属性", "bias risk、review_required", "提供偏差审计与人工复核优先级，不夸大为准确率提升来源。"],
            ["证据增强事实性", "context/reference 与回答文本", "evidence features、factuality signals", "增强事实性判断的证据一致性和可解释性。"],
            ["融合校准", "多源特征向量", "最终标签、置信度、ECE、复核阈值", "整合多源信号并改善校准质量与风险控制。"],
        ],
        [2.8, 4.2, 4.1, 5.3],
    )

    doc.add_heading("8. 可信性声明", level=1)
    doc.add_paragraph(
        "为保证学术真实性，本文档未声称完成新的大语言模型训练，未声称偏差模块单独提升整体准确率，未将未完成的 3-seed 复现实验写成正式结论，也未将外部 holdout 泛化能力作为已验证事实。当前可可靠陈述的结论是：seed42 epoch1 checkpoint pilot 已验证 QLoRA 可显著改善基础 Judge，且接入 BEA-Judge 四模块融合后在该内部 test split 上进一步提升 accuracy、macro-F1、ECE 和 Tie recall。正式投稿前仍需补充 3-seed mean ± std、外部 holdout 和统计显著性检验。"
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("附录 A：建议用于论文的方法段落", level=1)
    doc.add_paragraph(
        "We propose BEA-Judge, a modular and calibrated evaluation framework for lightweight LLM-as-a-judge systems. BEA-Judge consists of four components: a base judge scoring module, a bias-aware audit module, an evidence-enhanced factuality module, and a fusion calibration module with confidence estimation. The base judge module produces structured pairwise scores from a 3B Prometheus-style judge backbone. The bias-aware module estimates non-semantic risks such as position, length, format, rubric sensitivity, and source bias. The evidence-enhanced factuality module extracts evidence-grounded support signals for factuality assessment. Finally, the fusion calibration module integrates these signals and outputs calibrated predictions, confidence scores, and review indicators. This design separates raw judging ability, bias auditing, evidence grounding, and calibrated decision-making, thereby improving both predictive performance and interpretability under a lightweight parameter-efficient setting."
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
    print(DOCX_PATH)
