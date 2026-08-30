from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "paper"
FIG_DIR = OUT_DIR / "figures_bea_judge_10k"
MD_PATH = OUT_DIR / "bea_judge_sci_manuscript.md"
DOCX_PATH = OUT_DIR / "bea_judge_sci_manuscript.docx"


TITLE = "BEA-Judge-10K：面向大语言模型评价的偏差感知与证据增强校准框架"
SHORT_TITLE = "BEA-Judge-10K SCI中文论文初稿"


MAIN_RESULTS = [
    ["任务头", "划分", "样本数", "Accuracy", "Macro-F1", "ECE", "Brier", "Tie Recall", "Review Rate"],
    ["Pairwise", "Dev", "1087", "0.7443", "0.6887", "0.0401", "0.3150", "0.6558", "0.4471"],
    ["Factuality", "Dev", "491", "0.7251", "0.6976", "0.0409", "0.3469", "-", "0.5132"],
    ["Pairwise", "Test", "1053", "0.7512", "0.6730", "0.0558", "0.3048", "0.5231", "0.4131"],
    ["Factuality", "Test", "485", "0.7649", "0.7405", "0.0377", "0.3257", "-", "0.5340"],
]

CI_TABLE = [
    ["任务头", "划分", "指标", "点估计", "95% CI下界", "95% CI上界", "样本数"],
    ["Pairwise", "Test", "Accuracy", "0.7512", "0.7265", "0.7768", "1053"],
    ["Pairwise", "Test", "Macro-F1", "0.6730", "0.6412", "0.7023", "1053"],
    ["Pairwise", "Test", "ECE", "0.0558", "0.0406", "0.0796", "1053"],
    ["Pairwise", "Test", "Tie Recall", "0.5231", "0.4453", "0.6111", "1053"],
    ["Factuality", "Test", "Accuracy", "0.7649", "0.7278", "0.8021", "485"],
    ["Factuality", "Test", "Macro-F1", "0.7405", "0.6990", "0.7821", "485"],
    ["Factuality", "Test", "ECE", "0.0377", "0.0202", "0.0741", "485"],
]

ABLATION_TABLE = [
    ["模型变体", "任务头", "划分", "Accuracy", "Macro-F1", "ECE", "Brier", "Tie Recall"],
    ["Full BEA-Judge", "Pairwise", "Test", "0.7512", "0.6730", "0.0558", "0.3048", "0.5231"],
    ["Full BEA-Judge", "Factuality", "Test", "0.7649", "0.7405", "0.0377", "0.3257", "-"],
    ["w/o Bias Module", "Pairwise", "Test", "0.7654", "0.6892", "0.0385", "0.3001", "0.5462"],
    ["w/o Bias Module", "Factuality", "Test", "0.7649", "0.7405", "0.0377", "0.3257", "-"],
    ["w/o Evidence Module", "Pairwise", "Test", "0.7531", "0.6711", "0.0540", "0.3043", "0.5000"],
    ["w/o Evidence Module", "Factuality", "Test", "0.6928", "0.6542", "0.0221", "0.3962", "-"],
    ["w/o Calibration", "Pairwise", "Test", "0.7407", "0.6402", "0.0459", "0.3055", "0.3923"],
    ["w/o Calibration", "Factuality", "Test", "0.7649", "0.7405", "0.0377", "0.3257", "-"],
]

SIGNIFICANCE_TABLE = [
    ["对照变体", "任务头", "配对样本数", "ΔAccuracy", "95% CI", "ΔMacro-F1", "95% CI", "McNemar p"],
    ["w/o Bias Module", "Pairwise", "1053", "-0.0142", "[-0.0257, -0.0038]", "-0.0162", "[-0.0315, -0.0021]", "0.016674"],
    ["w/o Calibration", "Pairwise", "1053", "0.0105", "[0.0009, 0.0199]", "0.0328", "[0.0162, 0.0509]", "0.034690"],
    ["w/o Evidence Module", "Factuality", "485", "0.0721", "[0.0371, 0.1093]", "0.0863", "[0.0475, 0.1290]", "0.000224"],
]

RAGTRUTH_TABLE = [
    ["划分", "样本数", "Accuracy", "Macro-F1", "ECE", "Brier", "Review Rate", "Supported→Unsupported", "Unsupported→Supported"],
    ["Dev", "375", "0.6507", "0.5776", "0.0517", "0.4334", "0.6667", "60", "71"],
    ["Test", "372", "0.6962", "0.6363", "0.0360", "0.4130", "0.6909", "52", "61"],
]

PER_DATASET_TABLE = [
    ["任务头", "数据集", "样本数", "Accuracy", "Macro-F1", "ECE", "Brier", "Tie Recall", "Review Rate"],
    ["Pairwise", "helpsteer2", "307", "0.6678", "0.5373", "0.0540", "0.3982", "0.2625", "0.3974"],
    ["Pairwise", "judgebench", "17", "0.7059", "0.6886", "0.1636", "0.4404", "-", "1.0000"],
    ["Pairwise", "mt_bench", "72", "0.3750", "0.2072", "0.0842", "0.6857", "1.0000", "0.9861"],
    ["Pairwise", "oasst1", "119", "1.0000", "1.0000", "0.0070", "0.0001", "-", "0.0000"],
    ["Pairwise", "offsetbias", "229", "0.8690", "0.8690", "0.0752", "0.1664", "-", "0.0655"],
    ["Pairwise", "pandalm", "42", "0.0952", "0.0580", "0.2206", "0.6143", "1.0000", "1.0000"],
    ["Pairwise", "zh_professional", "157", "0.9809", "0.9661", "0.2388", "0.2013", "0.8500", "0.4268"],
    ["Factuality", "ares_nq", "113", "0.9912", "0.9911", "0.0536", "0.0383", "-", "0.0177"],
    ["Factuality", "ragtruth", "372", "0.6962", "0.6363", "0.0360", "0.4130", "-", "0.6909"],
]

DATASET_TABLE = [
    ["维度", "类别", "数量"],
    ["Task Type", "factuality_rag", "3500"],
    ["Task Type", "open_qa", "4000"],
    ["Task Type", "pairwise_bias", "2700"],
    ["Dataset", "helpsteer2", "2000"],
    ["Dataset", "oasst1", "800"],
    ["Dataset", "offsetbias", "1500"],
    ["Dataset", "ragtruth", "2500"],
    ["Dataset", "zh_professional", "1000"],
    ["Split", "train/dev/test", "7084 / 1578 / 1538"],
    ["Language", "en/zh", "9148 / 1052"],
    ["Human Label", "A>B / B>A / Tie", "4136 / 1845 / 965"],
    ["Factuality Label", "supported / unsupported", "2144 / 1110"],
]

LICENSE_TABLE = [
    ["数据源", "许可证", "是否允许重分发", "是否进入训练", "采纳记录数", "采集时间", "SHA-256"],
    ["HelpSteer2", "CC-BY-4.0", "是", "是", "2000", "2026-05-18", "有"],
    ["OASST1", "Apache-2.0", "是", "是", "800", "2026-05-18", "有"],
    ["OffsetBias", "BSD-3-Clause", "是", "是", "1500", "2026-05-18", "有"],
    ["RAGTruth", "MIT", "是", "是", "2500", "2026-05-18", "有"],
    ["RewardBench", "mixed-subset-license", "否", "否，仅外部评估规划", "0", "-", "无"],
]

BIAS_TABLE = [
    ["偏差组", "样本数", "Accuracy", "Macro-F1", "ECE", "Review Rate", "Avg Bias Risk"],
    ["format", "259", "0.7336", "0.7264", "0.1256", "0.2664", "0.5232"],
    ["length", "251", "0.8287", "0.8268", "0.1829", "0.0956", "0.2622"],
    ["none", "5746", "0.7624", "0.7125", "0.0476", "0.0935", "0.2145"],
    ["position", "293", "0.6962", "0.6962", "0.0966", "0.3038", "0.5474"],
    ["reasoning_difficulty", "153", "0.5686", "0.3748", "0.0618", "0.1046", "0.2726"],
    ["rubric_sensitivity", "244", "0.7500", "0.7469", "0.1043", "0.2500", "0.5125"],
]

EVIDENCE_ABLATION_TABLE = [
    ["特征组", "加权校准", "特征数", "Accuracy", "Macro-F1", "ECE", "Brier"],
    ["overlap-only", "否", "16", "0.7175", "0.6780", "0.0302", "0.3625"],
    ["+numeric/date/entity", "否", "24", "0.7361", "0.6900", "0.0300", "0.3485"],
    ["+sentence/local-risk", "否", "45", "0.7629", "0.6789", "0.0379", "0.3292"],
    ["+weighted calibration", "是", "46", "0.7649", "0.7405", "0.0377", "0.3257"],
]

BASELINE_COMPARISON_TABLE = [
    ["系统设置", "类别", "任务头", "划分", "n", "Accuracy", "Macro-F1", "ECE", "Brier", "Tie Recall", "Review Rate"],
    ["Full BEA-Judge", "module_variant", "pairwise", "test", "1053", "0.7512", "0.6730", "0.0558", "0.3048", "0.5231", "0.4131"],
    ["Full BEA-Judge", "module_variant", "factuality", "test", "485", "0.7649", "0.7405", "0.0377", "0.3257", "-", "0.5340"],
    ["w/o Bias Module", "module_variant", "pairwise", "test", "1053", "0.7654", "0.6892", "0.0385", "0.3001", "0.5462", "0.4017"],
    ["w/o Evidence Module", "module_variant", "factuality", "test", "485", "0.6928", "0.6542", "0.0221", "0.3962", "-", "0.5959"],
    ["w/o Calibration", "module_variant", "pairwise", "test", "1053", "0.7407", "0.6402", "0.0459", "0.3055", "0.3923", "0.3485"],
    ["w/o Base Judge Scores", "module_variant", "pairwise", "test", "1053", "0.6942", "0.6184", "0.0325", "0.3388", "0.5308", "0.5508"],
    ["w/o Tie Policy", "module_variant", "pairwise", "test", "1053", "0.7664", "0.6300", "0.0410", "0.3048", "0.2308", "0.4131"],
    ["w/o Review Threshold", "module_variant", "pairwise", "test", "1053", "0.7512", "0.6730", "0.0558", "0.3048", "0.5231", "0.0000"],
    ["Raw M-Prometheus-3B only", "control", "pairwise", "test", "1053", "0.5632", "0.4079", "0.4368", "0.8737", "0.0692", "0.0000"],
    ["Text/metadata-only", "control", "pairwise", "test", "1053", "0.5489", "0.5424", "0.1118", "0.4083", "0.8846", "0.5859"],
    ["Base + fusion calibration only", "control", "pairwise", "test", "1053", "0.6135", "0.5757", "0.0819", "0.3747", "0.6462", "0.4473"],
]

CALIBRATION_METHODS_TABLE = [
    ["方法", "划分", "Accuracy", "ECE", "MCE", "Brier", "NLL", "Coverage", "Avg Set Size"],
    ["conformal", "dev", "0.7553", "0.0414", "0.1135", "0.3427", "0.5477", "0.9006", "-"],
    ["conformal", "test", "0.7664", "0.0484", "0.1521", "0.3175", "0.5041", "0.9174", "1.4131"],
    ["isotonic", "dev", "0.7626", "0.0213", "0.3891", "0.3195", "0.5034", "-", "-"],
    ["isotonic", "test", "0.7692", "0.0263", "0.3926", "0.3083", "0.5102", "-", "-"],
    ["platt", "dev", "0.7489", "0.0205", "0.2404", "0.3380", "0.5489", "-", "-"],
    ["platt", "test", "0.7616", "0.0194", "0.1807", "0.3123", "0.5026", "-", "-"],
    ["temperature", "dev", "0.7553", "0.0383", "0.1304", "0.3439", "0.5434", "-", "-"],
    ["temperature", "test", "0.7664", "0.0414", "0.1138", "0.3204", "0.5064", "-", "-"],
    ["vector_scaling", "dev", "0.7544", "0.0395", "0.1228", "0.3425", "0.5475", "-", "-"],
    ["vector_scaling", "test", "0.7654", "0.0513", "0.1886", "0.3174", "0.5039", "-", "-"],
]

BIAS_RISK_UTILITY_TABLE = [
    ["设置", "任务头", "划分", "n", "Accuracy", "Macro-F1", "ECE", "Review Rate", "Review Capture Rate"],
    ["bias_as_decision_features", "pairwise", "test", "1053", "0.7512", "0.6730", "0.0558", "0.4131", "0.7557"],
    ["no_bias_decision_features", "pairwise", "test", "1053", "0.7654", "0.6892", "0.0385", "0.4017", "0.7328"],
    ["bias_risk_only_review", "pairwise", "test", "1053", "0.7654", "0.6892", "0.0385", "0.1595", "0.1498"],
]

RISK_COVERAGE_TABLE = [
    ["任务头", "划分", "Review Rate", "Review Count", "Error Capture Rate", "Auto Accept Count", "Auto Accept Accuracy", "Risk Threshold"],
    ["pairwise", "test", "0.2004", "211", "0.5000", "842", "0.8444", "0.474207"],
    ["pairwise", "test", "0.3998", "421", "0.7405", "632", "0.8924", "0.351688"],
    ["pairwise", "test", "0.4995", "526", "0.8511", "527", "0.9260", "0.215020"],
    ["factuality", "test", "0.2000", "97", "0.3509", "388", "0.8093", "0.425574"],
    ["factuality", "test", "0.4000", "194", "0.6491", "291", "0.8625", "0.349946"],
    ["factuality", "test", "0.4990", "242", "0.7544", "243", "0.8848", "0.301048"],
]

REFERENCES = [
    "Zheng L., Chiang W.-L., Sheng Y., et al. Judging LLM-as-a-Judge with MT-Bench and Chatbot Arena. arXiv:2306.05685, 2023. https://arxiv.org/abs/2306.05685",
    "Liu Y., Iter D., Xu Y., et al. G-Eval: NLG Evaluation using GPT-4 with Better Human Alignment. arXiv:2303.16634, 2023. https://arxiv.org/abs/2303.16634",
    "Kim S., Suk J., Longpre S., et al. Prometheus 2: An Open Source Language Model Specialized in Evaluating Other Language Models. arXiv:2405.01535, 2024. https://arxiv.org/abs/2405.01535",
    "Pombal J., Cortez A., Cardon R., et al. M-Prometheus: A Suite of Open Multilingual LLM Judges. arXiv:2504.04953, 2025. https://arxiv.org/abs/2504.04953",
    "Niu C., Wu J., Zhang Y., et al. RAGTruth: A Hallucination Corpus for Developing Trustworthy Retrieval-Augmented Language Models. arXiv:2401.00396, 2024. https://arxiv.org/abs/2401.00396",
    "Wang Z., Dong Y., Zeng O., et al. HelpSteer2: Open-source Dataset for Training Top-performing Reward Models. arXiv:2406.08673, 2024. https://arxiv.org/abs/2406.08673",
    "Köpf A., Kilcher Y., von Rütte D., et al. OpenAssistant Conversations: Democratising Large Language Model Alignment. arXiv:2304.07327, 2023. https://arxiv.org/abs/2304.07327",
    "Guo C., Pleiss G., Sun Y., Weinberger K. Q. On Calibration of Modern Neural Networks. Proceedings of ICML, 2017. https://proceedings.mlr.press/v70/guo17a.html",
    "Liang P., Bommasani R., Lee T., et al. Holistic Evaluation of Language Models. Transactions on Machine Learning Research, 2023. https://arxiv.org/abs/2211.09110",
    "Bai Y., Kadavath S., Kundu S., et al. Constitutional AI: Harmlessness from AI Feedback. arXiv:2212.08073, 2022. https://arxiv.org/abs/2212.08073",
    "Huang H., Qu Y., Liu J., et al. A Survey on LLM-as-a-Judge. arXiv:2411.15594, 2024. https://arxiv.org/abs/2411.15594",
    "Ouyang L., Wu J., Jiang X., et al. Training language models to follow instructions with human feedback. Advances in Neural Information Processing Systems, 2022. https://arxiv.org/abs/2203.02155",
    "Christiano P. F., Leike J., Brown T., et al. Deep reinforcement learning from human preferences. Advances in Neural Information Processing Systems, 2017. https://arxiv.org/abs/1706.03741",
    "Rafailov R., Sharma A., Mitchell E., et al. Direct Preference Optimization: Your Language Model is Secretly a Reward Model. Advances in Neural Information Processing Systems, 2023. https://arxiv.org/abs/2305.18290",
    "Stiennon N., Ouyang L., Wu J., et al. Learning to summarize with human feedback. Advances in Neural Information Processing Systems, 2020. https://arxiv.org/abs/2009.01325",
    "Thoppilan R., De Freitas D., Hall J., et al. LaMDA: Language Models for Dialog Applications. arXiv:2201.08239, 2022. https://arxiv.org/abs/2201.08239",
    "OpenAI. GPT-4 Technical Report. arXiv:2303.08774, 2023. https://arxiv.org/abs/2303.08774",
    "Brown T. B., Mann B., Ryder N., et al. Language Models are Few-Shot Learners. Advances in Neural Information Processing Systems, 2020. https://arxiv.org/abs/2005.14165",
    "Lin S., Hilton J., Evans O. TruthfulQA: Measuring How Models Mimic Human Falsehoods. Proceedings of ACL, 2022. https://arxiv.org/abs/2109.07958",
    "Manakul P., Liusie A., Gales M. J. F. SelfCheckGPT: Zero-Resource Black-Box Hallucination Detection for Generative Large Language Models. Proceedings of EMNLP, 2023. https://arxiv.org/abs/2303.08896",
    "Min S., Krishna K., Lyu X., et al. FActScore: Fine-grained Atomic Evaluation of Factual Precision in Long Form Text Generation. Proceedings of EMNLP, 2023. https://arxiv.org/abs/2305.14251",
    "Li J., Cheng X., Zhao W. X., Nie J.-Y., Wen J.-R. HaluEval: A Large-Scale Hallucination Evaluation Benchmark for Large Language Models. Proceedings of EMNLP, 2023. https://arxiv.org/abs/2305.11747",
    "Thorne J., Vlachos A., Christodoulopoulos C., Mittal A. FEVER: a Large-scale Dataset for Fact Extraction and VERification. Proceedings of NAACL-HLT, 2018. https://aclanthology.org/N18-1074/",
    "Honovich O., Scialom T., Levy O., Schick T. QAFactEval: Improved QA-Based Factual Consistency Evaluation for Summarization. Proceedings of NAACL, 2022. https://aclanthology.org/2022.naacl-main.187/",
    "Fabbri A. R., Kryściński W., McCann B., et al. SummEval: Re-evaluating Summarization Evaluation. Transactions of the Association for Computational Linguistics, 2021. https://aclanthology.org/2021.tacl-1.24/",
    "Zhang T., Kishore V., Wu F., Weinberger K. Q., Artzi Y. BERTScore: Evaluating Text Generation with BERT. International Conference on Learning Representations, 2020. https://arxiv.org/abs/1904.09675",
    "Sellam T., Das D., Parikh A. P. BLEURT: Learning Robust Metrics for Text Generation. Proceedings of ACL, 2020. https://aclanthology.org/2020.acl-main.704/",
    "Rei R., Stewart C., Farinha A. C., Lavie A. COMET: A Neural Framework for MT Evaluation. Proceedings of EMNLP, 2020. https://aclanthology.org/2020.emnlp-main.213/",
    "Pillutla K., Swayamdipta S., Zellers R., et al. MAUVE: Measuring the Gap Between Neural Text and Human Text using Divergence Frontiers. Advances in Neural Information Processing Systems, 2021. https://arxiv.org/abs/2102.01454",
    "Yuan W., Neubig G., Liu P. BARTScore: Evaluating Generated Text as Text Generation. Advances in Neural Information Processing Systems, 2021. https://arxiv.org/abs/2106.11520",
    "Papineni K., Roukos S., Ward T., Zhu W.-J. BLEU: a Method for Automatic Evaluation of Machine Translation. Proceedings of ACL, 2002. https://aclanthology.org/P02-1040/",
    "Banerjee S., Lavie A. METEOR: An Automatic Metric for MT Evaluation with Improved Correlation with Human Judgments. ACL Workshop on Intrinsic and Extrinsic Evaluation Measures, 2005. https://aclanthology.org/W05-0909/",
    "Lin C.-Y. ROUGE: A Package for Automatic Evaluation of Summaries. ACL Workshop on Text Summarization Branches Out, 2004. https://aclanthology.org/W04-1013/",
    "Naeini M. P., Cooper G. F., Hauskrecht M. Obtaining Well Calibrated Probabilities Using Bayesian Binning. Proceedings of AAAI, 2015. https://ojs.aaai.org/index.php/AAAI/article/view/9602",
    "Efron B. Bootstrap Methods: Another Look at the Jackknife. The Annals of Statistics, 1979. https://doi.org/10.1214/aos/1176344552",
    "McNemar Q. Note on the sampling error of the difference between correlated proportions or percentages. Psychometrika, 1947. https://doi.org/10.1007/BF02295996",
    "Brier G. W. Verification of forecasts expressed in terms of probability. Monthly Weather Review, 1950. https://doi.org/10.1175/1520-0493(1950)078%3C0001:VOFEIT%3E2.0.CO;2",
    "Kadavath S., Conerly T., Askell A., et al. Language Models (Mostly) Know What They Know. arXiv:2207.05221, 2022. https://arxiv.org/abs/2207.05221",
    "Jiang Z., Araki J., Ding H., Neubig G. How Can We Know When Language Models Know? On the Calibration of Language Models for Question Answering. Transactions of the Association for Computational Linguistics, 2021. https://aclanthology.org/2021.tacl-1.57/",
]


SECTIONS = [
    (
        1,
        "摘要",
        [
            "大语言模型评价正在从人工评分辅助工具转向开放式模型能力比较、偏好学习数据筛选和检索增强生成质量控制的核心基础设施。然而，直接使用单一 Judge 模型进行结论判定仍存在三个主要问题：评分偏差难以解释，事实性错误缺少证据支撑，模型输出置信度与真实错误风险之间存在校准偏差。",
            "针对上述问题，本文提出 BEA-Judge-10K，一个面向 LLM-as-a-Judge 场景的偏差感知与证据增强校准框架。该框架不微调新的大语言模型，而是在真实 M-Prometheus-3B 基础评分之上，构建偏差画像、证据特征、双任务 softmax 校准和置信度输出层。",
            "在数据层面，本文构建了 10200 条样本的 BEA-Judge-10K v2 数据包，覆盖 pairwise preference、pairwise bias 和 RAG factuality 三类任务。数据集采用 license-first 与 provenance-tracked 构建策略，记录数据来源、许可、采集时间、SHA-256 与划分信息，避免将许可不明或重分发受限的数据混入训练集。",
            "在正式实验中，M-Prometheus-3B 对 6946 条 pairwise 样本实现 6946/6946 全覆盖，启发式 fallback 行数为 0。测试集上，pairwise head 获得 0.7512 accuracy、0.6730 macro-F1 和 0.5231 Tie recall；factuality head 获得 0.7649 accuracy、0.7405 macro-F1 和 0.0377 ECE。",
            "消融实验表明，证据增强模块是事实性任务的主要性能贡献：移除该模块后 factuality macro-F1 从 0.7405 降至 0.6542，配对检验 McNemar p=0.000224。融合校准模块主要改善 pairwise 决策稳定性：移除校准后 pairwise macro-F1 从 0.6730 降至 0.6402，Tie recall 从 0.5231 降至 0.3923，McNemar p=0.034690。",
            "进一步分析显示，RAGTruth 作为 response-level hallucination 泛化集仍是主要挑战，其测试集 macro-F1 为 0.6363，错误主要集中在 supported 与 unsupported 的双向混淆。本文据此将 BEA-Judge 定位为可解释、可复现、轻量级的 Judge 增强与校准框架，而非端到端重训练的大模型评价器。",
        ],
    ),
    (
        1,
        "关键词",
        [
            "大语言模型评价；LLM-as-a-Judge；证据增强；偏差感知；置信度校准；RAG factuality；Tie policy",
        ],
    ),
    (
        1,
        "1 引言",
        [
            "随着大语言模型在开放问答、代码生成、检索增强生成和多轮对话中的广泛应用，模型输出质量的评价已经成为人工智能研究与工程部署中的关键问题。传统基于精确匹配或 n-gram overlap 的自动指标难以覆盖开放式生成结果中的语义质量、回答偏好和事实一致性，因此 LLM-as-a-Judge 方法被广泛用于模型比较和数据筛选[1-2]。",
            "尽管大模型 Judge 能在许多开放任务上获得与人类偏好较高的一致性，但其输出并不天然具备可解释性与可靠校准。现有研究已经指出，Judge 模型可能受到选项位置、回答长度、格式规范、评价 rubric 和任务难度的影响。在 pairwise preference 场景中，这类偏差会改变胜负判断；在事实性评估中，缺少显式证据建模会导致 Judge 将流畅但不被上下文支持的回答误判为可靠。",
            "另一个被低估的问题是置信度校准。Judge 模型通常给出离散评分或文本理由，但研究使用者最终需要的是可比较、可复核、可阈值化的决策结果。未校准的模型输出可能在总体 accuracy 看似可接受时，对 Tie、低置信样本或困难数据集产生系统性风险。",
            "本文研究问题可以概括为：在不微调大语言模型、不引入外部检索和不使用启发式 fallback 的前提下，能否基于开放 Judge 的真实输出构建一个可解释、可复现且面向 SCI 实验叙述的校准框架？",
            "为回答这一问题，本文提出 BEA-Judge-10K。该框架固定使用 M-Prometheus-3B 作为基础 Judge backbone，在其输出上叠加偏差感知模块、证据增强事实性模块以及融合校准与置信度输出模块。方法目标不是替代基础 Judge，而是将基础评分转化为更可解释、更适合审计、更便于人工复核的结构化结果。",
            "本文贡献包括三点。第一，构建并审计 BEA-Judge-10K v2 数据包，覆盖 10200 条样本、6946 条 pairwise 样本和 3500 条 evidence profile，并提供许可、来源和划分审计。第二，提出 deterministic evidence proxy 与 factuality calibration 组合，显著改善事实性任务的 macro-F1。第三，提出双任务融合校准与 Tie policy，使 pairwise 决策在 Tie recall 和宏平均性能上优于未校准基线。",
            "需要强调的是，本文不将自身表述为重新训练的 Judge 大模型，也不声称证据模块完成了原子事实级验证。BEA-Judge 的价值在于通过轻量、透明和可复现的后处理机制，提高开放 Judge 输出的可解释性、事实性可靠性和复核优先级分配能力。",
        ],
    ),
    (
        1,
        "2 相关工作",
        [
            "LLM-as-a-Judge 已成为开放式生成评价的重要范式。MT-Bench 与 Chatbot Arena 显示，基于大模型的偏好判断可以在多轮对话模型比较中提供高效评价信号[1]。G-Eval 进一步利用 GPT 系列模型与 chain-of-thought 风格评价提示改善自然语言生成评价与人类判断的一致性[2]。HELM 则从模型能力、鲁棒性、公平性和效率等维度提出更系统的评价框架[9]。",
            "开放 Judge 模型的发展降低了评价过程对闭源 API 的依赖。Prometheus 系列模型通过专门的评价数据与 rubric 学习，为研究者提供了可部署的评价基础模型[3]。M-Prometheus 进一步扩展了多语言评价能力，为本文选择其 3B 版本作为基础 Judge 提供了直接动机[4]。LLM-as-a-Judge 综述研究也指出，该方向仍面临位置偏差、长度偏差、评分尺度不稳定、任务迁移脆弱和校准不足等问题[11]。",
            "偏好学习和奖励建模为本文的 pairwise 设置提供了方法背景。RLHF、InstructGPT、人类反馈摘要学习、Constitutional AI 和 Direct Preference Optimization 表明，人类偏好或 AI 反馈可以显著影响模型行为[10,12-15]。但这些方法通常关注训练或对齐阶段，而本文关注固定 Judge 输出之后的可解释校准与复核。",
            "在数据层面，HelpSteer2 提供了用于 reward model 训练的开放偏好与质量标注资源[6]，OASST1 提供开放对话与人类协作数据[7]。GPT-3、LaMDA 和 GPT-4 等大模型研究说明开放式生成任务具有语义多样性和评价主观性[16-18]，这也是 pairwise Tie 与低置信复核不可忽略的原因。",
            "自动文本评价长期依赖 BLEU、ROUGE、METEOR 等词面指标[31-33]，随后出现 BERTScore、BLEURT、COMET、MAUVE 和 BARTScore 等语义或模型化指标[26-30]。这些工作改善了生成文本评价，但仍难以直接处理开放式偏好、Judge 偏差和事实性证据缺口。",
            "事实性与 hallucination 评价方面，TruthfulQA、SelfCheckGPT、FActScore、HaluEval、FEVER、QAFactEval 和 SummEval 分别从问答真实性、黑箱自一致性、原子事实、幻觉基准、事实验证和摘要一致性等角度推进了相关研究[19-25]。RAGTruth 则提供了面向检索增强生成的 response-level hallucination 语料[5]，其标签粒度高于简单句级 support 判断。",
            "校准研究表明，神经网络预测概率往往无法直接代表真实正确率，temperature scaling、Bayesian binning、ECE、Brier score 和语言模型问答校准等研究为本文的置信度建模提供了方法依据[8,34,37-39]。统计可信度方面，bootstrap、McNemar 检验和概率预测评分为本文的置信区间与消融显著性分析提供了基础[35-37]。",
            "与已有方法相比，本文不依赖闭源模型二次打分，也不引入外部检索或 LLM claim extraction。本文聚焦于一个更保守但可复现的设定：在固定基础 Judge 输出的前提下，通过偏差、证据和校准模块提升结果的可解释性和实验可信度。",
        ],
    ),
    (
        1,
        "3 数据集构建与审计",
        [
            "BEA-Judge-10K v2 的构建遵循 license-first、schema-stable 和 model-compatible 三项原则。所有新增数据源在进入训练、开发和测试划分前均进行许可证与来源审计；字段转换保持与既有 canonical schema 兼容；新增样本不会改变四模块模型的输入输出接口。",
            "最终数据集包含 10200 条样本，训练集、开发集和测试集分别为 7084、1578 和 1538 条。任务类型包括 3500 条 factuality_rag、4000 条 open_qa 和 2700 条 pairwise_bias。语言分布以英文为主，同时保留 1052 条中文样本以评估跨语言场景下的稳定性。",
            "Pairwise 标签包括 A>B、B>A 和 Tie，其中 human label distribution 分别为 4136、1845 和 965。事实性标签包括 supported 与 unsupported，数量分别为 2144 和 1110。该分布使模型既能学习明确胜负判断，也能学习对平局与低置信样本的保守处理。",
            "数据来源包括 HelpSteer2、OASST1、OffsetBias、RAGTruth 以及项目内既有正式样本。RewardBench 因混合子集许可和重分发风险，在当前版本中仅保留为 external evaluation 规划，不进入正式 train/dev/test 划分。",
            "每个训练准入数据源均记录 URL、license、acquisition date、revision/ref、record count 与 SHA-256。license audit 显示，HelpSteer2、OASST1、OffsetBias 与 RAGTruth 均满足当前训练准入条件；RewardBench 被标记为 external_eval_only 与 redistribution_restricted。",
            "数据清洗阶段执行 duplicate ID 检查、cross-split duplicate content 检查、required fields 非空检查和 factuality context 缺失检查。该策略的目的不是追求最大样本数，而是在接近 10K 规模的同时保持可追踪、可复现和可投稿审计。",
        ],
    ),
    (
        1,
        "4 方法",
        [
            "BEA-Judge 的整体链路为：BEA-Judge-10K v2 → M-Prometheus-3B base scores → bias/evidence features → dual-head softmax calibration → calibrated_results → SCI tables。该链路将基础 Judge 输出视为一个强但未校准的中间表征，再由轻量特征与校准层生成最终预测、置信度和复核标记。",
            "为便于复现，设第 i 个样本为 x_i=(q_i,a_i,b_i,c_i,r_i)，其中 q_i 为问题或指令，a_i 与 b_i 为候选回答，c_i 为上下文或检索证据，r_i 为评价 rubric 或任务说明。该定义覆盖 pairwise 与 factuality 两类任务。",
            "式(1)：x_i=(q_i,a_i,b_i,c_i,r_i)。",
            "基础 Judge 对样本 x_i 输出评分向量 s_i=[s_A,s_B,s_T]，分别对应 A>B、B>A 与 Tie 的基础偏好信号。对于事实性任务，基础 pairwise 输出不直接替代 factuality 标签，而作为可选的质量与冲突特征。",
            "式(2)：s_i=[s_A,s_B,s_T]。",
            "基础评分差距用于刻画 Judge 对 A/B 胜负的确信程度。margin 越小，样本越可能进入 Tie policy 或人工复核路径。",
            "式(3)：m_i=|s_A-s_B|。",
            "基础 Judge 评分模块固定使用 M-Prometheus-3B。正式运行中，系统对 6946 条 pairwise 样本生成真实基础评分，初始 154 条解析失败记录通过 retry repair 修复，最终 coverage 达到 6946/6946，backend error 行数与 heuristic fallback 行数均为 0。",
            "偏差感知模块构建 position、length、format、rubric_sensitivity 与 reasoning_difficulty 风险画像。该模块在训练中作为数值特征输入，在输出中用于 review_flag 和 review_reason。根据消融结果，偏差模块不被解释为主 accuracy 提升来源，而被定位为风险识别和人工复核优先级机制。",
            "偏差风险定义为多个可解释风险因子的加权组合，并裁剪到 [0,1] 区间。该定义使偏差模块成为显式审计变量，而不是隐式改变最终标签的黑箱规则。",
            "式(4)：r_i^bias=clip(Σ_j α_j r_{ij},0,1), j∈{pos,len,fmt,rubric,reason}。",
            "证据增强事实性模块使用 deterministic proxy v1。特征包括 context/reference overlap、句级最大支持度、低支持句比例、numeric/date/entity gap、negation mismatch、comparative mismatch、entity alias gap、pairwise support delta 和 local hallucination risk。这些特征均保持为范围在 [0,1] 的浮点数。",
            "句级证据支持度使用回答句子 u 与上下文句子 v 的 token overlap 衡量。该指标并不等价于语义蕴含，但能以可复现方式捕获显式证据覆盖程度。",
            "式(5)：S(u,v)=|tok(u)∩tok(v)| / max(|tok(u)|,1)。",
            "对于回答中的每个句子 u，最大支持度定义为其与所有上下文句子的最大 overlap。若最大支持度低于阈值 τ_s，则该句被视为低支持句。",
            "式(6)：S_max(u)=max_{v∈C_i} S(u,v)。",
            "低支持句比例刻画回答中潜在 unsupported 内容的密度。该特征对 RAGTruth 这类 response-level hallucination 标签尤其重要。",
            "式(7)：ρ_i^low=|{u∈A_i:S_max(u)<τ_s}| / max(|A_i|,1)。",
            "实体、数字与时间缺口分别衡量回答中的关键 token 是否缺失于 context/reference。局部幻觉风险取低支持句比例与这些缺口特征的最大值。",
            "式(8)：r_i^evidence=max(ρ_i^low,g_i^num,g_i^date,g_i^entity,g_i^neg,g_i^comp)。",
            "该证据模块并不声称完成原子事实验证。它的作用是为 factuality head 提供可解释的上下文支持信号，尤其在 RAGTruth 这类 response-level hallucination 任务中，帮助模型区分整体语义流畅但局部证据不足的回答。",
            "融合校准与置信度输出模块采用双 head 结构。Pairwise head 输出 A>B、B>A 和 Tie；factuality head 输出 supported、unsupported 与 ambiguous。特征输入包括基础 Judge 分数、score gap、pred label indicators、文本长度/重叠特征、bias risk、evidence risk 以及 dataset/task one-hot。",
            "融合层将基础评分、偏差风险、证据风险和数据集特征拼接为特征向量 φ_i。该向量进入 pairwise 与 factuality 两个任务头，各自学习轻量 softmax 分类器。",
            "式(9)：φ_i=[s_i,m_i,r_i^bias,r_i^evidence,text_i,overlap_i,onehot(dataset_i),onehot(task_i)]。",
            "双任务校准使用 temperature scaling。T 仅由 dev 集选择，test 集不参与温度或阈值搜索。",
            "式(10)：p_i=softmax((W_h φ_i+b_h)/T_h), h∈{pairwise,factuality}。",
            "校准过程仅使用 dev 集选择 temperature、review threshold、class/source weight 和 Tie policy。Test 集仅用于最终评估，不参与任何阈值或权重选择。Pairwise Tie policy 基于 P(Tie)、A/B 概率差距和基础评分 margin 选择，目标是在不显著损害 macro-F1 的前提下提高 Tie recall。",
            "Tie policy 被定义为一个 dev-only 决策规则。当 P(Tie) 足够高、A/B 概率差足够小且基础 margin 满足阈值时，输出 Tie；否则输出概率最大的非 Tie 标签。",
            "式(11)：ŷ_i=Tie, if p_i(Tie)≥τ_T and |p_i(A>B)-p_i(B>A)|≤τ_M and m_i≤τ_B。",
            "置信度定义为最大类别概率，风险分数由不确定性、偏差风险和证据风险共同决定。该分数用于 review_flag，而不是替换模型预测。",
            "式(12)：conf_i=max_k p_{ik}; risk_i=clip(1-conf_i+λ_b r_i^bias+λ_e r_i^evidence,0,1)。",
            "输出接口保持为 calibrated_results.json，其中包含 final_score、predicted_label、confidence、risk_score、review_flag、review_reason 和 label_probabilities。该固定 schema 保证后续审计表和论文结果表可以稳定复用。",
        ],
    ),
    (
        1,
        "5 实验设置",
        [
            "正式 baseline 固定为 bea_judge_20260521_110114。所有实验均使用 datasets/processed/bea_judge_cleaned_10000.json 与 datasets/judge_outputs/m_prometheus_3b_bea10k_v2/base_scores.repaired.json。模型输出存储于 datasets/model_outputs/bea_judge_20260521_110114。",
            "本文评估两个任务头。Pairwise head 主要报告 accuracy、macro-F1、ECE、Brier、Tie recall 和 review rate。Factuality head 主要报告 accuracy、macro-F1、ECE、Brier 和 review rate。由于类别分布与 Tie 判断均具有实际意义，macro-F1 与 Tie recall 在解释中优先级高于单一 accuracy。",
            "统计可信度通过 bootstrap 95% confidence interval、paired bootstrap delta 和 McNemar 检验补强。消融实验包括 Full BEA-Judge、w/o Bias Module、w/o Evidence Module 与 w/o Calibration 四组。",
            "Accuracy 衡量预测标签与 gold label 的总体一致性，但在类别不平衡和 Tie 样本较少的场景下容易掩盖少数类问题。",
            "式(13)：Accuracy=(1/N)Σ_i 1[ŷ_i=y_i]。",
            "Macro-F1 对每个类别计算 F1 后取平均，因此比 accuracy 更能反映 Tie、unsupported 等关键类别的表现。",
            "式(14)：Macro-F1=(1/K)Σ_{k=1}^K 2·Precision_k·Recall_k/(Precision_k+Recall_k)。",
            "ECE 将样本按置信度划分为 B 个 bin，衡量平均置信度与真实准确率之间的差距。该指标用于判断置信度是否可作为复核阈值依据。",
            "式(15)：ECE=Σ_{b=1}^B (|B_b|/N)|acc(B_b)-conf(B_b)|。",
            "Brier score 衡量概率分布与 one-hot gold label 的平方误差，数值越低表示概率预测越可靠。",
            "式(16)：Brier=(1/N)Σ_iΣ_k(p_{ik}-1[y_i=k])^2。",
            "Tie recall 单独衡量 gold Tie 样本被正确预测为 Tie 的比例，是 pairwise head 的关键稳健性指标。",
            "式(17)：TieRecall=TP_Tie/(TP_Tie+FN_Tie)。",
            "Bootstrap 置信区间通过对测试样本进行有放回重采样获得。本文报告 2.5% 和 97.5% 分位数作为 95% CI。",
            "式(18)：CI_95=[quantile_{0.025}({θ_b^*}), quantile_{0.975}({θ_b^*})]。",
            "McNemar 检验用于比较 Full 与消融模型在同一测试样本上的配对错误差异。n_01 表示 Full 错而对照正确，n_10 表示 Full 正确而对照错误。",
            "式(19)：χ²=(|n_01-n_10|-1)^2/(n_01+n_10)。",
            "本文将 RAGTruth 单独作为关键泛化集报告。其原因是 RAGTruth 的标签更接近 response-level hallucination，而不是简单句级证据支持；因此其错误分析对理解事实性模块的边界更重要。",
            "所有实验均不微调 M-Prometheus-3B，不使用 heuristic fallback 进入正式结果，不接入外部检索，也不使用 LLM claim extraction。该约束保证本文贡献集中在可解释增强与校准，而不是通过额外大模型推理掩盖方法边界。",
        ],
    ),
    (
        1,
        "6 实验结果",
        [
            "主结果显示，BEA-Judge 在 pairwise test 上获得 0.7512 accuracy、0.6730 macro-F1、0.0558 ECE、0.3048 Brier 和 0.5231 Tie recall。在 factuality test 上获得 0.7649 accuracy、0.7405 macro-F1、0.0377 ECE 和 0.3257 Brier。",
            "从式(10)和式(12)看，本文的改进并非来自改变基础 Judge 的原始判断，而是来自对特征向量 φ_i 的概率校准和风险建模。因此，结果解释应关注每个模块如何改变类别概率、置信度与复核风险，而不是将其理解为新的大模型推理能力。",
            "与 Raw M-Prometheus-3B only 对照相比，Full BEA-Judge 在 pairwise test 上将 macro-F1 从 0.4079 提升至 0.6730，将 Tie recall 从 0.0692 提升至 0.5231，并将 ECE 从 0.4368 降至 0.0558。该对照说明本文的贡献主要体现在基础 Judge 输出的结构化校准、Tie 决策和风险输出，而不是训练新的 Judge backbone。",
            "从置信区间看，pairwise test macro-F1 的 95% CI 为 [0.6412, 0.7023]，Tie recall 的 95% CI 为 [0.4453, 0.6111]。Factuality test macro-F1 的 95% CI 为 [0.6990, 0.7821]，ECE 的 95% CI 为 [0.0202, 0.0741]。这些区间为论文中讨论结果稳定性提供了边界。",
            "消融结果显示，证据增强模块对 factuality 的贡献最明确。移除 evidence module 后，factuality accuracy 从 0.7649 降至 0.6928，macro-F1 从 0.7405 降至 0.6542。配对显著性检验显示 Δmacro-F1=0.0863，95% CI 为 [0.0475, 0.1290]，McNemar p=0.000224。",
            "证据模块主要影响 factuality 而非 pairwise 的原因可由式(5)至式(8)解释：这些特征直接刻画上下文支持、低支持句比例和实体/数字/日期缺口，与 supported/unsupported 判断更接近，而与开放偏好任务中的写作质量、风格和偏好差异并不完全同构。",
            "融合校准模块对 pairwise 决策稳定性贡献明显。移除 calibration 后，pairwise macro-F1 从 0.6730 降至 0.6402，Tie recall 从 0.5231 降至 0.3923。对应 Δmacro-F1=0.0328，95% CI 为 [0.0162, 0.0509]，McNemar p=0.034690。",
            "Tie recall 的提升来自式(11)定义的 dev-only Tie policy。该规则将 P(Tie)、A/B 概率差和基础 margin 联合考虑，避免未校准模型在低差距样本上过早输出 A>B 或 B>A。",
            "校准方法对照显示，Platt 与 isotonic 在测试集 ECE 上可分别达到 0.0194 和 0.0263，conformal 输出可提供 0.9174 的覆盖率与 1.4131 的平均预测集合大小。这些结果可作为附加诊断和附录对照，但不替换 frozen baseline 中的 dual-head softmax 与 dev-only Tie policy。",
            "偏差模块的消融结果需要谨慎解释。w/o Bias 在 pairwise test 上的 accuracy 和 macro-F1 略高于 Full，因此不能将 Bias 模块写成整体准确率提升模块。但 Bias 模块提供了 position、length、format、rubric_sensitivity 与 reasoning_difficulty 的风险画像，可用于解释错误来源和安排人工复核优先级。",
            "从式(4)和式(12)看，bias risk 的主要作用是提高高风险样本的 review probability，而不是直接改变 argmax 标签。这解释了为何 Bias 模块在总体 accuracy 上不一定占优，却仍可作为风险控制与审计模块保留。",
            "分数据集结果显示，pairwise 难点集中在 mt_bench 与 pandalm。mt_bench 的 accuracy 为 0.3750，macro-F1 为 0.2072；pandalm 的 accuracy 为 0.0952，macro-F1 为 0.0580。这说明当前 Tie policy 虽然提升了总体 Tie recall，但在部分开放偏好数据集上可能过度保守，需要后续通过 dataset-aware policy 或 order-swap probe 分析。",
            "事实性任务呈现明显的数据集差异。ARES-NQ 上 accuracy 和 macro-F1 均约为 0.991，而 RAGTruth 上 accuracy 为 0.6962、macro-F1 为 0.6363。这一差异支持本文对 RAGTruth 难度的解释：response-level hallucination 判断比局部 support 判断更依赖细粒度证据与标签语义一致性。",
            "证据特征组消融进一步说明，简单 overlap-only 特征的 factuality macro-F1 为 0.6780；加入 numeric/date/entity 后提升至 0.6900；加入 sentence/local-risk 后 accuracy 提升至 0.7629，但 macro-F1 未同步提升；最终通过 weighted calibration 使 macro-F1 达到 0.7405。该结果表明证据特征与校准策略需要联合设计。",
        ],
    ),
    (
        1,
        "7 RAGTruth 难点分析",
        [
            "RAGTruth 是本文事实性模块的主要压力测试。其测试集 n=372，accuracy=0.6962，macro-F1=0.6363，ECE=0.0360。相较 ARES-NQ 的高准确率，RAGTruth 暴露了 response-level hallucination 评价的复杂性。",
            "错误类型上，测试集中 supported→unsupported 错误为 52 条，unsupported→supported 错误为 61 条。前者说明模型有时将上下文可支持但表述复杂的回答误判为不支持；后者则说明模型仍可能放过局部幻觉或隐含事实缺口。",
            "RAGTruth 的标签粒度与很多局部事实性数据不同。一个回答可能大部分内容与上下文一致，但只要关键实体、数字、时间或关系出现错误，response-level 标签就可能被标记为 unsupported。deterministic evidence proxy 能捕获一部分风险，但难以完全替代原子命题抽取和细粒度验证。",
            "本文新增的 numeric/date/entity gap、negation mismatch、comparative mismatch 和 local hallucination risk 对 RAGTruth 有帮助，但仍存在边界。特别是实体别名、同义改写、跨句推理和隐式比较关系，容易造成支持度估计偏差。",
            "因此，本文在论文叙事中不将 RAGTruth 表述为已经解决的问题，而是将其作为 BEA-Judge 进一步发展的关键泛化挑战。当前结果更适合说明：证据增强与校准能显著改善事实性可靠性，但 response-level hallucination 仍需要更细粒度的证据建模。",
            "后续可在不破坏主实验可复现性的前提下增加三类扩展：实体别名规范化，原子 claim extraction 的附录实验，以及基于轻量 verifier 的二级复核。但这些扩展应作为下一版方法，而不应混入当前冻结 baseline。",
        ],
    ),
    (
        1,
        "8 讨论",
        [
            "本文最稳定的结论是：证据增强模块对 factuality head 具有显著贡献。该贡献不仅体现在 macro-F1 提升，也体现在错误捕获和 evidence subtype 分析中。由于该模块由 deterministic 特征构成，其可解释性和复现性优于黑箱式二次 Judge。",
            "融合校准模块的主要价值体现在 pairwise 任务，尤其是 Tie policy。未校准模型倾向于低估 Tie 或将不确定样本强行归入 A>B/B>A。通过 dev-only policy 搜索，Full BEA-Judge 将 test Tie recall 提升至 0.5231。",
            "偏差模块的价值应被准确表述。它不是当前版本的主性能提升模块，甚至在部分消融指标上不优于 w/o Bias。其贡献在于将潜在偏差显式化，使研究者能识别 position、format、rubric_sensitivity 和 reasoning_difficulty 等高风险样本，并决定是否进入人工复核。",
            "偏差风险效用实验进一步支持这一定位。no_bias_decision_features 的 pairwise macro-F1 为 0.6892，高于 bias_as_decision_features 的 0.6730；但后者的 review_capture_rate 为 0.7557，高于 no_bias_decision_features 的 0.7328。因此，Bias 模块的合理叙事是提高复核资源分配效率，而不是追求最高整体 accuracy。",
            "风险覆盖实验说明 risk_score 具有实际审计价值。当 pairwise review_rate 约为 0.3998 时，系统可捕获 0.7405 的错误，剩余自动接收样本 accuracy 为 0.8924；当 factuality review_rate 为 0.4000 时，可捕获 0.6491 的错误，自动接收样本 accuracy 为 0.8625。这一结果支撑将 BEA-Judge 输出用于人工复核队列排序。",
            "从 SCI 论文角度看，本文的可信度来自四点。第一，基础 Judge 全覆盖且无 heuristic fallback。第二，10K 数据包有 license 与 provenance 审计。第三，消融与显著性检验明确区分性能贡献与解释贡献。第四，RAGTruth 难点被单独分析，没有用总体指标掩盖事实性泛化问题。",
            "本文也存在局限。首先，BEA-Judge 依赖 M-Prometheus-3B 的基础判断，基础 Judge 的系统性错误可能传递至后续校准层。其次，deterministic evidence proxy 无法处理复杂语义蕴含、跨句推理和实体别名的全部情况。第三，当前并未引入外部检索，因此对开放域事实核验的覆盖有限。",
            "另一个局限是 pairwise 难数据集表现仍不均衡。mt_bench 和 pandalm 的结果提示，当前 Tie policy 对某些数据源可能过于激进，导致正确胜负判断被替换为 Tie。未来应在 dev 集上开展 order-swap consistency probe，并将 position inconsistency 作为偏差证据，而不是直接覆盖原预测。",
            "尽管存在上述限制，本文的实验设计符合当前阶段目标：构建一个可解释、可复现、轻量的 Judge 增强与校准框架，并清晰呈现每个模块的真实贡献边界。",
        ],
    ),
    (
        1,
        "9 结论",
        [
            "本文提出 BEA-Judge-10K，一个面向大语言模型评价的偏差感知与证据增强校准框架。该框架在不微调大模型、不使用启发式 fallback、不引入外部检索的约束下，将 M-Prometheus-3B 的基础评分转化为包含预测标签、置信度、风险分数和复核建议的结构化输出。",
            "实验表明，证据增强模块显著改善事实性任务，融合校准模块提升 pairwise 与 Tie 判断稳定性，偏差模块为风险识别和复核优先级提供解释支持。BEA-Judge-10K v2 的数据规模、许可审计、全覆盖基础评分和统计显著性分析，使其具备 SCI 论文初稿的实验基础。",
            "后续工作将聚焦三方面：第一，针对 RAGTruth 的 response-level hallucination 进一步增强实体、时间、比较关系和否定关系建模；第二，引入 dev-only order-swap probe 量化位置敏感性；第三，在不改变主模型轻量性的前提下，探索附录式 verifier 或 claim-level audit。",
        ],
    ),
    (
        1,
        "数据可用性声明",
        [
            "BEA-Judge-10K v2 的 processed dataset、split files、model outputs、validation reports 与 SCI-ready tables 均在本地项目目录中生成。当前 processed dataset 包含 10200 条记录，train/dev/test 数量为 7084/1578/1538。",
            "进入训练的数据源均完成许可与 provenance 审计。HelpSteer2、OASST1、OffsetBias 与 RAGTruth 保留 acquisition date、revision 或 version reference、source URL、license 和 SHA-256 metadata。RewardBench 与其他混合许可或重分发受限来源仅作为 external evaluation 规划，不进入正式划分。",
            "正式基础 Judge 结果位于 datasets/judge_outputs/m_prometheus_3b_bea10k_v2/base_scores.repaired.json，仅包含真实 Prometheus-family 输出。校准结果位于 datasets/model_outputs/bea_judge_20260521_110114/calibrated_results.json。启发式 fallback 输出被排除在正式 SCI 结果之外。",
        ],
    ),
    (
        1,
        "伦理与利益冲突声明",
        [
            "本文使用公开或本地构建的研究数据，不涉及新增人类受试者实验。所有公开数据源均按其许可证与重分发限制进行审计和记录。",
            "本文方法不生成新的闭源标签，不使用未披露的人工标注结果，不以启发式 fallback 替代正式模型输出。作者应在投稿前根据目标期刊要求补充具体数据仓库链接、代码版本、计算环境和潜在利益冲突声明。",
            "本文使用 AI 辅助写作工具进行论文初稿组织和语言整理，但实验指标、数据来源和方法边界均来自用户提供的项目结果与本地可核验文件。投稿时应根据目标期刊政策披露 AI 辅助写作情况。",
        ],
    ),
    (
        1,
        "附录A 论点证据映射与投稿前审查清单",
        [
            "核心论点一：BEA-Judge 是轻量级 Judge 增强与校准框架，而非新的大模型训练方案。证据为基础 Judge 全覆盖、heuristic rows=0、M-Prometheus-3B repaired base scores 被冻结为正式输入。",
            "该论点应放在摘要、引言和方法开头反复约束。写作时应避免使用“训练新 Judge”“端到端评价模型”或“替代人工标注”的措辞。",
            "核心论点二：证据增强模块是事实性任务的主要性能贡献。证据为 Full factuality macro-F1=0.7405，w/o Evidence macro-F1=0.6542，Δmacro-F1=0.0863，McNemar p=0.000224。",
            "该论点适合放在结果和讨论的主线位置。对应图表可设计为事实性消融柱状图，突出 evidence module 对 accuracy、macro-F1 和 Brier 的影响。",
            "核心论点三：融合校准模块提高 pairwise 决策稳定性，尤其是 Tie recall。证据为 Full Tie recall=0.5231，w/o Calibration Tie recall=0.3923，pairwise Δmacro-F1=0.0328，McNemar p=0.034690。",
            "该论点应与置信度输出和复核阈值一起解释。写作时应强调 dev-only threshold search，避免让读者误解为使用 test 集调参。",
            "核心论点四：Bias 模块主要贡献在风险识别和复核优先级，而不是整体准确率提升。证据为 w/o Bias pairwise macro-F1=0.6892，高于 Full 的 0.6730。",
            "该论点需要在讨论中主动说明。若审稿人质疑 Bias 模块必要性，应回应其目标不是最大化单一 accuracy，而是为偏差诊断、审计和 review_flag 提供结构化依据。",
            "核心论点五：RAGTruth 是当前 factuality 泛化的主要难点。证据为 RAGTruth test macro-F1=0.6363，supported→unsupported 错误 52 条，unsupported→supported 错误 61 条。",
            "该论点应被写成局限与未来工作，而不是负面结果。RAGTruth 的 response-level hallucination 标签使其比 ARES-NQ 类局部 support 任务更难。",
            "投稿前应补充目标期刊格式。需要确认标题页、作者贡献、基金项目、利益冲突、数据可用性、代码可用性、AI 辅助写作披露和补充材料格式。",
            "投稿前应补充图1：BEA-Judge 四模块流程图。图中应包含 cleaned data、base scores、bias/evidence features、dual-head calibration 和 calibrated_results 五个节点。",
            "投稿前应补充图2：主结果与消融结果。建议以 pairwise macro-F1、Tie recall、factuality macro-F1 和 ECE 为四个 panel，减少单图中过多指标造成的信息拥挤。",
            "投稿前应补充图3：RAGTruth 错误分类。建议展示 supported→unsupported、unsupported→supported 与 review capture 的关系，用于解释 response-level hallucination 难点。",
            "投稿前应补充表格脚注。所有表格都应标注 split、n、metric 定义、是否使用 bootstrap CI，以及 threshold 是否来自 dev 集。",
            "投稿前应核验全部参考文献。当前参考文献以 arXiv、ICML、TMLR 和公开论文页面为依据，正式投稿时应按目标期刊格式补全 DOI、页码和出版信息。",
            "投稿前应复查数据许可。HelpSteer2、OASST1、OffsetBias 与 RAGTruth 的许可条款应随数据版本固化；RewardBench 继续保持 external evaluation only，除非完成子集级许可审计。",
            "投稿前应复查实验可复现性。需确认 baseline manifest、repair report、validation report、calibrated_results、ablation report 与 SCI tables 均被纳入归档。",
            "投稿前应复查统计叙事。主文中应报告点估计、95% CI 和显著性检验；附录可列出 paired bootstrap 与 McNemar 检验细节。",
            "投稿前应复查风险边界。本文不能宣称证据模块完成原子事实验证，也不能宣称 RAGTruth 已被解决；合理表述是证据增强显著改善 factuality reliability，但 response-level hallucination 仍是开放挑战。",
            "投稿前应复查模型边界。本文不微调 M-Prometheus-3B，不接外部检索，不做 LLM claim extraction，不使用 heuristic fallback 作为正式结果。",
            "投稿前应准备审稿回复策略。若审稿人要求更强 baseline，可在附录增加 LightGBM 或 MLP 对照，但主模型仍建议保持 dual-head softmax，以维护可解释性。",
            "投稿前应准备 RAGTruth 案例分析。建议从三类样本中各选 2-3 个代表案例：unsupported→supported、supported→unsupported、低置信正确样本。",
            "投稿前应准备偏差案例分析。建议展示 position、format、rubric_sensitivity 和 reasoning_difficulty 四类高风险样本各 1 个，并说明 review_reason 如何辅助人工复核。",
            "投稿前应准备 BEA-Judge-10K v2 内部对照。重点比较 Raw M-Prometheus、Full、w/o Evidence、w/o Calibration、w/o Bias、calibration methods 与 risk coverage，不引入早期版本对照。",
            "投稿前应控制结论强度。建议使用“表明”“支持”“提示”“在当前设置下”而非“证明”“彻底解决”“显著优于所有方法”等过强措辞。",
            "投稿前应进行语言润色。中文初稿可先用于组内审阅，正式 SCI 投稿通常需要英文稿或双语摘要，届时应进行专业英文改写。",
            "投稿前应保存生成脚本。当前 Word 和 Markdown 由 scripts/generate_bea_judge_sci_cn_docx.py 生成，便于后续指标或表格更新后重新生成稿件。",
        ],
    ),
]


TABLES_IN_TEXT = [
    ("表1 BEA-Judge-10K v2 数据分布", DATASET_TABLE),
    ("表2 数据来源与许可审计", LICENSE_TABLE),
    ("表3 主实验结果", MAIN_RESULTS),
    ("表4 测试集指标的 bootstrap 95% 置信区间", CI_TABLE),
    ("表5 BEA-Judge-10K v2 内部基线对照", BASELINE_COMPARISON_TABLE),
    ("表6 校准方法对照", CALIBRATION_METHODS_TABLE),
    ("表7 模块消融结果", ABLATION_TABLE),
    ("表8 消融显著性检验", SIGNIFICANCE_TABLE),
    ("表9 分数据集结果", PER_DATASET_TABLE),
    ("表10 RAGTruth 错误分类结果", RAGTRUTH_TABLE),
    ("表11 偏差 subgroup calibration", BIAS_TABLE),
    ("表12 证据特征组消融", EVIDENCE_ABLATION_TABLE),
    ("表13 偏差风险效用分析", BIAS_RISK_UTILITY_TABLE),
    ("表14 风险覆盖与自动接收准确率", RISK_COVERAGE_TABLE),
]

FIGURES_IN_TEXT = {
    "3 数据集构建与审计": [
        (
            "图2 BEA-Judge-10K v2 数据构成。该图展示任务类型、训练/开发/测试划分、标签分布和数据来源，说明 10K v2 数据包兼顾 pairwise preference、pairwise bias 与 RAG factuality。",
            "fig2_dataset_distribution.png",
        )
    ],
    "4 方法": [
        (
            "图1 BEA-Judge 四模块公式化流程。基础 Judge 输出 s_i，经偏差风险、证据风险与双任务 softmax 校准后生成 predicted_label、confidence、risk_score 与 review_flag。",
            "fig1_pipeline.png",
        )
    ],
    "6 实验结果": [
        (
            "图3 主实验结果与 bootstrap 95% 置信区间。Pairwise head 重点报告 macro-F1 与 Tie recall，factuality head 重点报告 macro-F1 与 ECE。",
            "fig3_main_results_ci.png",
        ),
        (
            "图4 消融结果与显著性证据。Evidence 模块对 factuality 的贡献最稳定，Calibration/Tie policy 对 pairwise 决策稳定性贡献最明显。",
            "fig4_ablation_significance.png",
        ),
    ],
    "7 RAGTruth 难点分析": [
        (
            "图5 RAGTruth response-level hallucination 难点。RAGTruth 与 ARES-NQ 的差异说明事实性泛化不仅是局部 overlap 问题，还涉及回答级别的实体、数字、时间和关系一致性。",
            "fig5_ragtruth_analysis.png",
        )
    ],
    "8 讨论": [
        (
            "图6 偏差复核与证据特征诊断。Bias subgroup 图支持将偏差模块定位为复核优先级机制，证据特征组消融显示 evidence feature 与 weighted calibration 需要联合设计。",
            "fig6_bias_evidence_diagnostics.png",
        )
    ],
}


def _set_font(run, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def md_table(rows: list[list[str]]) -> str:
    header = "| " + " | ".join(rows[0]) + " |"
    sep = "| " + " | ".join(["---"] * len(rows[0])) + " |"
    body = ["| " + " | ".join(row) + " |" for row in rows[1:]]
    return "\n".join([header, sep, *body])


def add_md_figures(lines: list[str], section_title: str) -> None:
    for caption, filename in FIGURES_IN_TEXT.get(section_title, []):
        fig_path = FIG_DIR / filename
        lines.append(f"![{caption}]({fig_path.as_posix()})")
        lines.append("")
        lines.append(f"**{caption}**")
        lines.append("")


def build_markdown() -> str:
    lines: list[str] = []
    lines.append(f"# {TITLE}")
    lines.append("")
    lines.append("作者信息：投稿前由作者团队按目标期刊要求补充")
    lines.append("")
    lines.append("版本说明：本稿基于 frozen formal baseline `bea_judge_20260521_110114` 生成。")
    lines.append("")
    for level, title, paragraphs in SECTIONS:
        lines.append("#" * (level + 1) + f" {title}")
        lines.append("")
        for paragraph in paragraphs:
            lines.append(paragraph)
            lines.append("")
        add_md_figures(lines, title)
        if title == "3 数据集构建与审计":
            for caption, table in TABLES_IN_TEXT[:2]:
                lines.append(f"**{caption}**")
                lines.append("")
                lines.append(md_table(table))
                lines.append("")
        if title == "6 实验结果":
            for caption, table in TABLES_IN_TEXT[2:10]:
                lines.append(f"**{caption}**")
                lines.append("")
                lines.append(md_table(table))
                lines.append("")
        if title == "8 讨论":
            for caption, table in TABLES_IN_TEXT[10:]:
                lines.append(f"**{caption}**")
                lines.append("")
                lines.append(md_table(table))
                lines.append("")
    lines.append("## 参考文献")
    lines.append("")
    for i, ref in enumerate(REFERENCES, 1):
        lines.append(f"[{i}] {ref}")
        lines.append("")
    return "\n".join(lines)


def add_paragraph(doc: Document, text: str, style: str | None = None, align=None) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _set_font(run, 10 if style == "Body Text" else 10)
    p.paragraph_format.first_line_indent = Pt(21) if style == "Body Text" else None
    p.paragraph_format.line_spacing = 1.25


def add_table(doc: Document, caption: str, rows: list[list[str]]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    _set_font(r, 10, True)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(value))
            _set_font(run, 8, row_idx == 0)
            if row_idx == 0:
                _set_cell_shading(cell, "D9EAF7")
    doc.add_paragraph()


def add_figure(doc: Document, caption: str, filename: str) -> None:
    path = FIG_DIR / filename
    if not path.exists():
        add_paragraph(doc, f"[图文件缺失，待生成：{filename}]", style="Body Text")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.3))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    _set_font(cap_run, 9, True)
    doc.add_paragraph()


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(TITLE)
    _set_font(title_run, 16, True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("中文SCI论文初稿（公式、图表与文献增强版）")
    _set_font(sub_run, 11, False)

    add_paragraph(doc, "作者信息：投稿前由作者团队按目标期刊要求补充", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "版本说明：本稿基于 frozen formal baseline bea_judge_20260521_110114 生成。", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    toc = doc.add_paragraph()
    toc_run = toc.add_run("目录：投稿排版时可由 Word 自动生成。")
    _set_font(toc_run, 10, False)
    doc.add_page_break()

    for level, title_text, paragraphs in SECTIONS:
        doc.add_heading(title_text, level=level)
        for paragraph in paragraphs:
            add_paragraph(doc, paragraph, style="Body Text")
        for caption, filename in FIGURES_IN_TEXT.get(title_text, []):
            add_figure(doc, caption, filename)
        if title_text == "3 数据集构建与审计":
            for caption, table_rows in TABLES_IN_TEXT[:2]:
                add_table(doc, caption, table_rows)
        if title_text == "6 实验结果":
            for caption, table_rows in TABLES_IN_TEXT[2:10]:
                add_table(doc, caption, table_rows)
        if title_text == "8 讨论":
            for caption, table_rows in TABLES_IN_TEXT[10:]:
                add_table(doc, caption, table_rows)

    doc.add_heading("参考文献", level=1)
    for i, ref in enumerate(REFERENCES, 1):
        add_paragraph(doc, f"[{i}] {ref}", style="Body Text")

    props = doc.core_properties
    props.title = SHORT_TITLE
    props.subject = "BEA-Judge-10K v2 SCI Chinese manuscript draft"
    props.keywords = "LLM-as-a-Judge, calibration, factuality, RAGTruth, bias-aware evaluation"
    props.comments = "Generated from frozen BEA-Judge v2 SCI tables. No experimental metrics were altered."

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    MD_PATH.write_text(build_markdown(), encoding="utf-8")
    build_docx()
    print(f"markdown={MD_PATH}")
    print(f"docx={DOCX_PATH}")


if __name__ == "__main__":
    main()
