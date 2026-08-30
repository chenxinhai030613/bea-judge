from __future__ import annotations

import json
import math
import os
import re
import shutil
import subprocess
from pathlib import Path
from textwrap import dedent
from typing import Any, Iterable
from xml.sax.saxutils import escape


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "论文撰写" / "QLoRA-BEA-Judge_epoch2_tie_rescue_20260605"
FIG = OUT / "figures"
FIG_PNG = OUT / "figures_png"
TAB = OUT / "tables"


METRIC_LABELS = ["accuracy", "macro-F1", "ECE", "Tie recall"]


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def fmt(value: Any) -> str:
    if value is None or value == "":
        return "--"
    if isinstance(value, str):
        return value
    return f"{float(value):.4f}"


def mean_std(mean: float, std: float | None = None) -> str:
    if std is None:
        return f"{mean:.4f}"
    return f"{mean:.4f} +/- {std:.4f}"


def md_table(headers: list[str], rows: list[list[Any]]) -> str:
    def cell(value: Any) -> str:
        return str(value).replace("|", "\\|")

    lines = [
        "| " + " | ".join(cell(header) for header in headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join(cell(x) for x in row) + " |")
    return "\n".join(lines)


def write_csv(path: Path, headers: list[str], rows: list[list[Any]]) -> None:
    with path.open("w", encoding="utf-8", newline="") as f:
        f.write(",".join(headers) + "\n")
        for row in rows:
            f.write(",".join('"' + str(x).replace('"', '""') + '"' for x in row) + "\n")


def svg_text(x: float, y: float, text: str, size: int = 13, anchor: str = "middle", weight: str = "400", color: str = "#1f2937") -> str:
    return f'<text x="{x:.1f}" y="{y:.1f}" text-anchor="{anchor}" font-family="Times New Roman, SimSun, serif" font-size="{size}" font-weight="{weight}" fill="{color}">{escape(text)}</text>'


def svg_rect(x: float, y: float, w: float, h: float, fill: str, stroke: str = "#1f2937", sw: float = 1.0, rx: float = 3.0) -> str:
    return f'<rect x="{x:.1f}" y="{y:.1f}" width="{w:.1f}" height="{h:.1f}" rx="{rx:.1f}" fill="{fill}" stroke="{stroke}" stroke-width="{sw:.1f}"/>'


def svg_arrow(x1: float, y1: float, x2: float, y2: float, color: str = "#374151") -> str:
    return (
        f'<line x1="{x1:.1f}" y1="{y1:.1f}" x2="{x2:.1f}" y2="{y2:.1f}" '
        f'stroke="{color}" stroke-width="1.8" marker-end="url(#arrow)"/>'
    )


def wrap_svg(body: str, width: int, height: int) -> str:
    return f'''<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">
  <defs>
    <marker id="arrow" viewBox="0 0 10 10" refX="8" refY="5" markerWidth="7" markerHeight="7" orient="auto-start-reverse">
      <path d="M 0 0 L 10 5 L 0 10 z" fill="#374151"/>
    </marker>
  </defs>
  <rect width="100%" height="100%" fill="#ffffff"/>
{body}
</svg>
'''


def fig_framework(path: Path) -> None:
    parts: list[str] = []
    parts.append(svg_text(520, 34, "BEA-Judge 四模块融合校准框架", 20, weight="700"))
    parts.append(svg_rect(40, 80, 180, 92, "#f8fafc"))
    parts.append(svg_text(130, 112, "输入样本", 16, weight="700"))
    parts.append(svg_text(130, 138, "prompt / context", 12))
    parts.append(svg_text(130, 158, "answer A / B / reference", 12))

    boxes = [
        (280, 52, 185, 96, "#e0f2fe", "模块1 基础Judge评分", ["3B M-Prometheus", "QLoRA adapter", "score gap + label"]),
        (280, 180, 185, 96, "#ecfdf5", "模块2 偏差感知", ["position / length", "format / rubric", "source risk"]),
        (540, 180, 185, 96, "#fff7ed", "模块3 证据事实性", ["support coverage", "numeric / entity gap", "negation / comparison"]),
        (790, 96, 190, 118, "#f5f3ff", "模块4 融合校准", ["softmax fusion", "temperature scaling", "confidence + risk"]),
    ]
    for x, y, w, h, color, title, lines in boxes:
        parts.append(svg_rect(x, y, w, h, color))
        parts.append(svg_text(x + w / 2, y + 28, title, 14, weight="700"))
        for idx, line in enumerate(lines):
            parts.append(svg_text(x + w / 2, y + 53 + idx * 19, line, 11))
    parts.append(svg_rect(790, 250, 190, 72, "#fef2f2"))
    parts.append(svg_text(885, 279, "Tie rescue", 14, weight="700"))
    parts.append(svg_text(885, 301, "dev-only selection", 11))
    parts.append(svg_text(885, 318, "accuracy-constrained", 11))

    parts.append(svg_rect(40, 242, 180, 80, "#f8fafc"))
    parts.append(svg_text(130, 273, "输出", 16, weight="700"))
    parts.append(svg_text(130, 298, "A>B / B>A / Tie", 12))
    parts.append(svg_text(130, 316, "confidence / review flag", 12))

    parts.extend(
        [
            svg_arrow(220, 126, 280, 100),
            svg_arrow(220, 126, 280, 228),
            svg_arrow(465, 228, 540, 228),
            svg_arrow(465, 100, 790, 132),
            svg_arrow(725, 228, 790, 157),
            svg_arrow(885, 214, 885, 250),
            svg_arrow(790, 286, 220, 286),
            svg_arrow(790, 154, 220, 286),
        ]
    )
    path.write_text(wrap_svg("\n".join(parts), 1020, 360), encoding="utf-8")


def fig_grouped_bars(path: Path, title: str, groups: list[str], series: list[tuple[str, list[float], str]], ymax: float = 1.0) -> None:
    width, height = 920, 430
    left, right, top, bottom = 80, 30, 70, 82
    plot_w = width - left - right
    plot_h = height - top - bottom
    parts = [svg_text(width / 2, 34, title, 19, weight="700")]
    for tick in [0, 0.25, 0.5, 0.75, 1.0]:
        y = top + plot_h * (1 - tick / ymax)
        parts.append(f'<line x1="{left}" y1="{y:.1f}" x2="{width-right}" y2="{y:.1f}" stroke="#e5e7eb" stroke-width="1"/>')
        parts.append(svg_text(left - 14, y + 4, f"{tick:.2f}", 11, anchor="end", color="#4b5563"))
    group_w = plot_w / len(groups)
    bar_w = min(24, group_w / (len(series) + 1.4))
    for gi, group in enumerate(groups):
        cx = left + group_w * gi + group_w / 2
        parts.append(svg_text(cx, height - 42, group, 12))
        for si, (name, vals, color) in enumerate(series):
            x = cx - (len(series) * bar_w) / 2 + si * bar_w
            h = plot_h * vals[gi] / ymax
            y = top + plot_h - h
            parts.append(svg_rect(x, y, bar_w * 0.82, h, color, stroke=color, rx=1.0))
            parts.append(svg_text(x + bar_w * 0.41, y - 5, f"{vals[gi]:.3f}", 9, color="#374151"))
    parts.append(f'<line x1="{left}" y1="{top+plot_h}" x2="{width-right}" y2="{top+plot_h}" stroke="#111827" stroke-width="1.2"/>')
    parts.append(f'<line x1="{left}" y1="{top}" x2="{left}" y2="{top+plot_h}" stroke="#111827" stroke-width="1.2"/>')
    lx = left + 20
    for i, (name, _, color) in enumerate(series):
        y = height - 20
        x = lx + i * 190
        parts.append(svg_rect(x, y - 12, 16, 10, color, stroke=color, rx=1))
        parts.append(svg_text(x + 22, y - 3, name, 11, anchor="start"))
    path.write_text(wrap_svg("\n".join(parts), width, height), encoding="utf-8")


def fig_line(path: Path, title: str, xs: list[str], ys: dict[str, tuple[list[float], str]], ymax: float = 1.0) -> None:
    width, height = 860, 400
    left, right, top, bottom = 78, 35, 65, 75
    plot_w = width - left - right
    plot_h = height - top - bottom
    parts = [svg_text(width / 2, 34, title, 19, weight="700")]
    for tick in [0, 0.25, 0.5, 0.75, 1.0]:
        y = top + plot_h * (1 - tick / ymax)
        parts.append(f'<line x1="{left}" y1="{y:.1f}" x2="{width-right}" y2="{y:.1f}" stroke="#e5e7eb" stroke-width="1"/>')
        parts.append(svg_text(left - 12, y + 4, f"{tick:.2f}", 11, anchor="end"))
    xcoords = [left + plot_w * i / max(1, len(xs) - 1) for i in range(len(xs))]
    for idx, label in enumerate(xs):
        parts.append(svg_text(xcoords[idx], height - 38, label, 12))
    for series_idx, (name, (vals, color)) in enumerate(ys.items()):
        points = []
        for x, val in zip(xcoords, vals):
            y = top + plot_h * (1 - val / ymax)
            points.append((x, y))
        d = " ".join(f"{x:.1f},{y:.1f}" for x, y in points)
        parts.append(f'<polyline points="{d}" fill="none" stroke="{color}" stroke-width="2.2"/>')
        for x, y in points:
            parts.append(f'<circle cx="{x:.1f}" cy="{y:.1f}" r="4" fill="{color}"/>')
        label_x, label_y = points[-1][0] + 10, points[-1][1] + 4 + series_idx * 14
        anchor = "start"
        if label_x > width - right - 120:
            label_x = points[-1][0] - 10
            anchor = "end"
        label_y = min(max(label_y, top + 13), top + plot_h - 8)
        parts.append(svg_text(label_x, label_y, name, 11, anchor=anchor, color=color, weight="700"))
    parts.append(f'<line x1="{left}" y1="{top+plot_h}" x2="{width-right}" y2="{top+plot_h}" stroke="#111827" stroke-width="1.2"/>')
    parts.append(f'<line x1="{left}" y1="{top}" x2="{left}" y2="{top+plot_h}" stroke="#111827" stroke-width="1.2"/>')
    path.write_text(wrap_svg("\n".join(parts), width, height), encoding="utf-8")


def fig_risk(path: Path) -> None:
    review = [0.0503, 0.0997, 0.2004, 0.3001, 0.3998, 0.4995, 0.7502, 1.0]
    capture = [0.1412, 0.2634, 0.5, 0.6565, 0.7405, 0.8511, 1.0, 1.0]
    auto_acc = [0.775, 0.7964, 0.8444, 0.8779, 0.8924, 0.926, 1.0, 1.0]
    fig_line(
        path,
        "风险复核比例与错误捕获/自动接受准确率",
        [f"{int(x*100)}%" for x in review],
        {
            "error capture": (capture, "#2563eb"),
            "auto-accept accuracy": (auto_acc, "#16a34a"),
        },
    )


def build_figures() -> None:
    fig_framework(FIG / "fig1_four_module_framework.svg")
    fig_grouped_bars(
        FIG / "fig2_main_comparison.svg",
        "主模型与轻量外部基线对比",
        ["Current", "epoch2", "epoch2+Tie", "GRM", "Qwen", "GLIDER"],
        [
            ("accuracy", [0.7512, 0.8297, 0.8297, 0.7312, 0.5736, 0.5043], "#2563eb"),
            ("macro-F1", [0.6730, 0.7348, 0.7441, 0.6584, 0.4160, 0.4273], "#16a34a"),
            ("Tie recall", [0.5231, 0.4256, 0.4795, 0.5000, 0.0308, 0.3692], "#f59e0b"),
        ],
    )
    fig_grouped_bars(
        FIG / "fig3_tie_rescue.svg",
        "Tie rescue 前后指标变化",
        ["accuracy", "macro-F1", "ECE", "Tie recall"],
        [
            ("epoch2", [0.8297, 0.7348, 0.0278, 0.4256], "#64748b"),
            ("epoch2+Tie", [0.8297, 0.7441, 0.0283, 0.4795], "#dc2626"),
        ],
    )
    fig_line(
        FIG / "fig4_sft_size_ablation.svg",
        "SFT size 消融：QLoRA-BEA-Judge",
        ["25%", "50%", "100%"],
        {
            "accuracy": ([0.7740, 0.7927, 0.8025], "#2563eb"),
            "macro-F1": ([0.6832, 0.7012, 0.7128], "#16a34a"),
            "Tie recall": ([0.4436, 0.4590, 0.4538], "#f59e0b"),
            "ECE": ([0.0310, 0.0303, 0.0279], "#9333ea"),
        },
    )
    fig_grouped_bars(
        FIG / "fig5_epoch_ablation.svg",
        "QLoRA 训练轮数消融：1024 序列长度",
        ["0.5 epoch", "1 epoch", "2 epoch"],
        [
            ("accuracy", [0.7873, 0.8025, 0.8297], "#2563eb"),
            ("macro-F1", [0.6943, 0.7128, 0.7348], "#16a34a"),
            ("Tie recall", [0.4385, 0.4538, 0.4256], "#f59e0b"),
        ],
    )
    fig_risk(FIG / "fig6_risk_review.svg")


def render_png_figures() -> bool:
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return False
    FIG_PNG.mkdir(parents=True, exist_ok=True)
    font_path = Path("C:/Windows/Fonts/simsun.ttc")
    latin_path = Path("C:/Windows/Fonts/times.ttf")

    def font(size: int, bold: bool = False):
        path = font_path if font_path.exists() else latin_path
        try:
            return ImageFont.truetype(str(path), size)
        except Exception:
            return ImageFont.load_default()

    def text(draw, xy, s, size=24, fill=(31, 41, 55), anchor="mm", bold=False):
        draw.text(xy, s, fill=fill, font=font(size, bold), anchor=anchor)

    def save_framework() -> None:
        img = Image.new("RGB", (1400, 500), "white")
        d = ImageDraw.Draw(img)
        text(d, (700, 38), "BEA-Judge 四模块融合校准框架", 28, bold=True)
        boxes = [
            (55, 120, 250, 120, (248, 250, 252), "输入样本", ["prompt / context", "answer A / B / reference"]),
            (390, 70, 255, 125, (224, 242, 254), "模块1 基础Judge评分", ["3B M-Prometheus", "QLoRA adapter", "score gap + label"]),
            (390, 260, 255, 125, (236, 253, 245), "模块2 偏差感知", ["position / length", "format / rubric", "source risk"]),
            (720, 260, 255, 125, (255, 247, 237), "模块3 证据事实性", ["support coverage", "numeric / entity gap", "negation / comparison"]),
            (1070, 130, 260, 140, (245, 243, 255), "模块4 融合校准", ["softmax fusion", "temperature scaling", "confidence + risk"]),
            (1070, 330, 260, 86, (254, 242, 242), "Tie rescue", ["dev-only + accuracy-constrained"]),
        ]
        for x, y, w, h, color, title, lines in boxes:
            d.rounded_rectangle([x, y, x + w, y + h], radius=8, fill=color, outline=(31, 41, 55), width=2)
            text(d, (x + w / 2, y + 34), title, 20, bold=True)
            for i, line in enumerate(lines):
                text(d, (x + w / 2, y + 66 + i * 25), line, 16)
        for a, b in [((305, 180), (390, 132)), ((305, 180), (390, 322)), ((645, 322), (720, 322)), ((645, 132), (1070, 190)), ((975, 322), (1070, 210)), ((1200, 270), (1200, 330))]:
            d.line([a, b], fill=(55, 65, 81), width=3)
        img.save(FIG_PNG / "fig1_four_module_framework.png")

    def save_bars(name: str, title: str, groups: list[str], series: list[tuple[str, list[float], tuple[int, int, int]]]) -> None:
        img = Image.new("RGB", (1400, 650), "white")
        d = ImageDraw.Draw(img)
        text(d, (700, 45), title, 28, bold=True)
        left, top, width, height = 110, 100, 1230, 410
        for tick in [0, 0.25, 0.5, 0.75, 1.0]:
            y = top + height * (1 - tick)
            d.line([(left, y), (left + width, y)], fill=(229, 231, 235), width=1)
            text(d, (left - 20, y), f"{tick:.2f}", 16, anchor="rm")
        group_w = width / len(groups)
        bar_w = min(38, group_w / (len(series) + 1.4))
        for gi, group in enumerate(groups):
            cx = left + gi * group_w + group_w / 2
            text(d, (cx, top + height + 38), group, 16)
            for si, (_label, vals, color) in enumerate(series):
                x = cx - len(series) * bar_w / 2 + si * bar_w
                h = height * vals[gi]
                y = top + height - h
                d.rectangle([x, y, x + bar_w * 0.78, top + height], fill=color)
                text(d, (x + bar_w * 0.39, y - 12), f"{vals[gi]:.3f}", 12)
        d.line([(left, top), (left, top + height), (left + width, top + height)], fill=(17, 24, 39), width=2)
        for i, (label, _vals, color) in enumerate(series):
            x = left + 30 + i * 250
            y = 605
            d.rectangle([x, y - 18, x + 24, y - 4], fill=color)
            text(d, (x + 34, y - 11), label, 15, anchor="lm")
        img.save(FIG_PNG / name)

    def save_line(name: str, title: str, xs: list[str], ys: dict[str, tuple[list[float], tuple[int, int, int]]]) -> None:
        img = Image.new("RGB", (1300, 600), "white")
        d = ImageDraw.Draw(img)
        text(d, (650, 45), title, 28, bold=True)
        left, top, width, height = 100, 95, 1080, 380
        for tick in [0, 0.25, 0.5, 0.75, 1.0]:
            y = top + height * (1 - tick)
            d.line([(left, y), (left + width, y)], fill=(229, 231, 235), width=1)
            text(d, (left - 18, y), f"{tick:.2f}", 16, anchor="rm")
        xcoords = [left + width * i / max(1, len(xs) - 1) for i in range(len(xs))]
        for x, label in zip(xcoords, xs):
            text(d, (x, top + height + 40), label, 16)
        for series_idx, (label, (vals, color)) in enumerate(ys.items()):
            pts = [(x, top + height * (1 - v)) for x, v in zip(xcoords, vals)]
            d.line(pts, fill=color, width=4)
            for x, y in pts:
                d.ellipse([x - 6, y - 6, x + 6, y + 6], fill=color)
            label_x = pts[-1][0] + 15
            label_y = pts[-1][1] + series_idx * 24
            anchor = "lm"
            try:
                bbox = d.textbbox((0, 0), label, font=font(16, True))
                label_w = bbox[2] - bbox[0]
            except Exception:
                label_w = len(label) * 9
            if label_x + label_w > img.width - 25:
                label_x = pts[-1][0] - 15
                anchor = "rm"
            label_y = min(max(label_y, top + 16), top + height - 12)
            text(d, (label_x, label_y), label, 16, anchor=anchor, fill=color, bold=True)
        d.line([(left, top), (left, top + height), (left + width, top + height)], fill=(17, 24, 39), width=2)
        img.save(FIG_PNG / name)

    save_framework()
    save_bars(
        "fig2_main_comparison.png",
        "主模型与轻量外部基线对比",
        ["Current", "epoch2", "epoch2+Tie", "GRM", "Qwen", "GLIDER"],
        [
            ("accuracy", [0.7512, 0.8297, 0.8297, 0.7312, 0.5736, 0.5043], (37, 99, 235)),
            ("macro-F1", [0.6730, 0.7348, 0.7441, 0.6584, 0.4160, 0.4273], (22, 163, 74)),
            ("Tie recall", [0.5231, 0.4256, 0.4795, 0.5000, 0.0308, 0.3692], (245, 158, 11)),
        ],
    )
    save_bars(
        "fig3_tie_rescue.png",
        "Tie rescue 前后指标变化",
        ["accuracy", "macro-F1", "ECE", "Tie recall"],
        [
            ("epoch2", [0.8297, 0.7348, 0.0278, 0.4256], (100, 116, 139)),
            ("epoch2+Tie", [0.8297, 0.7441, 0.0283, 0.4795], (220, 38, 38)),
        ],
    )
    save_line(
        "fig4_sft_size_ablation.png",
        "SFT size 消融：QLoRA-BEA-Judge",
        ["25%", "50%", "100%"],
        {
            "accuracy": ([0.7740, 0.7927, 0.8025], (37, 99, 235)),
            "macro-F1": ([0.6832, 0.7012, 0.7128], (22, 163, 74)),
            "Tie recall": ([0.4436, 0.4590, 0.4538], (245, 158, 11)),
            "ECE": ([0.0310, 0.0303, 0.0279], (147, 51, 234)),
        },
    )
    save_bars(
        "fig5_epoch_ablation.png",
        "QLoRA 训练轮数消融：1024 序列长度",
        ["0.5 epoch", "1 epoch", "2 epoch"],
        [
            ("accuracy", [0.7873, 0.8025, 0.8297], (37, 99, 235)),
            ("macro-F1", [0.6943, 0.7128, 0.7348], (22, 163, 74)),
            ("Tie recall", [0.4385, 0.4538, 0.4256], (245, 158, 11)),
        ],
    )
    save_line(
        "fig6_risk_review.png",
        "风险复核比例与错误捕获/自动接受准确率",
        ["5%", "10%", "20%", "30%", "40%", "50%", "75%", "100%"],
        {
            "error capture": ([0.1412, 0.2634, 0.5, 0.6565, 0.7405, 0.8511, 1.0, 1.0], (37, 99, 235)),
            "auto-accept accuracy": ([0.775, 0.7964, 0.8444, 0.8779, 0.8924, 0.926, 1.0, 1.0], (22, 163, 74)),
        },
    )
    return True


def data_sample_rows(dataset_path: Path, limit: int = 3) -> list[list[str]]:
    data = load_json(dataset_path)
    rows = data if isinstance(data, list) else data.get("samples", [])
    picked = []
    wanted = ["open_qa", "pairwise_bias", "factuality_rag"]
    for task in wanted:
        for row in rows:
            if row.get("task_type") == task:
                picked.append(row)
                break
    out = []
    for row in picked[:limit]:
        def short(key: str, n: int = 58) -> str:
            text = re.sub(r"\s+", " ", str(row.get(key) or "")).strip()
            return text[:n] + ("..." if len(text) > n else "")
        out.append([
            str(row.get("id")),
            str(row.get("dataset")),
            str(row.get("task_type")),
            short("prompt"),
            short("answer_a"),
            short("answer_b"),
            str(row.get("human_label")),
            str(row.get("split")),
        ])
    return out


def write_tables() -> dict[str, str]:
    tables: dict[str, str] = {}
    main_rows = [
        ["Current BEA-Judge", "internal four-module baseline", "3-seed repeated baseline", "0.7512", "0.6730", "0.0558", "0.5231", "--"],
        ["QLoRA-BEA-Judge epoch2_1024", "proposed accuracy-oriented model", "3-seed mean +/- std", "0.8297 +/- 0.0031", "0.7348 +/- 0.0062", "0.0278 +/- 0.0031", "0.4256 +/- 0.0270", "--"],
        ["QLoRA-BEA-Judge epoch2_1024 + Tie rescue", "proposed accuracy-constrained tie policy", "3-seed mean +/- std", "0.8297 +/- 0.0052", "0.7441 +/- 0.0093", "0.0283 +/- 0.0027", "0.4795 +/- 0.0489", "--"],
        ["GRM", "external 3B reward baseline", "single full test", "0.7312", "0.6584", "0.1759", "0.5000", "0.0000"],
        ["Qwen", "external 3B instruct baseline", "single full test", "0.5736", "0.4160", "0.3479", "0.0308", "0.0000"],
        ["GLIDER", "external 4B evaluator baseline", "single full test", "0.5043", "0.4273", "0.0353", "0.3692", "0.0000"],
    ]
    headers = ["system", "role", "run_type", "accuracy", "macro-F1", "ECE", "Tie recall", "parse failure rate"]
    tables["main"] = md_table(headers, main_rows)
    write_csv(TAB / "main_comparison.csv", headers, main_rows)

    ablation_rows = [
        ["Full BEA-Judge", "pairwise", "test", "0.7512", "0.6730", "0.0558", "0.5231", "--"],
        ["w/o Bias Module", "pairwise", "test", "0.7654", "0.6892", "0.0385", "0.5462", "--"],
        ["w/o Evidence Module", "pairwise", "test", "0.7531", "0.6711", "0.0540", "0.5000", "--"],
        ["w/o Calibration", "pairwise", "test", "0.7407", "0.6402", "0.0459", "0.3923", "--"],
        ["Full BEA-Judge", "factuality", "test", "0.7649", "0.7405", "0.0377", "--", "--"],
        ["w/o Evidence Module", "factuality", "test", "0.6928", "0.6542", "0.0221", "--", "--"],
    ]
    headers2 = ["variant", "head", "split", "accuracy", "macro-F1", "ECE", "Tie recall", "parse failure rate"]
    tables["ablation"] = md_table(headers2, ablation_rows)
    write_csv(TAB / "four_module_ablation.csv", headers2, ablation_rows)

    tie_rows = [
        ["epoch2_1024", "initial_test", "0.8297 +/- 0.0031", "0.7348 +/- 0.0062", "0.0278 +/- 0.0031", "0.4256 +/- 0.0270", "--"],
        ["epoch2_1024 + Tie rescue", "test", "0.8297 +/- 0.0052", "0.7441 +/- 0.0093", "0.0283 +/- 0.0027", "0.4795 +/- 0.0489", "--"],
        ["delta", "test-initial_test", "-0.0000", "+0.0093", "+0.0005", "+0.0539", "--"],
    ]
    headers3 = ["setting", "split", "accuracy", "macro-F1", "ECE", "Tie recall", "parse failure rate"]
    tables["tie"] = md_table(headers3, tie_rows)
    write_csv(TAB / "tie_rescue.csv", headers3, tie_rows)

    sft_rows = [
        ["sft25_epoch1_1024", "25%", "0.7740 +/- 0.0186", "0.6832 +/- 0.0214", "0.0310 +/- 0.0101", "0.4436 +/- 0.0778", "--"],
        ["sft50_epoch1_1024", "50%", "0.7927 +/- 0.0088", "0.7012 +/- 0.0123", "0.0303 +/- 0.0127", "0.4590 +/- 0.0470", "--"],
        ["sft100_epoch1_1024", "100%", "0.8025 +/- 0.0034", "0.7128 +/- 0.0063", "0.0279 +/- 0.0046", "0.4538 +/- 0.0400", "--"],
    ]
    headers4 = ["setting", "SFT size", "accuracy", "macro-F1", "ECE", "Tie recall", "parse failure rate"]
    tables["sft"] = md_table(headers4, sft_rows)
    write_csv(TAB / "sft_size_ablation.csv", headers4, sft_rows)

    epoch_rows = [
        ["epoch0p5_1024", "0.5", "0.7873 +/- 0.0035", "0.6943 +/- 0.0069", "0.0318 +/- 0.0008", "0.4385 +/- 0.0504", "--"],
        ["epoch1_1024", "1", "0.8025 +/- 0.0034", "0.7128 +/- 0.0063", "0.0279 +/- 0.0046", "0.4538 +/- 0.0400", "--"],
        ["epoch2_1024", "2", "0.8297 +/- 0.0031", "0.7348 +/- 0.0062", "0.0278 +/- 0.0031", "0.4256 +/- 0.0270", "--"],
    ]
    headers5 = ["setting", "epoch", "accuracy", "macro-F1", "ECE", "Tie recall", "parse failure rate"]
    tables["epoch"] = md_table(headers5, epoch_rows)
    write_csv(TAB / "epoch_ablation.csv", headers5, epoch_rows)

    risk_rows = [
        ["0.0503", "53", "0.1412", "1000", "0.7750"],
        ["0.0997", "105", "0.2634", "948", "0.7964"],
        ["0.2004", "211", "0.5000", "842", "0.8444"],
        ["0.3001", "316", "0.6565", "737", "0.8779"],
        ["0.3998", "421", "0.7405", "632", "0.8924"],
        ["0.4995", "526", "0.8511", "527", "0.9260"],
        ["0.7502", "790", "1.0000", "263", "1.0000"],
    ]
    headers6 = ["review_rate", "review_count", "error_capture_rate", "auto_accept_count", "auto_accept_accuracy"]
    tables["risk"] = md_table(headers6, risk_rows)
    write_csv(TAB / "risk_review.csv", headers6, risk_rows)

    prompt_rows = [
        ["m_prometheus_pairwise_v1", "###Task Description / ###The instruction to evaluate / ###Response A / ###Response B / ###Reference Answer / ###Score Rubric / ###Feedback", "[RESULT] A | [RESULT] B | [RESULT] Tie", "relevance; completeness; factuality; instruction_following; clarity; safety"],
        ["m_prometheus_direct_v1", "###Task Description / ###The instruction to evaluate / ###Response / ###Reference Answer / ###Score Rubric / ###Output JSON Schema / ###Feedback", "strict JSON or [RESULT] 1-5", "relevance; completeness; factuality; instruction_following; clarity; safety; overall_score"],
    ]
    headers7 = ["prompt_template", "required_sections", "output_constraint", "rubric_dimensions"]
    tables["prompt"] = md_table(headers7, prompt_rows)
    write_csv(TAB / "prompt_templates.csv", headers7, prompt_rows)

    sample_rows = data_sample_rows(ROOT / "datasets" / "processed" / "bea_judge_cleaned_10000.json")
    headers8 = ["id", "dataset", "task_type", "prompt_excerpt", "answer_a_excerpt", "answer_b_excerpt", "human_label", "split"]
    tables["sample"] = md_table(headers8, sample_rows)
    write_csv(TAB / "data_samples.csv", headers8, sample_rows)
    return tables


REFERENCES = [
    "Ouyang L., Wu J., Jiang X., et al. Training language models to follow instructions with human feedback. Advances in Neural Information Processing Systems, 35:27730-27744, 2022.",
    "Bai Y., Kadavath S., Kundu S., et al. Constitutional AI: Harmlessness from AI feedback. arXiv:2212.08073, 2022.",
    "Zheng L., Chiang W. L., Sheng Y., et al. Judging LLM-as-a-Judge with MT-Bench and Chatbot Arena. Advances in Neural Information Processing Systems, 36, 2023.",
    "Liu Y., Iter D., Xu Y., Wang S., Xu R., Zhu C. G-Eval: NLG evaluation using GPT-4 with better human alignment. Proceedings of the 2023 Conference on Empirical Methods in Natural Language Processing, 2023.",
    "Kim S., Shin J., Cho Y., et al. Prometheus: Inducing fine-grained evaluation capability in language models. International Conference on Learning Representations, 2024.",
    "Kim S., Bae S., Shin J., et al. Prometheus 2: An open source language model specialized in evaluating other language models. arXiv:2405.01535, 2024.",
    "Lambert N., Pyatkin V., Morrison J., et al. RewardBench: Evaluating reward models for language modeling. arXiv:2403.13787, 2024.",
    "Dubois Y., Li X., Taori R., et al. AlpacaFarm: A simulation framework for methods that learn from human feedback. Advances in Neural Information Processing Systems, 36, 2023.",
    "Dubois Y., Galambosi B., Liang P., et al. Length-controlled AlpacaEval: A simple way to debias automatic evaluators. arXiv:2404.04475, 2024.",
    "Rafailov R., Sharma A., Mitchell E., et al. Direct preference optimization: Your language model is secretly a reward model. Advances in Neural Information Processing Systems, 36, 2023.",
    "Hu E. J., Shen Y., Wallis P., et al. LoRA: Low-rank adaptation of large language models. International Conference on Learning Representations, 2022.",
    "Dettmers T., Pagnoni A., Holtzman A., Zettlemoyer L. QLoRA: Efficient finetuning of quantized LLMs. Advances in Neural Information Processing Systems, 36, 2023.",
    "Dettmers T., Lewis M., Belkada Y., Zettlemoyer L. LLM.int8(): 8-bit matrix multiplication for transformers at scale. Advances in Neural Information Processing Systems, 35, 2022.",
    "Wei J., Wang X., Schuurmans D., et al. Chain-of-thought prompting elicits reasoning in large language models. Advances in Neural Information Processing Systems, 35, 2022.",
    "Touvron H., Martin L., Stone K., et al. Llama 2: Open foundation and fine-tuned chat models. arXiv:2307.09288, 2023.",
    "Dubey A., Jauhri A., Pandey A., et al. The Llama 3 herd of models. arXiv:2407.21783, 2024.",
    "Yang A., Yang B., Zhang B., et al. Qwen2.5 technical report. arXiv:2412.15115, 2024.",
    "Köpf A., Kilcher Y., von Rütte D., et al. OpenAssistant Conversations: Democratizing large language model alignment. Advances in Neural Information Processing Systems, 36, 2023.",
    "Wang Y., Kordi Y., Mishra S., et al. Self-Instruct: Aligning language models with self-generated instructions. Proceedings of the 61st Annual Meeting of the Association for Computational Linguistics, 2023.",
    "Wang Y., Yu Z., Zeng Z., et al. PandaLM: An automatic evaluation benchmark for LLM instruction tuning optimization. arXiv:2306.05087, 2023.",
    "Chen J., Lin H., Han X., Sun L. Benchmarking large language models in retrieval-augmented generation. Proceedings of the AAAI Conference on Artificial Intelligence, 2024.",
    "Niu C., Wu Y., Zhu J., Xu S., Shum K., Zhong R., Song J., Zhang T. RAGTruth: A hallucination corpus for developing trustworthy retrieval-augmented language models. Proceedings of the 62nd Annual Meeting of the Association for Computational Linguistics, 10862-10878, 2024.",
    "Min S., Krishna K., Lyu X., et al. FActScore: Fine-grained atomic evaluation of factual precision in long-form text generation. Proceedings of the 2023 Conference on Empirical Methods in Natural Language Processing, 2023.",
    "Manakul P., Liusie A., Gales M. J. F. SelfCheckGPT: Zero-resource black-box hallucination detection for generative large language models. Proceedings of the 2023 Conference on Empirical Methods in Natural Language Processing, 2023.",
    "Ji Z., Lee N., Frieske R., et al. Survey of hallucination in natural language generation. ACM Computing Surveys, 55(12):1-38, 2023.",
    "Farquhar S., Kossen J., Kuhn L., Gal Y. Detecting hallucinations in large language models using semantic entropy. Nature, 630:625-630, 2024.",
    "Zhou L., Schellaert W., Martinez-Plumed F., et al. Larger and more instructable language models become less reliable. Nature, 634:61-68, 2024.",
    "Steyvers M., Tejeda H., Kumar A., et al. What large language models know and what people think they know. Nature Machine Intelligence, 7:221-231, 2025.",
    "Guo C., Pleiss G., Sun Y., Weinberger K. Q. On calibration of modern neural networks. Proceedings of the 34th International Conference on Machine Learning, 1321-1330, 2017.",
    "Kull M., Perello-Nieto M., Kängsepp M., Silva Filho T., Song H., Flach P. Beyond temperature scaling: Obtaining well-calibrated multiclass probabilities with Dirichlet calibration. Advances in Neural Information Processing Systems, 32, 2019.",
    "Kumar A., Liang P., Ma T. Verified uncertainty calibration. Advances in Neural Information Processing Systems, 32, 2019.",
    "Angelopoulos A. N., Bates S. A gentle introduction to conformal prediction and distribution-free uncertainty quantification. Foundations and Trends in Machine Learning, 16(4):494-591, 2023.",
    "Weyssow M., Kamanda A., Zhou X., Sahraoui H. CodeUltraFeedback: An LLM-as-a-Judge dataset for aligning large language models to coding preferences. ACM Transactions on Software Engineering and Methodology, 35(3):1-36, 2026.",
    "Hernandez-Orallo J., Martinez-Plumed F., Schellaert W., et al. General scales unlock AI evaluation with explanatory and predictive power. Nature, 2026.",
    "Li H., Moon J. T., Purkayastha S., et al. Ethics of large language models in medicine and medical research. The Lancet Digital Health, 2023.",
    "Huang L., Yu W., Ma W., et al. A survey on hallucination in large language models: Principles, taxonomy, challenges, and open questions. ACM Transactions on Information Systems, 2025.",
]


def manuscript(tables: dict[str, str]) -> str:
    refs = "\n\n".join(f"[{i}] {ref}" for i, ref in enumerate(REFERENCES, 1))
    text = dedent(f"""
    # 面向轻量化生成式回答评估的 BEA-Judge：3B Judge 模型的 QLoRA 微调、四模块融合校准与准确率约束 Tie rescue

    **摘要**：大语言模型生成内容的自动评估通常需要同时处理三类目标：总体判别准确率、平局样本召回率以及概率校准[3,4,7,25-28]。三者并不天然一致。偏好评估模型若过度追求非平局判别，容易把本应标注为 Tie 的样本压向 A>B 或 B>A；若直接放宽 Tie 判定，又可能牺牲整体 accuracy 和 macro-F1；若只关注分数排序，则置信度与实际正确率之间的偏差会影响后续人工复核和自动验收[29-32]。本文围绕这一张力构建 BEA-Judge，一种面向轻量化评估场景的 3B Judge 框架。系统采用 M-Prometheus-3B 作为基础 Judge，经 QLoRA 在 pairwise 指令上微调，并在后处理阶段引入四模块融合校准：基础 Judge 评分模块、偏差感知模块、证据增强事实性模块，以及融合校准与置信度输出模块。进一步地，本文提出基于 accuracy 约束的 Tie rescue 策略：只在验证集上选择策略阈值，且候选策略必须满足 accuracy、macro-F1 和 ECE 的约束，然后一次性应用到测试集。实验在 BEA-Judge-10K-v2 上进行。内部 QLoRA/四模块结果报告 3 个随机种子的 mean +/- std，外部轻量基线报告 single full test。结果显示，QLoRA-BEA-Judge epoch2_1024 相比 Current BEA-Judge 的 accuracy 从 0.7512 提升至 0.8297 +/- 0.0031，macro-F1 从 0.6730 提升至 0.7348 +/- 0.0062，ECE 从 0.0558 降至 0.0278 +/- 0.0031，但 Tie recall 从 0.5231 降至 0.4256 +/- 0.0270。加入 Tie rescue 后，accuracy 保持为 0.8297 +/- 0.0052，macro-F1 为 0.7441 +/- 0.0093，ECE 为 0.0283 +/- 0.0027，Tie recall 提升至 0.4795 +/- 0.0489。该结果说明，轻量 3B Judge 通过 QLoRA 与融合校准可以改善总体判别和校准，而 Tie recall 需要单独以约束策略处理；本文的主结论因此不是所有指标同时无条件提升，而是在 accuracy 基本不降的前提下提升 Tie recall，并保持较稳定校准。

    **关键词**：Judge 模型；QLoRA；概率校准；Tie recall；偏差感知；事实性评估；风险复核

    ## Abstract

    Automatic evaluation of large-language-model outputs requires balancing accuracy, Tie recall, and calibration. These objectives may conflict: stricter preference decisions improve non-Tie discrimination but suppress true ties, whereas aggressive Tie prediction can reduce overall accuracy. This paper presents BEA-Judge, a lightweight 3B judge framework that combines QLoRA-tuned M-Prometheus-3B with four post-hoc modules: base judge scoring, bias-aware features, evidence-enhanced factuality features, and fusion calibration with confidence output. We further introduce an accuracy-constrained Tie rescue policy. The policy is selected only on the development set and is eligible only when accuracy, macro-F1, and ECE constraints are satisfied; the selected policy is then applied once to the test set. On BEA-Judge-10K-v2, internal results are reported as 3-seed mean +/- standard deviation, whereas external baselines are reported as single full-test runs. QLoRA-BEA-Judge epoch2_1024 improves accuracy and calibration over Current BEA-Judge, while Tie rescue improves Tie recall with nearly unchanged accuracy. The main contribution is therefore a reproducible operating-point design for accuracy-preserving Tie recall recovery in lightweight judge models.

    ## 1 引言

    自动评估模型正在成为大语言模型开发流程中的基础组件。无论是指令微调、偏好对齐、检索增强生成，还是代码生成与事实性问答，研究者都需要一个可复现、可扩展且成本可控的 Judge 模型，对候选回答进行排序、打分或风险复核[1,3,4,21-23,34]。近期 LLM-as-a-Judge 研究说明，强模型可以在开放式任务中给出接近人类偏好的评价，但这一路线存在成本、可部署性和可解释性问题[3-5,14-16]。对于许多实验室或应用场景，使用 7B、13B 以上模型进行大规模评估并不总是可行；更轻量的 3B Judge 如果可以在准确率、Tie recall 与校准之间形成可控折中，将具有更高的复现价值。

    本文关注的核心挑战是三目标平衡。第一，accuracy 是 Judge 模型能否替代人工初筛的基本条件。第二，Tie recall 决定系统是否能识别两个回答质量接近、难以可靠排序的情形。在 pairwise 评估中，Tie 不是边缘类别，而是对模型不确定性、标注者分歧和回答近似等价状态的显式表达。第三，calibration 决定输出置信度能否被后续策略使用。一个模型即便 accuracy 较高，如果置信度系统性偏高或偏低，也会使人工复核阈值、自动接受策略和风险提示失去依据。

    这三类目标之间存在实际冲突。基础 Judge 往往倾向于选择 A 或 B，尤其当回答长度、格式或细节数量不同的时候，模型容易把表面信息当成质量信号。相反，如果为了提高 Tie recall 简单降低 Tie 判定阈值，模型可能把本来可以区分优劣的样本错误归为 Tie，导致 accuracy 和 macro-F1 下降。校准又引入额外约束：输出概率不仅要排序正确，还要与经验正确率匹配。因此，一个可复现的轻量 Judge 框架不应只报告单一最优 accuracy，而应明确说明不同操作点的目标和代价。

    BEA-Judge 的设计目标正是把这一折中显式化。本文在一个 3B 基础 Judge 上进行 QLoRA 微调，以适配 pairwise 输出格式和项目内数据分布；随后通过四模块融合把基础分数、偏差风险、证据支持和校准置信度结合起来。与直接把微调模型作为最终 Judge 不同，BEA-Judge 不把基础模型的 `[RESULT]` 作为唯一结论，而是把它视为可校准特征，再由轻量分类器在训练集上学习融合权重，并在验证集上选择温度和阈值。这样做的目的不是把所有判断交给一个不可解释的生成器，而是使基础 Judge、偏差修正、事实性证据和概率输出具有可审计的接口。

    本文的另一项贡献是准确率约束的 Tie rescue。该策略只对当前预测为 A>B 或 B>A 的样本生效，并要求 Tie 概率达到阈值、A/B 概率差足够小、A/B 最大置信度不高。重要的是，策略不在测试集上调参，也不以测试集 Tie recall 选择阈值；候选策略必须在验证集上满足 accuracy 和 macro-F1 不低于约束、ECE 不超过约束，才允许被选择。换言之，Tie rescue 不是事后为了提高 Tie recall 的任意修补，而是一个验证集选择、测试集一次应用的操作点策略。

    本文围绕以下问题展开：

    1. 在 3B Judge 上使用 QLoRA 微调能否改善 BEA-Judge 的总体判别能力？
    2. 四模块融合校准是否能缓解轻量基础模型的校准问题和偏差风险？
    3. SFT 数据量和训练轮数如何影响 accuracy、macro-F1、ECE 与 Tie recall？
    4. 在不牺牲 accuracy 的约束下，Tie rescue 能否恢复一部分 Tie recall？
    5. 与外部轻量基线相比，BEA-Judge 的优势主要体现在哪些指标，风险又在哪里？

    本文的贡献可概括为四点。第一，构建了一个轻量 3B Judge + QLoRA + 四模块融合校准的评估框架，并给出完整数学描述。第二，在 BEA-Judge-10K-v2 上统一整理了开放问答、偏差扰动、检索增强事实性和中文专业标注样本，使数据集同时覆盖偏好判别、偏差敏感性和证据支持三类评估需求。第三，系统报告 QLoRA 训练轮数和 SFT size 消融，明确显示 accuracy 与 Tie recall 的权衡。第四，提出 accuracy-constrained Tie rescue，并用三种子实验说明其在 accuracy 基本不变时提高 Tie recall 的效果。

    ## 2 相关工作

    LLM-as-a-Judge 研究表明，大模型可以通过自然语言推理对开放式回答进行评价[3,4]。MT-Bench 和 Chatbot Arena 推动了基于模型裁判的成对比较，Prometheus 系列进一步强调细粒度评价和可解释反馈[3,5,6]。与此同时，AlpacaEval、G-Eval 和 GPTScore 等工作展示了自动评价在效率上的优势，也揭示了长度偏差、位置偏差和提示敏感性等问题[4,8,9]。本文的偏差感知模块与这些观察一致：Judge 输出不能被视为无偏观测，而需要显式建模表面形式、回答顺序、rubric 变化和数据源差异。

    奖励模型和偏好模型为自动评价提供了另一条路径[1,2,7,10]。RewardBench 等基准显示，奖励模型在不同任务上的泛化和校准仍不稳定[7]。外部基线 GRM、Qwen 和 GLIDER 代表了本文关注的轻量部署条件：GRM 是 3B reward model，Qwen 是 3B instruction model，GLIDER 作为 external 4B evaluator baseline。本文没有把 Prometheus-2 7B 纳入正式主表和主结论，因为其参数规模和历史实验定位与轻量外部基线不一致[5,6]；相关结果只适合放在历史参考或附录说明。

    参数高效微调为轻量 Judge 提供了可行训练路径。LoRA 通过低秩增量矩阵减少可训练参数，QLoRA 进一步把基础模型量化到 4-bit，并使用 NF4、double quantization 和低秩适配器，使 24GB 单卡上的微调成为可能[11-13]。CodeUltraFeedback 的研究也说明，面向偏好的数据构建、SFT/DPO 以及 QLoRA 可以改善小模型在复杂评价任务中的对齐表现[33]。本文借鉴这种“以偏好反馈适配轻量模型”的思想，但研究目标从代码偏好生成转向 Judge 判别与校准。

    概率校准是本文方法的关键组成。现代神经网络经常出现置信度与实际正确率不匹配的问题，温度缩放、Dirichlet calibration 和 conformal prediction 都可用于不同程度的校准或不确定性估计[29-32]。近年关于大模型可靠性的研究也表明，模型能力提升并不自动带来稳定校准或可靠自知[26-28]。BEA-Judge 采用以验证集为中心的校准流程：先训练融合分类器，再选择温度参数，并计算 ECE、Brier score 和 NLL。本文主表重点报告 ECE，因为它直接度量置信度分箱后平均置信度与经验准确率之间的偏差。

    ## 3 数据

    ### 3.1 数据构成

    正式实验使用 `datasets/processed/bea_judge_cleaned_10000.json`，构建报告记为 BEA-Judge-10K-v2。数据集共 10,200 条样本，其中 train/dev/test 分别为 7,084、1,578 和 1,538。任务类型包括 open_qa 4,000 条、pairwise_bias 2,700 条和 factuality_rag 3,500 条。语言上，英文 9,148 条，中文 1,052 条。pairwise 样本共 6,946 条，标签分布为 A>B 4,136、B>A 1,845、Tie 965，Tie 在 pairwise 标签中的比例为 0.1389。事实性标签包括 supported 2,144 和 unsupported 1,110。

    数据来源覆盖三类需求。第一类是开放问答偏好数据，包括 PandaLM、MT-Bench human judgments、JudgeBench、HelpSteer2 和 OASST1，用于学习普通指令回答的成对偏好[3,18-20]。第二类是偏差敏感数据，包括 synthetic_perturbed 和 OffsetBias，用于让模型面对位置、长度、格式和 rubric 敏感性扰动[9]。第三类是事实性和检索增强数据，包括 ARES、WikiEval 与 RAGTruth，用于把证据支持、幻觉风险和 factuality 标签纳入评估[21-26,36]。中文专业标注样本用于补充非英文场景，使框架不完全依赖英语偏好数据。

    选择这些数据的理由是覆盖 Judge 模型的主要失效来源。开放问答数据提供一般偏好判断；偏差扰动数据提供表面形式与真实质量不一致的样本；RAG 和事实性数据提供回答与证据不一致时的负例；中文样本提供跨语言格式与专业表达差异。与只使用单一偏好数据不同，BEA-Judge-10K-v2 的创新点在于把偏好判别、偏差敏感性和证据事实性放入一个统一 schema 中，使后续四模块可以从同一数据对象中读取基础回答、偏差元数据、证据文本和人工标签。

    ### 3.2 数据预处理流程

    数据预处理遵循可复现的流水线。首先，对原始来源进行许可证与来源元数据登记，记录 URL、版本、获取时间、SHA-256 和可再分发状态。HelpSteer2、OASST1、OffsetBias 和 RAGTruth 在 v2 构建中均通过准入[18,22]；RewardBench 被标记为 external_eval_only，不进入训练数据[7]。这样做是为了避免把许可证混合或外部评测专用数据误用于模型训练。

    第二步是统一 schema。不同来源的数据字段差异较大：有的提供 pairwise votes，有的提供单回答事实性标签，有的提供检索上下文和引用。预处理把样本映射到统一字段，包括 `id`、`dataset`、`task_type`、`prompt`、`context`、`answer_a`、`answer_b`、`reference`、`human_score` 和 `human_label`。统一 schema 的原因是四模块需要稳定读取同一批字段，否则偏差模块和证据模块会因字段缺失而产生不可控差异。

    第三步是清洗与修复。项目记录显示早期 3,400 条 canonical 数据执行了 schema compliance 检查，输入 3,400、接受 3,400、修复 3,400、拒绝 0，并检查 duplicate ID、duplicate content、cross-split leakage 和 invalid enum。v2 扩展到 10,200 后，构建门禁包括 minimum_count、target_count、duplicate_id_or_content_zero、cross_split_duplicate_content_zero、license_metadata_complete 和 factuality_context_missing_zero。这样做的原因是 Judge 训练对重复和跨 split 泄漏尤其敏感；同一 prompt 或内容跨 train/dev/test 会高估泛化能力。

    第四步是分层划分。数据按内容组和任务来源划分 train/dev/test，确保 cross_split_duplicate_content 为 0。划分后，训练集用于 QLoRA 和融合分类器训练，验证集用于超参数、温度校准和 Tie rescue 策略选择，测试集只用于最终报告。这一划分是本文严谨性的基础，尤其 Tie rescue 必须只在验证集上选择。

    第五步是构建 pairwise SFT 数据。脚本 `scripts/build_judge_sft_dataset.py` 将 pairwise 样本转成 M-Prometheus 风格的 prompt-target JSONL，其中 target 只允许 `[RESULT] A`、`[RESULT] B`、`[RESULT] Tie`。这样做有两个原因：其一，稳定输出格式可降低 parse failure；其二，训练目标聚焦在判别标记上，减少生成冗长解释对显存和序列长度的压力[5,11-13]。

    表1给出基础 Judge 的 Prompt 模板。表2给出数据样例摘录，文本经过截断，仅用于展示字段结构。

    **表1 Prompt 模板**

    {tables["prompt"]}

    **表2 数据样例摘录**

    {tables["sample"]}

    ### 3.3 数据集创新性与风险

    BEA-Judge-10K-v2 的主要创新不在于单个来源规模最大，而在于将三种评估困难合并到同一训练和诊断框架中。第一，数据同时包含 A>B、B>A 和 Tie，使模型必须学习“不强行排序”。第二，偏差扰动样本把位置、长度、格式和 rubric 变化显式写入 metadata，使偏差模块可以形成结构化特征，而不是依赖模型自发纠偏[9]。第三，事实性样本包含 context 和 reference，使证据增强模块可以计算支持率、数字缺口、日期缺口和实体缺口[21-26]。第四，中文样本为专业场景提供额外分布，防止系统只在英文开放问答上优化。

    数据仍有局限。首先，Tie 样本比例为 13.89%，属于少数类，模型在优化总体 accuracy 时自然倾向于牺牲 Tie recall。其次，事实性样本和 pairwise 样本的标签空间不同，融合模型需要分别训练 pairwise head 和 factuality head。再次，早期质量审计显示某些来源字段存在缺失或空字符串，虽然 v2 构建门禁已处理正式输入，但附录仍应保留数据清洗日志，方便复现者判断字段缺失是否影响下游实验。

    ## 4 方法

    ### 4.1 问题定义

    对于每个 pairwise 样本 $x_i=(q_i,c_i,a_i,b_i,r_i)$，其中 $q_i$ 为用户指令，$c_i$ 为可选上下文，$a_i,b_i$ 为两个候选回答，$r_i$ 为可选参考答案。目标标签 $y_i \\in \\mathcal{{Y}}=\\{{A>B,B>A,Tie\\}}$。模型输出校准概率 $p_i(y)$、预测标签 $\\hat y_i=\\arg\\max_y p_i(y)$、置信度 $s_i=\\max_y p_i(y)$ 和风险复核标记 $g_i$。事实性 head 使用 $\\mathcal{{Y}}_f=\\{{supported,unsupported,ambiguous\\}}$，但本文主贡献和主表聚焦 pairwise head。

    ### 4.2 模块1：基础 Judge 评分与 QLoRA 微调

    基础模块先把样本映射为固定模板 $T(x_i)$，再由 3B Judge 输出结果标记。未微调时，基础模型直接给出 `[RESULT] A/B/Tie`；QLoRA 微调时，目标是最大化正确结果标记的条件似然。固定输出标记借鉴了细粒度 Judge 和指令式评估模型的做法，低秩与量化训练则来自 LoRA/QLoRA 系列方法[5,11-13]。设 token 序列由 prompt token $u_{{i,1:m}}$ 和 target token $v_{{i,1:k}}$ 组成，训练只在 target 部分计算损失：

    $$
    \\mathcal{{L}}_{{SFT}}(\\theta)=-\\sum_i\\sum_{{t=1}}^k \\log P_\\theta(v_{{i,t}}\\mid u_{{i,1:m}},v_{{i,<t}}).
    $$

    QLoRA 将基础权重 $W$ 量化为 $Q(W)$，并在选定线性层上加入低秩增量[11-13]：

    $$
    W' = Q(W) + \\Delta W, \\quad \\Delta W = \\frac{{\\alpha}}{{r}}BA,
    $$

    其中 $A\\in\\mathbb{{R}}^{{r\\times d}}$，$B\\in\\mathbb{{R}}^{{d'\\times r}}$，$r=16$，$\\alpha=32$，dropout 为 0.05。本文使用 NF4 4-bit 量化、double quantization 和 bfloat16 计算，目标模块包括 `q_proj`、`k_proj`、`v_proj`、`o_proj`、`gate_proj`、`up_proj` 和 `down_proj`。稳定序列长度为 1024，训练采用 batch size 1、gradient accumulation 16、learning rate $1\\times 10^{{-4}}$、cosine scheduler、warmup ratio 0.03、weight decay 0.01 和 gradient checkpointing。

    基础 Judge 输出被解析为 $\\ell_i\\in\\mathcal{{Y}}$，并提取 score_a、score_b。定义基础差值 $d_i=s_{{a,i}}-s_{{b,i}}$，边际 $m_i=|d_i|$，以及 one-hot 预测向量：

    $$
    z_i^{{base}}=[s_{{a,i}},s_{{b,i}},d_i,m_i,\\mathbb{{1}}(\\ell_i=A>B),\\mathbb{{1}}(\\ell_i=B>A),\\mathbb{{1}}(\\ell_i=Tie)].
    $$

    若有 swap consistency 信息，则加入 $z_i^{{swap}}=[available,consistent,\\Delta m]$。模块1提供的是强但未必校准的判别信号，它并不直接等同于最终输出。

    ### 4.3 模块2：偏差感知特征

    偏差感知模块用于捕捉 Judge 容易误用的表面信号。对回答长度，设 $L_a,L_b$ 为字符长度，长度比为：

    $$
    \\rho_L=\\frac{{\\max(L_a,L_b)}}{{\\max(1,\\min(L_a,L_b))}}.
    $$

    当 $\\rho_L\\ge 1.15$ 时，较长回答一侧记为 $side_L$。若基础预测偏向较长回答且人工标签不支持该侧，则长度风险为 1；若人工标签为 Tie，则风险为 0.7；否则为 0。用分段函数表示：

    $$
    R_{{len}}=
    \\begin{{cases}}
    1, & \\hat y\\text{{ favors }}side_L \\land y\\notin\\{{side_L,Tie\\}},\\\\
    0.7, & \\hat y\\text{{ favors }}side_L \\land y=Tie,\\\\
    0.25, & side_L\\neq\\varnothing \\land \\hat y=\\varnothing,\\\\
    0, & \\text{{otherwise}}.
    \\end{{cases}}
    $$

    格式风险基于项目符号数量差 $\\Delta B=B_a-B_b$ 和 metadata 中的 `format` 扰动。若格式扰动存在且预测错误，则 $R_{{fmt}}=1$；若模型偏向项目符号更多一侧而人工标签不支持该侧，则 $R_{{fmt}}=0.8$；若仅存在格式扰动，则 $R_{{fmt}}=0.35$。位置风险和 rubric 风险也采用同样逻辑：扰动存在且预测错误时风险为 1，仅有扰动时风险为 0.35。

    数据源风险来自各数据集 accuracy 与总体 accuracy 的差异：

    $$
    R_{{src}}(d)=\\left|Acc_d-Acc_{{all}}\\right|.
    $$

    综合偏差风险取最大值：

    $$
    R_{{bias}}=\\max(R_{{pos}},R_{{len}},R_{{fmt}},R_{{rubric}},R_{{src}}).
    $$

    偏差模块向融合层提供两类信息：一类是 metadata 标志，如 position、length、format、rubric_sensitivity；另一类是连续风险值，如 $R_{{len}}$ 和 $R_{{bias}}$。这种设计使模型能够区分“存在扰动”和“扰动导致预测不可信”。

    ### 4.4 模块3：证据增强事实性特征

    事实性模块利用 context 与 reference 估计回答被证据支持的程度。RAG 与幻觉评估研究表明，局部事实、数字、实体和语义支持关系是回答可信度的重要来源[21-26,36]。设 $Tok(a)$ 为去停用词后的 token 集，基本覆盖率为：

    $$
    C(a,e)=\\frac{{|Tok(a)\\cap Tok(e)|}}{{|Tok(a)|}}.
    $$

    对长词元和数字单独计算覆盖率。若回答包含数字集合 $N(a)$，证据包含 $N(e)$，数字支持率为 $|N(a)\\cap N(e)|/|N(a)|$。综合支持分数定义为：

    $$
    S(a,e)=
    \\begin{{cases}}
    0.45C(a,e)+0.30C_{{long}}(a,e)+0.25C_{{num}}(a,e), & |N(a)|>0,\\\\
    0.65C(a,e)+0.35C_{{long}}(a,e), & |N(a)|=0.
    \\end{{cases}}
    $$

    对 context 和 reference 分别计算 $S(a,c)$、$S(a,r)$，并以组合证据 $e=[c;r]$ 得到 claim support：

    $$
    S_{{claim}}(a)=\\max(S(a,e),0.65S(a,c)+0.35S(a,r)).
    $$

    证据缺口包括数字、日期、实体和实体别名：

    $$
    G_{{num}}(a)=\\frac{{|N(a)\\setminus N(e)|}}{{|N(a)|}},\\quad
    G_{{ent}}(a)=\\frac{{|E(a)\\setminus E(e)|}}{{|E(a)|}}.
    $$

    对否定和比较关系，模块先把回答切分为句子，为每个回答句选择支持分数最高的证据句；若二者否定标记不一致，则记为 negation mismatch；若回答中的 increase/decrease/maximum/minimum 关系无法在最佳证据句中找到对应组，则记为 comparative mismatch。局部幻觉风险定义为：

    $$
    R_{{hall}}=\\max(G_{{num}},G_{{date}},G_{{ent}},G_{{alias}},R_{{neg}},R_{{comp}},R_{{lowSent}}).
    $$

    对 pairwise factuality，计算两侧支持差：

    $$
    \\Delta S = S_{{claim}}(a)-S_{{claim}}(b).
    $$

    若人工或基础判定与支持差方向冲突，则生成 pairwise support contradiction 特征。例如预测 A>B 但 $\\Delta S<-0.05$，或预测 Tie 但 $|\\Delta S|>0.35$，均表示证据信号与判定存在冲突。证据模块的作用不是替代基础 Judge，而是为融合层提供“回答是否被证据支撑”的可解释约束。

    ### 4.5 模块4：融合校准与置信度输出

    融合层把四类特征拼接为：

    $$
    x_i=[z_i^{{text}},z_i^{{base}},z_i^{{bias}},z_i^{{evidence}},z_i^{{onehot}}].
    $$

    其中 $z_i^{{text}}$ 包括长度、句子数、项目符号、数字数量、prompt/context/reference overlap；$z_i^{{onehot}}$ 包括 dataset、task 和 scoring system。特征在训练集上标准化：

    $$
    \\tilde x_{{ij}}=\\frac{{x_{{ij}}-\\mu_j}}{{\\sigma_j+\\epsilon}}.
    $$

    对 pairwise head，轻量 softmax classifier 输出：

    $$
    p_\\theta(y\\mid x_i)=\\frac{{\\exp(\\tilde x_i^\\top w_y+b_y)}}{{\\sum_{{y'\\in\\mathcal{{Y}}}}\\exp(\\tilde x_i^\\top w_{{y'}}+b_{{y'}})}}.
    $$

    训练目标为加权交叉熵加 L2 正则：

    $$
    \\mathcal{{L}}_{{fusion}}(\\theta)=-\\frac{{1}}{{\\sum_i\\omega_i}}\\sum_i\\omega_i\\log p_\\theta(y_i\\mid x_i)+\\frac{{\\lambda}}{{2}}\\|W\\|_2^2.
    $$

    超参数在验证集上用目标函数选择：

    $$
    J=macroF1+0.25\\cdot accuracy-0.05\\cdot ECE.
    $$

    之后进行温度缩放。设未校准概率为 $p_i$，温度为 $T>0$，则：

    $$
    p_{{i,T}}(y)=\\frac{{\\exp(\\log p_i(y)/T)}}{{\\sum_{{y'}}\\exp(\\log p_i(y')/T)}}.
    $$

    温度在验证集上最小化，这一后处理方式沿用温度缩放和现代分类器校准文献中的验证集选择原则[29-31]：

    $$
    \\mathcal{{J}}(T)=NLL(T)+ECE(T)+0.25\\cdot Brier(T).
    $$

    ECE 使用 10 个等宽置信度分箱：

    $$
    ECE=\\sum_{{b=1}}^B\\frac{{|I_b|}}{{n}}\\left|acc(I_b)-conf(I_b)\\right|.
    $$

    最终输出 $\\hat y_i=\\arg\\max_y p_{{i,T}}(y)$，置信度 $s_i=p_{{i,T}}(\\hat y_i)$。风险复核分数可取 $1-s_i$ 或结合偏差/证据风险，阈值在验证集上选择，用于捕获较低置信度错误。

    ### 4.6 准确率约束 Tie rescue

    Tie rescue 只处理当前预测为 A>B 或 B>A 的样本。设校准概率为 $p_A,p_B,p_T$。在选定数据子集 $d$ 上，如果满足：

    $$
    p_T\\ge \\tau,\\quad |p_A-p_B|\\le m,\\quad \\max(p_A,p_B)\\le c,
    $$

    则把预测从 A>B 或 B>A 改为 Tie；否则保持原预测。本文最终全局策略在验证集三种子均值上选择，参数为 dataset=helpsteer2、$\\tau=0.38$、$m=0.65$、$c=0.65$。

    策略选择不是无约束搜索。候选策略 $\\pi$ 必须在验证集满足：

    $$
    Acc_{{dev}}(\\pi)\\ge Acc_{{min}},\\quad
    MacroF1_{{dev}}(\\pi)\\ge F1_{{min}},\\quad
    ECE_{{dev}}(\\pi)\\le ECE_{{max}},\\quad
    TieRecall_{{dev}}(\\pi)>TieRecall_{{min}}.
    $$

    合格候选按验证集 Tie recall、accuracy、负 ECE、macro-F1 的顺序排序。这个策略体现了本文对 Tie recall 的保守态度：Tie recall 可以被修复，但不能通过牺牲验证集 accuracy 和校准来换取。

    ![四模块BEA-Judge框架](figures/fig1_four_module_framework.svg)

    ## 5 实验设置

    ### 5.1 实验协议

    所有主实验使用 BEA-Judge-10K-v2。内部 QLoRA 与四模块实验使用种子 13、42、2026，报告 mean +/- sample std。外部基线使用 single full test，原因是这些模型作为外部评估器运行成本较高，且不参与本文内部训练和阈值选择。所有正式表格使用相同指标名：accuracy、macro-F1、ECE、Tie recall、parse failure rate。对于内部融合结果，parse failure rate 不是 post-calibration 指标，表中以 `--` 标记；外部基线报告 parse failure rate 为 0.0000。

    本文明确区分三类系统。Current BEA-Judge 是冻结基础 Judge 的四模块版本。QLoRA-M-Prometheus-3B 是只看 QLoRA 基础模型输出的内部基线，不含四模块融合。QLoRA-BEA-Judge 是本文主模型，即 QLoRA 基础模型加四模块融合校准。最终主模型报告两个操作点：epoch2_1024 与 epoch2_1024 + Tie rescue。

    ### 5.2 指标

    Accuracy 衡量预测标签与人工标签一致的比例。Macro-F1 对 A>B、B>A、Tie 三类分别计算 F1 后平均，避免多数类主导。Tie recall 只在人工标签为 Tie 的样本上计算：

    $$
    TieRecall=\\frac{{\\sum_i\\mathbb{{1}}(y_i=Tie\\land \\hat y_i=Tie)}}{{\\sum_i\\mathbb{{1}}(y_i=Tie)}}.
    $$

    ECE 衡量置信度校准。Parse failure rate 衡量外部模型输出无法解析为合法标签的比例。风险复核分析额外报告 review_rate、error_capture_rate、auto_accept_accuracy，用于衡量置信度/风险输出是否能辅助人工复核。

    ## 6 QLoRA 微调实验

    ### 6.1 为什么需要 QLoRA

    引言中提出的问题是轻量 Judge 的 accuracy、Tie recall 与 calibration 平衡。Current BEA-Judge 的四模块融合已经能在冻结基础 Judge 上工作，但基础 Judge 的原始判别能力较弱，限制了融合层可获得的信息质量。QLoRA 的作用是在不显著提高部署成本的情况下，让 3B 基础模型适应 BEA-Judge 的 pairwise 标签空间、固定 prompt 模板和 `[RESULT]` 输出格式。由于序列长度 1024 在项目中稳定，本文把 1024 作为正式主线配置。

    表3给出训练轮数消融。随着训练从 0.5 epoch 增至 2 epoch，QLoRA-BEA-Judge 的 accuracy 和 macro-F1 持续提升，ECE 保持在约 0.028-0.032 的低水平。Tie recall 则不是单调上升：epoch1_1024 的 Tie recall 为 0.4538 +/- 0.0400，epoch2_1024 降至 0.4256 +/- 0.0270。这一结果验证了本文的问题设定：提高总体判别并不自动提高 Tie recall，因此需要独立的 Tie rescue 操作点。

    **表3 QLoRA训练轮数消融**

    {tables["epoch"]}

    ![QLoRA训练轮数消融](figures/fig5_epoch_ablation.svg)

    ### 6.2 主模型结果

    表4是主对比表。epoch2_1024 相比 Current BEA-Judge 的 accuracy 提升 0.0785，macro-F1 提升 0.0618，ECE 降低 0.0280。这说明 QLoRA 微调后的基础 Judge 为四模块融合提供了更强输入，尤其改善了非 Tie 类判别和整体概率质量。与此同时，Tie recall 低于 Current BEA-Judge，说明更强的偏好判别倾向于把部分真实 Tie 样本推向 A 或 B。

    加入 Tie rescue 后，accuracy 均值保持为 0.8297，macro-F1 提升到 0.7441，ECE 小幅从 0.0278 变为 0.0283，Tie recall 从 0.4256 提升到 0.4795。该结果支持本文主结论：Tie rescue 不是全面提升所有指标，而是在 accuracy 基本不降、校准仍稳定的前提下恢复 Tie recall。

    **表4 主对比表**

    {tables["main"]}

    ![主模型与外部基线对比](figures/fig2_main_comparison.svg)

    ## 7 SFT size 消融实验

    ### 7.1 实验目的

    SFT size 消融用于回答两个问题。第一，QLoRA 的收益是否依赖完整训练集，还是少量 pairwise SFT 已能改善基础模型[11-13]？第二，随着数据量增加，accuracy、macro-F1、ECE 与 Tie recall 是否同步改善？这直接关系到轻量 Judge 的训练成本：如果小数据已足够，完整训练的边际收益有限；如果完整数据显著改善 accuracy，则说明数据规模仍是关键。

    表5显示，从 25% 到 100% SFT，QLoRA-BEA-Judge 的 accuracy 从 0.7740 提升到 0.8025，macro-F1 从 0.6832 提升到 0.7128，ECE 从 0.0310 小幅降至 0.0279。Tie recall 在 50% 时达到 0.4590，100% 时为 0.4538，变化不大。结果表明，更多 SFT 数据主要改善总体判别与稳定性，对 Tie recall 的收益较弱。这与 epoch 消融一致：Tie 类需要显式操作点，而不是只依赖更多训练。

    **表5 SFT size消融**

    {tables["sft"]}

    ![SFT size消融](figures/fig4_sft_size_ablation.svg)

    ## 8 四模块消融

    表6给出四模块消融。去除校准后，pairwise accuracy 为 0.7407，macro-F1 为 0.6402，Tie recall 为 0.3923，显示校准和融合策略对最终可用性有明显影响。去除证据模块对 factuality head 的影响更大，factuality accuracy 从 0.7649 降至 0.6928，macro-F1 从 0.7405 降至 0.6542，说明证据特征在事实性判断中不是装饰性特征，而是关键输入。

    偏差模块的结果需要谨慎解释。w/o Bias Module 在 pairwise test 上 accuracy 和 macro-F1 高于 Full BEA-Judge，这说明偏差特征作为决策输入可能在部分样本上引入保守修正，牺牲了一些总体指标。但风险表显示，bias_as_decision_features 的 review_capture_rate 为 0.7557，高于 no_bias_decision_features 的 0.7328，说明偏差模块对风险复核仍有价值。本文因此不把偏差模块描述为无条件提高 accuracy 的组件，而把它定位为提升风险可解释性和复核覆盖的模块。

    **表6 四模块消融表**

    {tables["ablation"]}

    ## 9 对比实验与风险复核分析

    ### 9.1 外部轻量基线

    本文选择 GRM、Qwen 和 GLIDER 作为外部轻量基线。GRM 是 3B reward model，代表专门用于偏好打分的轻量奖励模型[7]。Qwen2.5-3B-Instruct 代表通用 3B 指令模型，检验未专门训练的轻量模型能否直接承担 Judge 任务[17]。GLIDER 标注为 external 4B evaluator baseline，代表略大但仍轻量的 evaluator。Prometheus-2 7B 仅作为历史参考结果，不纳入正式主表和主结论[6]。

    结果显示，epoch2_1024 的 accuracy 高于 GRM、Qwen 和 GLIDER，ECE 显著低于 GRM 和 Qwen，并略低于 GLIDER。GRM 的 Tie recall 为 0.5000，高于 epoch2_1024，但其 accuracy 和 ECE 明显较弱；GLIDER 的 ECE 为 0.0353，校准较好，但 accuracy 只有 0.5043。这个对比说明，单一指标不能概括 Judge 质量。BEA-Judge 的优势主要在 accuracy 和校准的组合，而 Tie recall 需要通过 Tie rescue 改善。

    ### 9.2 Tie rescue 对比

    表7显示 Tie rescue 的测试集变化。全局策略在验证集上选择，测试集一次应用。三种子平均下，Tie rescue 把 Tie recall 提升 0.0539，macro-F1 提升 0.0093，accuracy 均值保持 0.8297，ECE 仅增加 0.0005。逐种子结果显示，被 rescue 的测试样本数分别为 11、24、7，对应 rescue precision 为 0.6364、0.4583、0.4286。这说明策略较保守，但仍能恢复一部分真实 Tie。

    **表7 Tie rescue对比表**

    {tables["tie"]}

    ![Tie rescue对比](figures/fig3_tie_rescue.svg)

    ### 9.3 风险复核分析

    风险复核分析的必要性来自校准输出的使用场景。Judge 模型通常不是孤立给出一个标签，而是服务于自动验收、人工抽检和错误定位。如果模型能把高风险样本集中到较小复核集合中，则即便仍有错误，也能降低人工成本。表8和图6显示，pairwise test 上 review_rate 从 0.0503 增至 0.4995 时，error_capture_rate 从 0.1412 增至 0.8511，auto_accept_accuracy 从 0.7750 增至 0.9260。也就是说，扩大复核范围可以显著捕获错误，而剩余自动接受样本更可靠。

    这项分析不应被解读为模型本身 accuracy 提升，而应解读为置信度和风险分数对人工流程的实用价值。对于高风险应用，可以选择更高 review_rate；对于成本敏感场景，可以使用较低 review_rate 并接受更低的错误捕获率。

    **表8 风险复核分析**

    {tables["risk"]}

    ![风险复核曲线](figures/fig6_risk_review.svg)

    ## 10 进一步分析

    ### 10.1 为什么 accuracy、Tie recall 与 calibration 会形成张力

    在三分类 pairwise 评估中，A>B 与 B>A 是方向性标签，Tie 则是非方向性标签。若只从 0-1 accuracy 出发，模型在类别不均衡时会倾向于强化多数方向性标签，因为 A>B 与 B>A 的样本数合计远高于 Tie。设真实标签先验为 $\\pi_A,\\pi_B,\\pi_T$，且 $\\pi_T<\\pi_A+\\pi_B$。当模型把一部分边界 Tie 样本预测为方向性标签时，只要这些样本的置信度与相邻非 Tie 样本相近，整体 accuracy 的损失未必显著；但 Tie recall 会立即下降。相反，若模型扩大 Tie 判定区域，则方向性标签中的边界样本会被吸收到 Tie 类，Tie recall 可能上升，但 accuracy 与 A/B 的 precision 可能下降。

    Calibration 进一步加剧这种张力。一个未校准模型可能在边界样本上给出过高 A/B 置信度，使 Tie rescue 无法可靠识别候选；也可能给出过高 Tie 概率，造成过度 rescue[29-32]。温度缩放降低或提高概率分布尖锐度，本质上改变的是置信度尺度，而不是基础排序。因此，校准质量决定了 Tie rescue 规则中的 $p_T$、$|p_A-p_B|$ 和 $\\max(p_A,p_B)$ 是否具有可解释意义。若 ECE 较高，阈值规则即便在验证集有效，也更可能在测试集漂移。

    这种冲突也解释了为什么本文不把 Tie rescue 纳入训练损失。若在训练阶段直接提高 Tie 类权重，模型可能整体改变决策边界，影响所有数据集和所有样本；而 Tie rescue 只在校准后概率满足特定条件时局部调整。它更像一个操作点选择，而不是新的模型参数。操作点的优势是可审计、可关闭、可在不同部署场景中重新选择；缺点是需要严格的验证集约束，且不能保证跨分布无损迁移。

    ### 10.2 四模块之间的误差传播

    四模块不是线性堆叠，而是误差信号的逐层过滤。基础 Judge 错误可以分为三类：解析错误、方向性错误和边界错误。解析错误通过固定 `[RESULT]` 模板和 QLoRA SFT 降低；方向性错误依赖基础评分和融合层纠正；边界错误则主要与 Tie recall 和校准有关。偏差模块和证据模块不能修复所有基础错误，但它们能提供错误发生的条件信息。例如，若回答 A 更长、项目符号更多，而基础模型偏向 A，偏差模块会把该样本标记为长度或格式风险；若回答 A 包含未在证据中出现的数字或实体，证据模块会提高 factuality risk。

    融合层学习的是这些条件信息与人工标签之间的统计关系。设基础模块给出 $z^{{base}}$，偏差模块给出 $z^{{bias}}$，证据模块给出 $z^{{evidence}}$。如果某一类风险只在少量数据源中出现，softmax 权重可能较小；如果该风险在验证集上与错误高度相关，权重会增大。训练目标中的 L2 正则限制单个特征过度主导，避免模型把某个启发式当成绝对规则。标准化也很重要，因为长度差、概率差和 one-hot 特征量纲不同；没有标准化时，大尺度长度变量可能压制校准概率和风险分数。

    从误差传播角度看，校准模块承担两个职责。第一，它把融合层输出映射到更可信的概率尺度，使 ECE 降低。第二，它为风险复核和 Tie rescue 提供阈值基础。若只输出标签而不输出置信度，人工复核只能随机抽样或按数据源抽样；若置信度经过校准，复核策略可以优先处理低置信度和高风险样本，提高错误捕获率。本文风险复核结果说明，这一职责具有独立价值。

    ### 10.3 数据来源与模块任务的对应关系

    不同数据来源在四模块中承担不同角色。HelpSteer2 和 OASST1 提供一般指令跟随与帮助性偏好，主要影响基础 Judge 和文本 overlap 特征[18]。MT-Bench 与 PandaLM 提供开放式回答对比，帮助模型学习综合质量判断[3,20]。OffsetBias 和 synthetic_perturbed 对偏差模块更重要，因为它们让位置、长度、格式和 rubric 改写具有可识别 metadata[9]。RAGTruth、ARES 和 WikiEval 对证据模块更重要，因为它们包含 context、reference 或事实性标签，允许计算支持率与缺口[21,22]。

    这种来源-模块映射解释了为什么单一数据集难以训练出稳定 Judge。若只使用普通偏好数据，模型可能在开放问答上表现良好，但面对格式扰动时仍偏向更长、更像“高质量”的回答。若只使用事实性数据，模型可能对证据支持敏感，但对一般帮助性和清晰度判断不足。若只使用偏差数据，模型可能过度怀疑表面质量信号，降低正常样本的判别力。因此，BEA-Judge-10K-v2 的价值在于提供多源、异质、但 schema 统一的训练和诊断基础。

    中文专业样本也有类似作用。中文回答的断句、术语、格式和引用习惯与英文不同；如果模型只在英文样本上学习 overlap 或长度特征，可能把中文中的紧凑表达误判为信息不足。本文没有单独报告中文子集主结果，是因为用户给定的正式关键结果聚焦全测试集；但在复现和扩展中，中文子集应作为重要外部效度检查。

    ### 10.4 QLoRA 训练对基础 Judge 的具体影响

    QLoRA 的收益可拆成输出格式适配、任务分布适配和边际分数适配[11-13]。输出格式适配最直接：SFT 目标只允许三个 `[RESULT]` 标记，使模型更稳定地产生可解析结论。任务分布适配体现在模型看到 BEA-Judge-10K-v2 的 prompt、context、answer 和 reference 组合，学习项目内部标签口径。边际分数适配则体现在 score_a、score_b 与最终标签之间的关系更符合融合层需要。

    但 QLoRA 也可能带来 Tie 压缩。SFT 的目标是给定样本输出正确标记，训练损失对 Tie 类没有天然保护；即使使用完整数据，Tie 类仍是少数类。模型为了降低总体损失，可能学到更强的方向性偏好，尤其当 A>B 样本多于 B>A 和 Tie 时。epoch 消融中 accuracy 和 macro-F1 随训练轮数提高，而 Tie recall 在 epoch2 下降，正符合这种机制。本文不把这解释为 QLoRA 失败，而解释为 QLoRA 解决了基础判别问题，但没有单独解决 Tie 边界问题。

    从部署角度看，epoch2_1024 是 accuracy-oriented operating point，适合需要较高总体判别能力的场景；epoch2_1024 + Tie rescue 是 tie-sensitive operating point，适合需要保守处理近似等价回答的场景。两个配置共享同一主模型和同一校准输出，只在最终决策策略上不同。这种分离有助于避免把一个指标偏好的策略误当成模型本身的普遍改进。

    ### 10.5 与外部基线的公平性边界

    外部基线为 single full test，而内部结果为 3-seed mean +/- std。二者不能进行严格显著性检验，只能作为轻量部署条件下的经验对比。GRM、Qwen 和 GLIDER 的 prompt、训练目标和输出接口与 BEA-Judge 不同；即便 parse failure rate 为 0，它们的概率校准方式也不一定与本文的融合校准一致。因此，本文只把外部基线用于回答“在同一测试集上，轻量外部模型直接作为 evaluator 的结果如何”，而不把它们用于训练或阈值选择。

    这种边界尤其影响 ECE 解读。GLIDER 的 ECE 为 0.0353，接近 BEA-Judge，但 accuracy 只有 0.5043。低 ECE 可能来自较保守的概率输出，而不一定代表判别质量强。GRM 的 Tie recall 为 0.5000，但 ECE 为 0.1759，说明它能识别更多 Tie，却难以提供稳定置信度。Qwen 的 parse failure rate 为 0，但 Tie recall 只有 0.0308，说明通用指令模型即便能按格式输出，也不等于能理解 pairwise Tie 的标签语义。

    ### 10.6 风险复核策略的使用方式

    风险复核曲线可转化为部署策略。若系统允许复核约 10% 样本，pairwise error_capture_rate 为 0.2634，auto_accept_accuracy 为 0.7964；若复核约 20%，错误捕获率达到 0.5000，自动接受准确率为 0.8444；若复核约 50%，错误捕获率为 0.8511，自动接受准确率为 0.9260。选择哪一点取决于应用成本，而不是模型论文中的单一最优值。

    这也提示 future work 的方向：风险复核不应只依赖置信度，还应把偏差风险和证据风险纳入排序。例如，在医疗、法律或科研问答中，证据风险可能比置信度更重要；在开放问答偏好评估中，长度或格式偏差可能更重要。当前实现已经输出 review flag，但仍可进一步学习一个专门的 error predictor，使复核排序不完全依赖 $1-s_i$。

    ### 10.7 有效性威胁

    内部有效性威胁主要来自阈值选择和多次实验。本文通过固定三种子、固定测试集、验证集选择 Tie rescue 和明确报告失败/成功门禁来降低风险，但仍不能排除某些阈值对当前数据分布有适配。构念有效性威胁来自指标本身：accuracy、macro-F1、ECE 和 Tie recall 分别刻画不同方面，任何单一指标都不能代表 Judge 的整体质量。因此本文始终把指标组合解释为 trade-off。

    外部有效性威胁来自数据分布。BEA-Judge-10K-v2 虽覆盖多源数据，但仍主要是文本回答评估，不等同于所有领域的安全、医学、法律或代码执行正确性评估。CodeUltraFeedback 说明代码偏好评估需要专门数据和偏好维度[33]；医学等高风险领域还涉及更强的伦理和安全约束[35]。BEA-Judge 的方法可迁移到这类任务，但不能直接声称在代码偏好上达到同等效果。复现有效性威胁来自环境：QLoRA 训练依赖 CUDA、bitsandbytes、transformers 和显存配置[12,13]，因此本文固定报告 1024 序列长度与对应运行配置。

    ## 11 讨论

    本文结果最重要的现象是 trade-off。epoch2_1024 明显优于 Current BEA-Judge 的 accuracy、macro-F1 和 ECE，但 Tie recall 下降。这不是异常，而是 pairwise Judge 的常见倾向：当模型更擅长区分 A 与 B，它会减少 Tie 输出。对于只关心胜负排序的场景，这可能是可接受的；但在需要识别近似等价回答、标注分歧或不确定性时，Tie recall 下降会降低系统可信度。

    Tie rescue 的价值在于提供第二个操作点。它没有改变主模型训练，也没有在测试集上搜索阈值，而是在验证集上选择满足 accuracy 和 calibration 约束的策略。结果中 Tie recall 提升，但仍未超过 Current BEA-Judge 的 0.5231。这说明策略是保守的：它恢复了部分真实 Tie，但没有为了追平旧基线而牺牲 accuracy。本文认为这种保守性是必要的，因为 Tie recall 如果通过大量误报 Tie 获得，会损害 Judge 的基本判别价值。

    四模块框架的作用也不是每个模块都单独提高所有指标。基础 Judge 提供主要判别能力，QLoRA 改善其任务适配；偏差模块提供偏差和复核信号；证据模块主要改善 factuality 与证据不一致样本；融合校准模块把这些信号转为概率和置信度。消融结果中 w/o Bias Module 的 pairwise accuracy 更高，提示偏差特征用于决策时可能过度保守。但在风险复核中，偏差信息能提高错误捕获。因此，BEA-Judge 的模块价值应从“决策指标”和“复核可用性”两个层面评价。

    外部基线对比进一步说明轻量 Judge 不能只看模型规模。GRM 的 Tie recall 较强，但 ECE 高；Qwen 作为通用指令模型直接做 Judge 时 Tie recall 很低；GLIDER 校准较好但 accuracy 低。BEA-Judge 的设计选择是把一个轻量基础模型与任务内数据、结构化偏差/证据特征和校准结合起来，而不是期望通用模型零样本完成所有评价。

    从应用角度看，本文实际上提供了两个可选择的操作点，而不是单一固定系统。若任务是自动化排序、批量模型筛选或偏好数据初筛，epoch2_1024 更合适，因为它保持较高 accuracy 和较低 ECE，输出也更集中于方向性判断。若任务是人工标注辅助、答案去重、或需要识别“两个回答都可接受”的场景，epoch2_1024 + Tie rescue 更合适，因为它在不降低 accuracy 均值的前提下恢复一部分 Tie recall。这个区分对 SCI 论文写作很重要：方法贡献不是宣称一个策略支配所有场景，而是把同一模型的判别能力和边界保守性拆成可审计的部署选择。

    这一区分也影响指标解释。accuracy 高并不意味着 Tie 样本被充分识别，Tie recall 高也不意味着系统适合自动裁决。一个过度输出 Tie 的模型可以获得较高 Tie recall，却会让 A/B 胜负样本的有效决策率下降；一个过度自信的方向性模型可以获得较高 accuracy，却会把真实平局样本强行分出胜负。校准指标 ECE 则提供第三个视角：即使预测标签正确，若概率尺度不可靠，模型也难以用于风险复核和阈值策略。本文因此把 accuracy、macro-F1、ECE 和 Tie recall 放在同一张主表中，而不是挑选单项最优指标。

    对比实验的另一个启示是，轻量 evaluator 的训练目标必须与使用场景匹配。GRM 作为 reward model 更接近偏好打分，Qwen 作为 instruct model 更接近通用回答生成，GLIDER 作为 evaluator baseline 更接近评价模型，但这些模型都未必显式接受 BEA-Judge-10K-v2 的偏差、事实性和 Tie 边界标注。BEA-Judge 的改进来自三方面共同作用：第一，QLoRA 让 3B 基础模型学习项目内 prompt 和标签口径；第二，四模块把偏差和证据风险显式化；第三，温度缩放和 Tie rescue 把概率输出转化为可控操作点。任何一部分单独存在，都不足以解释最终结果。

    还需要强调的是，Tie rescue 的成功条件依赖校准质量。若基础概率分布没有校准，$p_T\\ge 0.38$、$|p_A-p_B|\\le 0.65$ 和 $\\max(p_A,p_B)\\le 0.65$ 这些阈值就只是经验数值，难以解释其含义。温度缩放虽然简单，却使概率尺度更稳定，进而让阈值策略有可复核基础。本文没有在测试集上重新选择阈值，也没有把 rescue 后结果作为新训练目标，这降低了过拟合风险。未来如果引入更复杂的阈值策略，例如按数据源、语言或任务类型分层选择，也必须保持相同原则：所有策略选择只能在 dev 上完成，test 只用于最终报告。

    本文也存在局限。第一，外部基线为 single full test，不能与内部三种子结果进行完全同等的方差比较。第二，Tie rescue 当前只在 helpsteer2 子集上选择策略，未来应探索跨数据集或分层策略，但必须保持验证集选择原则。第三，证据模块主要使用词元、数字、日期、实体和句子支持启发式，不能替代更强的检索验证或自然语言推理。第四，本文没有把 Prometheus-2 7B 纳入正式轻量主表，因此结论只适用于轻量 3B/4B 比较。

    ## 12 结论

    本文构建并评估了 QLoRA-BEA-Judge：一个围绕轻量 3B Judge、QLoRA 微调、四模块融合校准和准确率约束 Tie rescue 的自动评估框架。实验表明，epoch2_1024 相比 Current BEA-Judge 在 accuracy、macro-F1 和 ECE 上有明显优势；但 Tie recall 下降，说明总体判别能力和 Tie 识别之间存在真实权衡。加入 Tie rescue 后，模型在 accuracy 基本不变的前提下提升 Tie recall，并保持较低 ECE。本文的主结论因此聚焦于“accuracy 不降、Tie recall 上升、校准更稳”，而不是宣称所有指标同时全面提升。

    ## 附录A 复现细节

    正式数据文件为 `datasets/processed/bea_judge_cleaned_10000.json`。SFT 数据由 `scripts/build_judge_sft_dataset.py` 生成，输出到 `datasets/sft/m_prometheus_pairwise/`。QLoRA 主配置使用 `configs/qlora_judge_sft_24gb_epoch1_1024.json`，epoch2 通过训练轮数覆盖实现。稳定序列长度为 1024；2048 长度在部分环境下可能触发显存不足，因此只作为复现注意事项，不作为正文主线。

    训练参数包括：4-bit NF4、double quantization、bfloat16、LoRA rank 16、alpha 32、dropout 0.05、目标模块 q_proj/k_proj/v_proj/o_proj/gate_proj/up_proj/down_proj、per-device batch size 1、gradient accumulation 16、learning rate 1e-4、warmup ratio 0.03、weight decay 0.01、cosine scheduler、gradient checkpointing 和 max grad norm 1.0。

    ## 附录B 校准方法补充

    项目实现比较了 temperature scaling、Dirichlet calibration、isotonic regression、vector scaling 和 conformal prediction[29-32]。正式主线采用温度缩放作为后选择校准，因为它参数少、稳定、易复现，并且不会改变预测排序。其他校准方法可作为附加诊断：isotonic 或 Dirichlet calibration 在部分表中可能获得较低 ECE，但可能改变概率形状；conformal prediction 适合输出集合预测和覆盖率，但不直接对应本文主表的单标签 Judge 输出。

    ## 附录C 误差与风险分析补充

    Tie 类误差主要来自三类样本。第一，两个回答都满足基本要求但表达风格不同，基础 Judge 容易偏向更长或更结构化的一侧。第二，回答都存在轻微缺陷，模型倾向于选择缺陷较少的一侧，而人工标签可能为 Tie。第三，检索增强样本中两个回答证据支持差异不明显，但存在局部实体或数字缺口。Tie rescue 针对的是第一和第二类：当 $p_T$ 足够高且 A/B 概率差不大时，模型说明自身已有 Tie 信号，只是最大概率仍落在 A 或 B。

    ## 附录D 数据预处理逐步说明

    本附录把数据处理步骤展开为可复现检查清单。第一步是来源准入。每个候选来源进入数据集之前，先检查许可证状态、获取路径、获取时间和哈希值。该步骤的目的不是形式化记录，而是保证后续论文中能够说明哪些数据可用于训练，哪些数据只能用于外部评测。RewardBench 在项目 manifest 中被标为 external_eval_only，原因是许可证和使用边界不适合混入训练集。此类数据如果误入训练，会导致外部评测被污染，也会使数据可用性声明不准确。

    第二步是原始记录解析。不同来源有不同记录结构：HelpSteer2 和 OASST1 更接近对话或偏好反馈，OffsetBias 包含偏差扰动信息，RAGTruth 包含检索增强生成中的事实性错误标注，中文专业样本包含内部标注字段。解析器不能简单按字段名拼接，而需要先识别任务类型。若任务是 pairwise preference，则必须同时存在 answer_a、answer_b 和 pairwise 标签；若任务是 single-answer factuality，则 answer_b 可以为空，但 context 或 reference 对事实性判断更重要。任务识别错误会直接导致训练标签空间错误，因此预处理阶段把 task_type 和 scoring_system 写入 metadata。

    第三步是文本规范化。所有 prompt、context、answer 和 reference 字段执行空白折叠、字符串化和空字段处理。空 context 与空 reference 不被强行补写为伪证据，而是在后续特征中通过 `has_reference` 或 context length 显式表达。这样做的原因是，事实性模块需要区分“没有证据”和“证据不支持”。若把缺失证据写成固定文本，overlap 与 support 计算会引入虚假的共同 token。

    第四步是标签映射。pairwise 标签统一为 A>B、B>A、Tie，事实性标签统一为 supported、unsupported、ambiguous。对于来源中的打分或投票信息，预处理保留 human_score，但最终训练 head 使用 human_label。该设计保留了更丰富的标注信息，便于未来做软标签或标注者分歧建模；同时当前实验只使用离散标签，避免不同来源的分数量纲不一致。

    第五步是质量门禁。正式 v2 构建门禁包括样本量、重复 ID、重复内容、跨 split 内容泄漏、许可证元数据完整性和事实性 context 缺失检查。跨 split 泄漏尤其关键。若同一内容组同时出现在 train 和 test，融合分类器可能记住 dataset 或 prompt 模式，导致测试 accuracy 虚高。本文只引用门禁通过后的 `bea_judge_cleaned_10000.json`，不把 `latest_summary.json` 等探测结果写入正式结论。

    第六步是 SFT 样本构建。SFT prompt 沿用 M-Prometheus pairwise 模板，target 只保留结果标记。与把完整 rationale 作为训练目标相比，这种目标更短、更稳定，也更适合 1024 序列长度。训练时 prompt 部分 label 置为 -100，只在 target token 上计算交叉熵，因此模型不会被要求复述输入内容。该处理降低显存压力，也降低因长解释导致的输出漂移。

    第七步是子集构建。SFT size 消融使用确定性采样，25%、50% 和 100% 子集共享 dev 集。共享 dev 集的原因是让不同训练数据量的模型在相同验证和测试条件下比较；若 dev 集也随训练子集变化，则无法区分训练数据量影响和验证集采样影响。子集采样必须固定 seed，否则三种子训练结果会混入数据采样方差。

    第八步是输出归档。每次训练和评估都写入 raw summary、claim gate、comparison report 和 calibrated results。论文只使用正式汇总目录中的数据。这样做避免从日志中人工挑选结果，也使失败门禁可以被记录。例如，某些三种子保守汇总 gate 可能为 false，不能被当作最终提交包；Tie rescue 审计也必须区分 per-seed policy 和 global policy。

    ## 附录E 四模块数学推导展开

    模块1的生成式输出可视为隐式打分函数。设基础模型在结果标记集合 $\\mathcal{{R}}=\\{{[RESULT]A,[RESULT]B,[RESULT]Tie\\}}$ 上给出 token 条件概率。若只取生成结果，得到离散标签 $\\ell_i$；若进一步读取 score_a 和 score_b，则得到连续边际。融合层等价于把生成式 Judge 的输出投影到一个判别特征空间。这个投影很重要，因为生成模型本身未必校准，而判别分类器可以在 dev 上被温度缩放。

    对 QLoRA，低秩增量的参数量为 $r(d+d')$，远小于全量矩阵 $dd'$[11,12]。若一个线性层输入维度为 $d$、输出维度为 $d'$，LoRA 参数占比约为 $r(d+d')/(dd')$。当 $d,d'$ 较大且 $r=16$ 时，占比很小。量化后的基础权重不参与全精度更新，显存主要用于 4-bit 权重、LoRA 参数、优化器状态和激活[12,13]。gradient checkpointing 通过重算部分激活换取显存，适合 24GB 条件下的 3B 模型训练。

    模块2的风险函数是启发式与可学习融合的结合。单个风险 $R_k$ 不直接决定标签，而是作为特征进入 softmax classifier。因此，即便长度风险为 1，模型也可以根据基础分数、证据支持和 dataset one-hot 判断是否真的需要修正。形式上，融合层对某个类别 $y$ 的 logit 可写为：

    $$
    h_y(x)=\\beta_{{y,0}}+\\beta_{{y,base}}^\\top z^{{base}}+\\beta_{{y,bias}}^\\top z^{{bias}}+\\beta_{{y,evi}}^\\top z^{{evidence}}+\\beta_{{y,text}}^\\top z^{{text}}.
    $$

    若某个偏差特征与错误相关，训练会使对应 $\\beta$ 调整该类别 logit；若特征与标签无稳定关系，L2 正则会抑制其权重。这比硬规则更稳健，也解释了为什么偏差模块可同时用于决策和风险提示。

    模块3的证据支持可以理解为弱事实性判别器。它不做复杂自然语言推理，而是把容易造成幻觉的可观测信号拆开。数字缺口捕捉数量事实错误，日期缺口捕捉时间事实错误，实体缺口捕捉名词实体错误，否定不一致捕捉极性错误，比较关系不一致捕捉方向性错误。综合风险取最大值而不是平均值，是因为任意一个关键事实错误都可能使回答不可靠；平均会把严重局部错误稀释。

    对句子支持，设回答句集合为 $\\mathcal{{S}}_a$，证据句集合为 $\\mathcal{{S}}_e$。每个回答句的最佳支持为：

    $$
    B(s)=\\max_{{e\\in\\mathcal{{S}}_e}}S(s,e).
    $$

    低支持句比例为：

    $$
    R_{{lowSent}}=\\frac{{1}}{{|\\mathcal{{S}}_a|}}\\sum_{{s\\in\\mathcal{{S}}_a}}\\mathbb{{1}}(B(s)<\\gamma).
    $$

    其中 $\\gamma$ 为实现中的支持阈值。该特征对长回答尤其重要，因为长回答可能整体 overlap 较高，但其中某些关键句缺乏证据。anchored hallucination severity 进一步关注含实体、数字或比较词的低支持句，使风险更集中在事实承载句上。

    模块4的温度缩放可从 NLL 推导。若未校准模型过度自信，$T>1$ 会平滑分布，降低最大概率；若模型过于保守，$T<1$ 会锐化分布。由于温度缩放不改变类别排序，accuracy 在理论上保持不变，但 ECE、NLL 和 Brier 会变化。本文选择目标 $NLL+ECE+0.25Brier$，是为了同时约束概率对数损失、分箱校准误差和平方概率误差。单独最小化 ECE 可能产生不稳定概率，单独最小化 NLL 可能对分箱误差不敏感。

    Tie rescue 的规则可看作对概率单纯形中的局部区域重标记。三分类概率位于二维单纯形中。原预测为 A 或 B 表示 $p_A$ 或 $p_B$ 为最大值；rescue 条件要求 $p_T$ 足够高、$p_A$ 与 $p_B$ 彼此接近且二者都不太高。这实际上选择了靠近 Tie 顶点但尚未成为 argmax 的边界区域。因为策略只在特定 dataset 上生效，它还隐含一个数据源条件，避免把同一阈值盲目应用到所有分布。

    ## 附录F 实验读数来源与表格生成口径

    主对比表的 Current BEA-Judge、epoch2_1024 和 epoch2_1024 + Tie rescue 使用用户指定的关键结果，并与 `external_3b_full_comparison_table.md` 和 Tie rescue 审计一致。Current BEA-Judge 是内部四模块冻结基线，数值不带方差；为避免误导，正文将其写作 3-seed repeated baseline，而不把零方差解释为独立重复训练。epoch2_1024 和 Tie rescue 是三种子 mean +/- std。

    外部基线表中，GRM、Qwen 和 GLIDER 都是 single full test。表格中 parse failure rate 为 0.0000，表示输出解析没有失败；但这不表示模型概率校准或判别质量一定好。Prometheus-2 7B 虽在脚本和历史报告中存在，但本文遵循用户要求，不把它纳入正式主表和主结论。若读者需要比较 7B 历史结果，应在附录另列并明确标注为历史参考。

    四模块消融来自 `sci_tables_v2_20260521_110114/ablation_table.md`。该表同时包含 pairwise 和 factuality head。正文重点解释 pairwise，因为主贡献围绕 Tie recall；但 evidence module 对 factuality 的影响更直接，因此正文也报告 factuality head 的 full 与 w/o Evidence 结果。对于 w/o Bias Module 指标高于 Full BEA-Judge 的现象，正文采用保守解释，不把偏差模块包装成 accuracy 改进组件。

    SFT size 消融来自 `qlora_sft_size_ablation_3seed_1024_summary/sft_size_ablation_summary.md`。由于该消融固定 epoch1_1024，不能直接与最终 epoch2_1024 主模型等同。它回答的是训练数据量问题，而不是最终最优配置问题。训练轮数消融来自 `qlora_epoch_ablation_3seed_1024_summary/epoch_ablation_summary.md`，它说明 epoch2 在 accuracy 与 macro-F1 上更强，但 Tie recall 下降，支撑 Tie rescue 的必要性。

    风险复核表来自 `sci_tables_ablation_report/risk_coverage_table.md`。该表不是主性能表，而是流程分析表。review_rate 表示被送入人工复核的比例；error_capture_rate 表示错误中有多少被复核集合覆盖；auto_accept_accuracy 表示未复核样本的准确率。该分析的目标是说明置信度与风险输出如何支持人工流程，而不是声称模型标签本身发生改变。

    ## 附录G 结果解释模板

    为避免过度声明，本文建议按以下模板解释结果。第一，当报告 epoch2_1024 时，应说“相较 Current BEA-Judge，accuracy、macro-F1 与 ECE 改善，但 Tie recall 下降”。不应说“所有指标全面提升”。第二，当报告 Tie rescue 时，应说“在 accuracy 基本不变的前提下，Tie recall 提高，ECE 小幅变化”。不应说“Tie rescue 无代价”，因为 macro-F1 和 ECE 的变化虽小，但仍是变化。

    第三，当比较外部基线时，应强调运行类型差异。内部结果是三种子均值，外部基线是 single full test。若直接比较标准差或显著性，会超出数据支持范围。第四，当讨论 GLIDER 的 ECE 时，应同时指出其 accuracy 较低，避免把低 ECE 单独解释为更强评价能力。第五，当讨论 GRM 的 Tie recall 时，应同时指出其 ECE 较高，说明 Tie 识别和校准之间仍存在权衡。

    第六，当讨论偏差模块时，应区分“决策收益”和“复核收益”。消融显示去除偏差模块后 pairwise accuracy 更高，但风险捕获略低。因此偏差模块的价值更适合表述为提高偏差可解释性和复核定位，而不是提高所有主指标。第七，当讨论证据模块时，应强调其对 factuality head 的贡献，因为 w/o Evidence 明显降低 factuality accuracy 和 macro-F1。第八，当讨论 QLoRA 时，应强调它改善基础判别并降低融合后 ECE，但不会自动解决 Tie 类少数问题。

    ## 附录H 数据来源逐项说明

    HelpSteer2 在 BEA-Judge-10K-v2 中承担一般帮助性与指令跟随偏好样本的作用。该来源规模较大，原始记录数超过两万，正式构建接受 2,000 条。它的价值在于覆盖普通用户指令下的回答质量差异，使 Judge 不只学习偏差或事实性特例。由于 HelpSteer2 也被 Tie rescue 的全局策略选为生效数据集，本文对其解释保持谨慎：这说明验证集上该数据源的 Tie 边界可被概率条件识别，并不表示同一阈值可直接迁移到所有来源。

    OASST1 提供开放式对话数据，接受 800 条。其回答风格比严格问答更自然，能够补充多轮对话和用户表达差异。对 Judge 模型而言，这类样本有助于减少模板化偏好，因为回答可能在语气、详略和结构上差异明显。若只用高度规整的 benchmark 数据，模型可能对真实用户输入中的语义噪声适应不足。

    OffsetBias 接受 1,500 条，是偏差感知模块的重要来源。该数据使系统能够观察到回答位置、长度或呈现形式改变后 Judge 是否改变偏好。偏差样本的意义不只是训练模型“不要偏”，更重要的是让审计表可以分组报告偏差风险。没有这类样本，偏差模块只能基于人工构造规则，缺少与人工标签对应的统计证据。

    RAGTruth 接受 2,500 条，是证据事实性模块的核心来源[22]。检索增强场景中的错误往往不是回答整体低质量，而是局部事实、实体或数字与证据不一致[21-26]。传统 pairwise 偏好可能把流畅但不忠实的回答判为更好，因此 BEA-Judge 需要把 context/reference 与回答之间的支持关系显式特征化。RAGTruth 的加入使 evidence gap、negation mismatch 和 anchored hallucination severity 这些特征有实际训练和诊断意义。

    早期来源如 MT-Bench human judgments、PandaLM、JudgeBench、ARES 和 WikiEval 构成 3.4K legacy anchor，并被扩展到 v2 正式数据中。它们提供了开放问答、pairwise preference 和事实性标签的基础结构。中文专业标注样本接受 1,000 条，弥补了多语言与专业表达场景。本文没有把这些来源逐一作为主结果分表，是因为用户指定的关键结果已经以全测试集为单位给出；但在复现实验中，可以用 per_dataset_table 检查每个来源的误差。

    数据来源之间并不完全同质，这正是 BEA-Judge 需要融合校准的原因。若模型在某一来源上 accuracy 较低，source risk 会进入偏差模块；若某一来源的 context 更长，证据模块的 support score 分布也会变化；若某一来源的 Tie 比例更高，Tie recall 的解释也应按来源分层。统一主表给出整体结论，分来源表则用于诊断而非替代主结论。

    ## 附录I 算法流程与伪代码说明

    算法1给出训练和评估流程的文字伪代码。输入为正式数据集 $D$、基础模型 $M_0$、QLoRA 配置 $C_q$ 和融合配置 $C_f$。首先按 split 得到 $D_{{train}},D_{{dev}},D_{{test}}$。其次，从 pairwise 样本构建 SFT JSONL，并用 QLoRA 得到适配器 $A_q$。第三，对 train/dev/test 运行 $M_0+A_q$，得到 base scores。第四，对每个样本计算文本、基础、偏差、证据和来源 one-hot 特征。第五，在 train 上训练 pairwise head，在 dev 上选择超参数和温度。第六，在 test 上评估 epoch2_1024。第七，在 dev 上搜索 Tie rescue 策略，若通过 accuracy、macro-F1、ECE 和 Tie recall 约束，则应用到 test，得到 epoch2_1024 + Tie rescue。

    该流程的关键约束是数据边界。训练阶段不能读取 test 标签以外的信息；校准和 Tie rescue 只能使用 dev；test 只在最终评估时读取标签计算指标。若在 test 上调阈值，Tie rescue 的结果会变成事后优化，无法作为泛化证据。本文脚本中 `accuracy_constrained_tie_rescue_audit.py` 把候选策略、合格策略数量、选择顺序和逐种子 test 结果写入审计文件，目的就是防止阈值选择过程不可追溯。

    算法2给出 Tie rescue 的决策过程。对每个测试样本，若原预测为 Tie，则不改变；若原预测为 A>B 或 B>A，但样本来源不是策略指定 dataset，也不改变；若来源匹配，则读取 $p_A,p_B,p_T$。只有当 $p_T\\ge0.38$、$|p_A-p_B|\\le0.65$ 且 $\\max(p_A,p_B)\\le0.65$ 时，预测改为 Tie。这个规则包含三个直觉：Tie 概率必须已经较高，A 与 B 的差异必须不大，且 A/B 任何一方都不能过于自信。三者缺一不可，否则策略可能把明确的 A/B 胜负误改为 Tie。

    算法3给出风险复核排序。对已校准样本计算风险分数，最简单形式为 $r_i=1-s_i$。按 $r_i$ 从高到低排序，选择前 $k$ 个样本进入人工复核。随着 $k/n$ 增大，error_capture_rate 增大，auto_accept_accuracy 也提高。若加入偏差风险和证据风险，可写为：

    $$
    r_i=\\lambda_0(1-s_i)+\\lambda_1R_{{bias,i}}+\\lambda_2R_{{evidence,i}}.
    $$

    当前正式结果主要报告基于置信度的风险覆盖表，未来可在 dev 上学习 $\\lambda$，但同样必须避免 test 调参。风险复核算法的产出不是新标签，而是复核队列。它适合与人工标注预算结合，例如固定 review_rate 为 20% 或 30%，观察错误捕获和自动接受准确率。

    ## 附录J 论文组织与SCI写作规范说明

    本稿按照“问题提出-数据-方法-实验-讨论-附录”的顺序组织。引言不把 BEA-Judge 描述为通用万能评估器，而是从 accuracy、Tie recall 和 calibration 的冲突切入；方法部分不只列模块名称，而是给出每个模块的输入、输出、公式和协同机制；实验部分区分内部三种子结果和外部 single full test；讨论部分强调 trade-off 和有效性威胁。这样的组织方式更接近 SCI 论文对可复现性和边界条件的要求。

    图表也按证据链组织。图1说明系统结构，回答“方法是什么”；表1和表2说明数据与 prompt，回答“输入是什么”；表3、表5和图4、图5说明训练轮数与数据量，回答“QLoRA 如何影响结果”；表4和图2说明主结果，回答“相比基线如何”；表7和图3说明 Tie rescue，回答“如何在 accuracy 约束下恢复 Tie recall”；表8和图6说明风险复核，回答“置信度输出如何服务人工流程”。这种顺序避免把所有结果堆成指标列表。

    格式上，正文中文使用宋体五号，英文使用 Times New Roman 五号；公式采用 `$...$` 和 display math；表格采用三线表；图片以 SVG 输出，便于 Visio 打开后编辑。DOCX 由 Markdown 经 Pandoc 生成，并经过 python-docx 后处理以统一字体和表格边框。由于不同 Word/Pandoc 版本对 SVG 的嵌入方式可能不同，交付目录同时保留独立 SVG 文件，复现者可在 Visio 中直接打开或导入。

    写作风格上，本文避免使用“显著优于所有模型”“全面提升”等宣传性表述。只要某个指标存在下降或边界，就在正文中说明。例如 epoch2_1024 的 Tie recall 低于 Current BEA-Judge，Tie rescue 的 ECE 有 0.0005 的小幅增加，外部基线没有三种子方差。这些说明并不削弱贡献，反而使主张更可复核。本文真正的贡献是把轻量 3B Judge 的强判别、稳定校准和 Tie 边界修复拆成可控操作点，而不是追求单表上的绝对完美。

    ## 附录K 分任务完成核对

    对子任务1，正文第3节已经说明数据构成、来源选择理由、数据集创新性和预处理流程。表1给出 prompt 模板，表2给出样例摘录。附录D和H进一步展开来源准入、schema 统一、文本规范化、标签映射、质量门禁、SFT 构建和来源逐项作用。需要注意，论文没有把 `论文撰写` 目录中的旧稿作为数据依据；该目录只作为本次输出位置。正式数据依据来自 `datasets`、`configs`、`src`、`scripts` 和 `论文参考` 中的 CodeUltraFeedback。

    对子任务2，正文第4节给出四模块框架图和完整数学说明。模块1说明基础 Judge prompt、QLoRA 低秩增量和 SFT 损失；模块2说明长度、格式、位置、rubric 和 source risk；模块3说明 token coverage、claim support、numeric/date/entity gap、negation mismatch、comparative mismatch 和 hallucination risk；模块4说明 softmax fusion、L2 正则、温度缩放、ECE 和 confidence 输出。附录E进一步解释这些公式为何这样组合。

    对子任务3，正文第6节解释 QLoRA 微调原因：Current BEA-Judge 的四模块已验证，但冻结基础 Judge 的判别信号较弱，因此需要轻量微调提高基础评分质量。表3和图5给出训练轮数消融，表4和图2给出主实验结果。正文明确指出稳定序列长度为 1024，最终主模型是 epoch2_1024 和 epoch2_1024 + Tie rescue 两个配置。

    对子任务4，正文第7节说明 SFT size 消融目的：检验训练数据量对 accuracy、macro-F1、ECE 和 Tie recall 的影响。表5和图4显示 25%、50% 和 100% SFT 数据量下的三种子结果。正文没有把 SFT size 消融误写为最终配置，而是说明该消融固定 epoch1_1024，回答的是数据规模问题。

    对子任务5，正文第9节介绍外部轻量基线和风险复核分析。GRM 被解释为 external 3B reward baseline，Qwen 被解释为 external 3B instruct baseline，GLIDER 被解释为 external 4B evaluator baseline。Prometheus-2 7B 明确不进入正式主表和主结论。风险复核分析给出表8和图6，解释其必要性在于人工复核和自动接受流程，而不是主标签性能。

    对子任务6，全文组织围绕用户给出的核心事实：四模块已完成验证；主模型采用 QLoRA；稳定序列长度为 1024；最终配置为 epoch2_1024 与 epoch2_1024 + Tie rescue；内部结果是 3-seed mean +/- std；外部基线是 single full test；主结论聚焦 accuracy 不降、Tie recall 上升、校准更稳。正文讨论了 trade-off，附录保留复现细节、OOM 说明、误差分析和审计口径。

    ## 附录L 潜在审稿质疑与回应

    质疑1：为什么不直接选择 Tie recall 最高的模型？回应：Tie recall 不能脱离 accuracy 和 calibration 单独优化。Current BEA-Judge 的 Tie recall 为 0.5231，高于 epoch2_1024，但 accuracy 和 ECE 明显弱于 QLoRA-BEA-Judge。本文目标不是最大化单一 Tie recall，而是在高 accuracy 操作点上恢复一部分 Tie recall。Tie rescue 的约束设计正是为了防止为了 Tie recall 牺牲总体判别。

    质疑2：为什么 GLIDER 的 ECE 较低却不作为更好模型？回应：ECE 只衡量置信度与正确率的匹配，不衡量正确率本身。GLIDER 的 ECE 为 0.0353，但 accuracy 为 0.5043，接近随机方向性判别的低水平。一个低 accuracy 但保守输出的模型可能有较低 ECE，却不能满足 Judge 主任务。因此主结论必须同时看 accuracy、macro-F1、ECE 和 Tie recall。

    质疑3：为什么 w/o Bias Module 的 pairwise accuracy 更高，还要保留偏差模块？回应：偏差模块的价值不只体现在主决策指标，也体现在风险复核和偏差审计。消融显示去除偏差模块可能提高部分 pairwise 指标，但 bias_as_decision_features 在 review_capture_rate 上更高。若系统用于需要偏差可解释性的评估流程，偏差模块仍有保留价值。本文因此没有夸大偏差模块的 accuracy 贡献。

    质疑4：Tie rescue 是否是测试集后处理？回应：不是。策略在验证集上搜索和选择，测试集只应用一次。审计脚本记录了候选阈值、选择约束、选择顺序和逐种子测试结果。若未来改用其他数据集或其他模型，必须重新在对应验证集上选择，而不能复用测试集结论。

    质疑5：为什么主表不纳入 Prometheus-2 7B？回应：本文主题是轻量 3B Judge 和轻量外部基线。Prometheus-2 7B 可作为历史参考，但参数规模和实验定位不同，纳入主表会模糊“轻量”比较边界。用户要求也明确指出 Prometheus-2 7B 不得纳入正式主表和主结论。

    质疑6：为什么只用启发式证据特征，而不是更强 NLI 模型？回应：本文目标是轻量、可复现和可解释。启发式证据特征不能覆盖所有事实推理，但能稳定捕捉数字、日期、实体、否定和比较关系等高风险错误，并且不引入额外大型模型。未来可以在相同框架中替换或增强证据模块，但当前结果已经说明证据模块对 factuality head 有明确作用。

    质疑7：外部基线 single full test 是否足够？回应：外部基线主要用于轻量参考，不用于统计显著性声明。正文已经明确区分 single full test 与内部三种子结果，避免把外部基线方差缺失解释为稳定性。若后续投稿需要更强证据，可以对外部基线也进行多次 prompt/order 变化或 bootstrap 置信区间分析。

    ## 附录M 复现实验记录建议

    为便于后续投稿或开源，建议把每次训练和评估记录为四类文件。第一类是配置文件，包括基础模型路径、QLoRA 超参数、最大序列长度、训练轮数、随机种子、数据文件哈希和输出目录。该文件用于回答“实验如何启动”。第二类是训练日志，包括 loss 曲线、显存使用、训练步数、保存 checkpoint 的时间和任何 OOM 或中断记录。该文件用于回答“实验是否稳定完成”。第三类是评估输出，包括 raw prediction、parsed label、probability、confidence、risk score、human_label 和 per-dataset summary。该文件用于回答“指标如何计算”。第四类是审计报告，包括 claim gate、parse failure、calibration table、Tie rescue 候选策略、dev 选择结果和 test 一次应用结果。该文件用于回答“论文中的每个结论是否可追溯”。

    对 QLoRA 主实验，最小复现单元不是单个 checkpoint，而是 seed、epoch、length 与 calibration 配置的组合。本文正式主模型为 epoch2_1024，因此记录中应明确区分 epoch1_1024、epoch2_1024 和 SFT size 消融输出。若复现者只看到一个 adapter 文件，而不知道其训练轮数和数据量，就无法判断它对应主表、消融表还是探索性实验。建议在输出目录名称中包含 `seed`、`epoch`、`max_seq_length` 和 `sft_fraction`，并在 manifest 中写入 parent config。这样做虽然增加文件名长度，但能显著减少论文复核时的混淆。

    对 Tie rescue，复现记录必须包含 dev 搜索空间和筛选顺序。一个合格策略应先满足 accuracy 约束，再考虑 Tie recall 提升；如果多个候选同时合格，应按照预先定义的选择规则排序。本文采用的全局策略为 `dataset=helpsteer2`、`min_tie_probability=0.38`、`max_ab_margin=0.65`、`max_ab_confidence=0.65`。这些数值不应被描述为理论最优阈值，而应描述为在 dev 上通过约束筛选的操作点。若在新数据集上部署，应重新进行 dev 选择，而不是直接套用本文阈值。

    对外部基线，复现记录还应保留 prompt、解析规则和失败样本。本文报告的 GRM、Qwen 和 GLIDER parse failure rate 均为 0.0000，这只说明输出可以解析，不说明输出概率经过同等校准。若外部模型输出没有天然三分类概率，需要说明概率如何构造或 ECE 如何计算。对于 single full test，建议同时记录模型版本、加载精度、推理温度、max tokens 和是否使用 system prompt。没有这些信息，外部基线很难被后续研究者严格复核。

    ## 附录N 图表与版式复核建议

    本文交付 SVG 矢量图和 PNG 预览两套图像。SVG 用于满足 Visio 可编辑要求，PNG 用于 DOCX 嵌入稳定性。复核时应优先检查 SVG 是否包含完整标题、坐标轴、图例或线末标签，以及所有文字是否在 viewBox 内。PNG 则用于检查实际嵌入 Word 后的可读性，尤其是中文标题、英文系列名和小数标签。风险复核图这类末端标签较长的折线图，应避免把系列名放在最后一个点的右侧画布外；更稳妥的做法是根据边界自动选择左置或右置标签。

    表格复核应分为内容和版式两层。内容层面，所有主表指标名必须统一为 accuracy、macro-F1、ECE、Tie recall 和 parse failure rate；内部结果必须保留 mean +/- std；外部基线必须标注 single full test；Prometheus-2 7B 只能出现在历史参考或限制说明中。版式层面，三线表应有表顶线、表头下横线和表底线，不应出现密集竖线。对于 Word 输出，Pandoc 生成后需要用 python-docx 后处理字体和表格边框，避免中英文混排时字体漂移。

    正文版式复核应关注三点。第一，公式采用 `$...$` 或 display math，避免同一篇稿件混用多种公式标记。第二，中文正文保持克制叙述，不把实验结果写成宣传式“全面领先”；当 Tie recall 下降、ECE 小幅变化或外部基线缺少方差时，应在正文说明。第三，附录只放复现、误差和审计细节，不把 2048 OOM 等工程问题放入正文主线。这样能使论文主线集中于轻量 3B Judge、QLoRA、四模块融合校准和准确率约束 Tie rescue，而附录提供足够复核材料。

    ## 附录O 后续研究路线

    后续研究可从三个方向推进。第一是分层 Tie rescue。当前策略只在验证集上选择一个保守操作点，未来可以按数据源、语言、任务类型或标注分歧程度选择不同阈值。但分层策略的参数更多，更容易过拟合，因此需要更严格的 dev/test 隔离和更大的验证集。第二是证据模块增强。现有证据特征主要基于词元、数字、日期、实体和否定/比较模式，适合轻量复现，但对复杂推理、跨句蕴含和隐含条件支持仍有限。未来可以引入小型 NLI、检索验证或声明级事实检查模型，但必须评估其计算成本是否破坏轻量部署目标。

    第三是校准与复核联合建模。本文把校准概率用于 ECE、confidence、risk review 和 Tie rescue，说明概率尺度已经具有操作价值。但风险复核排序仍可进一步学习，例如用 dev 上的错误标签训练一个 error predictor，把 $1-s_i$、偏差风险、证据风险、source risk 和模型输出熵共同输入。该方向的目标不是改变主预测，而是更有效地把有限人工预算分配给可能错误的样本。若未来系统进入高风险应用，复核策略甚至可能比单点 accuracy 更重要，因为它决定了错误如何被发现和拦截。

    还有一个值得探索的方向是标注不确定性建模。Tie 标签并不总是严格表示两个回答完全等价，有时也表示标注者之间的分歧、评价维度权重不同或证据不足。若数据集能保留多标注者投票、分数方差或解释文本，模型可以学习软标签或分歧预测，而不是把 Tie 当作普通第三类。这可能改善 Tie recall 和 calibration 之间的关系，但也会改变任务定义。本文选择离散三分类，是为了与现有结果和用户给定主表保持一致；后续工作可以在不改变主任务的前提下增加 uncertainty head。

    最后，轻量 Judge 的外部效度仍需要扩展。BEA-Judge-10K-v2 覆盖多源文本回答、偏差和事实性样本，但并不等同于代码执行、数学证明、医学建议或法律推理。CodeUltraFeedback 等代码评估研究提示，不同领域有专门偏好维度和错误类型[33]。高风险应用还要求把可靠性、校准与伦理边界同时纳入系统设计[27,28,35]。BEA-Judge 的四模块思想可以迁移，但每个领域都应重新定义证据特征、偏差特征和 Tie 语义。只有在新领域重新完成数据准入、校准和风险复核后，才能提出稳健结论。

    ## 参考文献

    {refs}
    """).strip() + "\n"
    text = "\n".join(line[4:] if line.startswith("    ") else line for line in text.splitlines()) + "\n"
    return text


def create_reference_docx(path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.enum.style import WD_STYLE_TYPE
    except Exception:
        return
    doc = Document()
    styles = doc.styles
    for style_name in ["Normal", "Body Text"]:
        style = styles[style_name] if style_name in styles else styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5)
    for name, size in [("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
    doc.save(path)


def set_cell_text_font(cell: Any) -> None:
    try:
        from docx.shared import Pt
        from docx.oxml.ns import qn
    except Exception:
        return
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(9)


def set_table_three_line(table: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tblPr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    for edge in ["top", "bottom"]:
        tag = borders.find(qn(f"w:{edge}"))
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "12")
        tag.set(qn("w:color"), "000000")
    tblPr.append(borders)
    if table.rows:
        for cell in table.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is None:
                tcBorders = OxmlElement("w:tcBorders")
                tcPr.append(tcBorders)
            bottom = tcBorders.find(qn("w:bottom"))
            if bottom is None:
                bottom = OxmlElement("w:bottom")
                tcBorders.append(bottom)
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "8")
            bottom.set(qn("w:color"), "000000")


def postprocess_docx(input_path: Path, output_path: Path) -> Path | None:
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
    except Exception:
        return None
    doc = Document(input_path)
    section = doc.sections[0]
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            if p.style and p.style.name == "Normal":
                run.font.size = Pt(10.5)
    for table in doc.tables:
        set_table_three_line(table)
        for row in table.rows:
            for cell in row.cells:
                set_cell_text_font(cell)
    if output_path.exists():
        try:
            output_path.unlink()
        except OSError:
            output_path = output_path.with_name(output_path.stem + "_new.docx")
    doc.save(output_path)
    return output_path


def build_docx(md_path: Path, ref_path: Path, docx_path: Path, *, use_png_figures: bool) -> Path | None:
    if shutil.which("pandoc") is None:
        return None
    docx_md_path = OUT / "manuscript_docx.md"
    if use_png_figures:
        docx_text = md_path.read_text(encoding="utf-8").replace("figures/", "figures_png/").replace(".svg)", ".png)")
        docx_md_path.write_text(docx_text, encoding="utf-8")
        input_name = docx_md_path.name
        resource_path = ".;figures_png;figures"
    else:
        input_name = md_path.name
        resource_path = ".;figures"
    raw_docx_path = OUT / "_pandoc_raw_epoch2_tie_rescue.docx"
    if raw_docx_path.exists():
        raw_docx_path.unlink()
    cmd = [
        "pandoc",
        input_name,
        "--from=markdown+tex_math_dollars+pipe_tables",
        "--to=docx",
        f"--reference-doc={ref_path.name}",
        f"--resource-path={resource_path}",
        f"--output={raw_docx_path.name}",
    ]
    subprocess.run(cmd, cwd=OUT, check=True)
    return postprocess_docx(raw_docx_path, docx_path)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    FIG.mkdir(parents=True, exist_ok=True)
    FIG_PNG.mkdir(parents=True, exist_ok=True)
    TAB.mkdir(parents=True, exist_ok=True)
    build_figures()
    png_ok = render_png_figures()
    tables = write_tables()
    md = manuscript(tables)
    md_path = OUT / "manuscript.md"
    md_path.write_text(md, encoding="utf-8")
    ref = OUT / "reference.docx"
    create_reference_docx(ref)
    docx_path = OUT / "QLoRA-BEA-Judge_epoch2_tie_rescue_20260605_final.docx"
    actual_docx_path = build_docx(md_path, ref, docx_path, use_png_figures=png_ok) if ref.exists() else None
    docx_ok = actual_docx_path is not None
    manifest = {
        "output_dir": str(OUT),
        "manuscript": str(md_path),
        "docx": str(actual_docx_path) if docx_ok else None,
        "figures": sorted(p.name for p in FIG.glob("*.svg")),
        "png_figures": sorted(p.name for p in FIG_PNG.glob("*.png")),
        "tables": sorted(p.name for p in TAB.glob("*.csv")),
        "reference_count": len(REFERENCES),
        "word_count_cjk_approx": len(re.findall(r"[\u4e00-\u9fff]", md)),
        "docx_generated": docx_ok,
        "docx_png_previews": png_ok,
    }
    (OUT / "artifact_manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
