from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SCI_DIR = ROOT / "datasets" / "model_outputs" / "sci_tables_v2_20260521_110114"
QLORA_SUMMARY = (
    ROOT
    / "datasets"
    / "model_outputs"
    / "qlora_ablation_3seed_epoch1_1024_summary"
    / "ablation_3seed_summary.json"
)
OUT_DIR = ROOT / "paper"
OUT_PATH = OUT_DIR / "bea_judge_four_module_ablation_report.docx"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def fmt(value: str | float | int | None, digits: int = 4) -> str:
    if value in (None, ""):
        return "--"
    if isinstance(value, (int, float)):
        return f"{value:.{digits}f}"
    try:
        return f"{float(value):.{digits}f}"
    except ValueError:
        return value


def fmt_delta(value: float | str | None, digits: int = 4) -> str:
    if value in (None, ""):
        return "--"
    v = float(value)
    sign = "+" if v >= 0 else ""
    return f"{sign}{v:.{digits}f}"


def metric(summary: dict[str, Any], section: str, variant: str, head: str, name: str) -> str:
    node = summary["sections"][section][variant][head]["metrics"][name]
    mean = node["mean"]
    std = node["std"]
    if mean is None:
        return "--"
    if std is None:
        return fmt(mean)
    return f"{mean:.4f}±{std:.4f}"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:" + key), str(value))


def set_run_east_asia_font(run, font_name: str = "Microsoft YaHei") -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def style_document(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.color.rgb = RGBColor.from_string("1F2937")
    styles["Title"].font.size = Pt(20)
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.size = Pt(12)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.2
    p.add_run(text)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.15
        p.add_run(item)


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(title)
    run.bold = True
    run.font.size = Pt(9.5)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_shading(hdr[idx], "E5E7EB")
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in hdr[idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8.5)
                set_run_east_asia_font(run)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(8.2)
                    set_run_east_asia_font(run)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": "D1D5DB"},
                bottom={"val": "single", "sz": "4", "color": "D1D5DB"},
                left={"val": "single", "sz": "4", "color": "D1D5DB"},
                right={"val": "single", "sz": "4", "color": "D1D5DB"},
            )
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F4F6")
    p = cell.paragraphs[0]
    r = p.add_run(title + "：")
    r.bold = True
    r.font.size = Pt(9.5)
    set_run_east_asia_font(r)
    r = p.add_run(body)
    r.font.size = Pt(9.5)
    set_run_east_asia_font(r)
    doc.add_paragraph()


def build_formal_ablation_tables(doc: Document, ablation: list[dict[str, str]], significance: list[dict[str, str]]) -> None:
    rows = []
    for row in ablation:
        if row["split"] != "test":
            continue
        rows.append(
            [
                row["variant"],
                row["head"],
                fmt(row["accuracy"]),
                fmt(row["macro_f1"]),
                fmt(row["ece"]),
                fmt(row["brier"]),
                fmt(row["tie_recall"]),
            ]
        )
    add_table(
        doc,
        "表1 冻结四模块 BEA-Judge 模块级消融结果（test split）",
        ["变体", "任务头", "Accuracy", "Macro-F1", "ECE", "Brier", "Tie recall"],
        rows,
        [1.65, 0.8, 0.8, 0.85, 0.7, 0.7, 0.85],
    )

    sig_rows = []
    for row in significance:
        if row["head"] not in {"pairwise", "factuality"}:
            continue
        sig_rows.append(
            [
                row["variant"],
                row["head"],
                row["paired_n"],
                fmt_delta(row["delta_accuracy_full_minus_variant"]),
                fmt_delta(row["delta_macro_f1_full_minus_variant"]),
                row["mcnemar_full_only_correct"],
                row["mcnemar_variant_only_correct"],
                fmt(row["mcnemar_p"], 6),
            ]
        )
    add_table(
        doc,
        "表2 Full BEA-Judge 相对消融变体的配对差异与 McNemar 检验",
        ["对照变体", "任务头", "配对 n", "ΔAccuracy", "ΔMacro-F1", "Full only", "Variant only", "p 值"],
        sig_rows,
        [1.65, 0.75, 0.65, 0.8, 0.85, 0.65, 0.8, 0.75],
    )


def build_evidence_table(doc: Document, rows: list[dict[str, str]]) -> None:
    table_rows = [
        [
            row["feature_group"],
            row["weighted_calibration"],
            row["feature_count"],
            fmt(row["accuracy"]),
            fmt(row["macro_f1"]),
            fmt(row["ece"]),
            fmt(row["brier"]),
        ]
        for row in rows
    ]
    add_table(
        doc,
        "表3 事实性证据特征组细粒度消融",
        ["特征组", "加权校准", "特征数", "Accuracy", "Macro-F1", "ECE", "Brier"],
        table_rows,
        [1.75, 0.8, 0.65, 0.8, 0.85, 0.7, 0.7],
    )


def build_qlora_table(doc: Document, summary: dict[str, Any]) -> None:
    variants = [
        ("Full BEA-Judge", "pairwise"),
        ("w/o Bias Module", "pairwise"),
        ("w/o Evidence Module", "pairwise"),
        ("w/o Calibration", "pairwise"),
        ("w/o Base Judge Scores", "pairwise"),
        ("w/o Tie Policy", "pairwise"),
        ("Full BEA-Judge", "factuality"),
        ("w/o Evidence Module", "factuality"),
    ]
    rows = []
    for variant, head in variants:
        section = "variants"
        rows.append(
            [
                variant,
                head,
                metric(summary, section, variant, head, "accuracy"),
                metric(summary, section, variant, head, "macro_f1"),
                metric(summary, section, variant, head, "ece"),
                metric(summary, section, variant, head, "brier"),
                metric(summary, section, variant, head, "tie_recall"),
            ]
        )
    add_table(
        doc,
        "表4 QLoRA-BEA-Judge epoch1 三种子模块回放消融",
        ["变体", "任务头", "Accuracy", "Macro-F1", "ECE", "Brier", "Tie recall"],
        rows,
        [1.65, 0.8, 1.0, 1.0, 1.0, 1.0, 1.0],
    )


def build_report() -> Path:
    ablation = read_csv(SCI_DIR / "ablation_table.csv")
    significance = read_csv(SCI_DIR / "ablation_significance_table.csv")
    evidence_groups = read_csv(SCI_DIR / "evidence_feature_group_ablation_table.csv")
    main_results = read_csv(SCI_DIR / "main_results_table.csv")
    qlora_summary = json.loads(QLORA_SUMMARY.read_text(encoding="utf-8"))

    doc = Document()
    style_document(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("BEA-Judge 四模块消融实验报告")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("基于本项目正式 SCI 表格与 QLoRA 三种子模块回放结果")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("4B5563")
    set_run_east_asia_font(r)
    doc.add_paragraph()

    doc.add_heading("1. 实验目的", level=1)
    add_paragraph(
        doc,
        "本报告围绕 BEA-Judge 的四模块框架开展消融分析，目标是回答三个问题：第一，基础 Judge 分数、偏差感知、证据增强和融合校准分别承担什么功能；第二，去除特定模块后 pairwise 偏好判断和 factuality 判断的指标如何变化；第三，哪些模块应作为论文中的主要性能贡献，哪些模块应被定位为风险控制或审计机制。"
    )
    add_callout(
        doc,
        "报告边界",
        "主结论以 datasets/model_outputs/sci_tables_v2_20260521_110114 下的正式消融表为准；QLoRA 三种子模块回放用于补充验证模块作用，不与冻结四模块基线混写为同一实验。"
    )

    doc.add_heading("2. 四模块框架与实验设置", level=1)
    add_bullets(
        doc,
        [
            "基础 Judge 评分模块：使用 M-Prometheus-3B 对 pairwise 样本输出 A/B/Tie 等原始判断和分数信号，是偏好判断的核心语义来源。",
            "偏差感知模块：识别位置、长度、格式、rubric sensitivity 和数据源等风险，用于风险审计、复核优先级和保守决策控制。",
            "证据增强事实性模块：基于 context/reference 与回答之间的支持关系，提取数值、日期、实体、否定、比较关系和局部幻觉风险等证据特征。",
            "融合校准与置信度输出模块：将基础分数、偏差特征、证据特征和文本特征输入任务特定 softmax head，并在 dev 集上选择温度、阈值和 Tie policy，输出最终标签、置信度、ECE 与复核信号。",
        ],
    )
    add_paragraph(
        doc,
        "正式实验包含 pairwise 与 factuality 两个任务头。pairwise 任务报告 Accuracy、Macro-F1、ECE、Brier 和 Tie recall；factuality 任务报告 Accuracy、Macro-F1、ECE 与 Brier。显著性分析采用 Full BEA-Judge 与消融变体之间的配对差异，并报告 McNemar 检验。"
    )
    main_rows = [
        [
            row["head"],
            row["split"],
            row["n"],
            fmt(row["accuracy"]),
            fmt(row["macro_f1"]),
            fmt(row["ece"]),
            fmt(row["brier"]),
            fmt(row["tie_recall"]),
            fmt(row["review_rate"]),
        ]
        for row in main_results
        if row["split"] == "test"
    ]
    add_table(
        doc,
        "表0 Full BEA-Judge 正式测试集主结果",
        ["任务头", "Split", "n", "Accuracy", "Macro-F1", "ECE", "Brier", "Tie recall", "Review rate"],
        main_rows,
        [0.75, 0.65, 0.55, 0.8, 0.8, 0.65, 0.65, 0.8, 0.8],
    )

    doc.add_heading("3. 冻结四模块消融主结果", level=1)
    build_formal_ablation_tables(doc, ablation, significance)
    add_paragraph(
        doc,
        "在 pairwise 任务中，Full BEA-Judge 的 accuracy 为 0.7512，macro-F1 为 0.6730，ECE 为 0.0558，Tie recall 为 0.5231。去除 Calibration 后 accuracy 降至 0.7407，macro-F1 降至 0.6402，Full 相对该变体的 ΔAccuracy 为 +0.0105、ΔMacro-F1 为 +0.0328，McNemar p=0.034690，说明融合校准与决策策略对 pairwise 分类边界和类别均衡表现具有统计意义上的贡献。"
    )
    add_paragraph(
        doc,
        "去除 Evidence Module 对 pairwise 的影响较小：accuracy 为 0.7531，macro-F1 为 0.6711，Full 相对该变体的 ΔMacro-F1 仅为 +0.0019，McNemar p=0.726562。这说明在成对偏好判断中，基础 Judge 分数已经承担主要语义判别功能，证据特征更多是辅助信号，而不是 pairwise 主导因素。"
    )
    add_paragraph(
        doc,
        "去除 Bias Module 后，pairwise accuracy 和 macro-F1 分别变为 0.7654 和 0.6892，高于 Full BEA-Judge；Full 相对该变体的 ΔAccuracy 为 -0.0142，p=0.016674。该结果不能被解读为偏差模块“无效”，而应说明其目标不是单纯提高准确率。偏差模块引入了风险控制和复核优先级信号，可能牺牲部分测试集点估计性能来换取更保守、更可审计的评估流程。"
    )
    add_paragraph(
        doc,
        "在 factuality 任务中，证据增强是最关键模块。Full BEA-Judge 的 accuracy 和 macro-F1 分别为 0.7649 和 0.7405；去除 Evidence Module 后分别降至 0.6928 和 0.6542。Full 相对该变体的 ΔAccuracy 为 +0.0721、ΔMacro-F1 为 +0.0863，McNemar p=0.000224，说明事实性判断高度依赖显式证据支持特征。"
    )

    doc.add_heading("4. 证据特征组消融", level=1)
    build_evidence_table(doc, evidence_groups)
    add_paragraph(
        doc,
        "细粒度证据消融进一步说明，单纯 overlap-only 特征可达到 0.7175 accuracy 和 0.6780 macro-F1；加入 numeric/date/entity 特征后 accuracy 提升到 0.7361，macro-F1 提升到 0.6900；加入 sentence/local-risk 后 accuracy 达到 0.7629，但 macro-F1 为 0.6789。最终加入 weighted calibration 后，accuracy 为 0.7649，macro-F1 显著提升到 0.7405，表明事实性任务不仅需要更多证据特征，还需要校准机制将不同证据信号转化为稳定的类别概率。"
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("5. QLoRA 三种子模块回放补充结果", level=1)
    build_qlora_table(doc, qlora_summary)
    add_paragraph(
        doc,
        "QLoRA 三种子模块回放显示，Full BEA-Judge pairwise accuracy 为 0.8025±0.0034，macro-F1 为 0.7128±0.0063，ECE 为 0.0279±0.0046。去除 Base Judge Scores 后 accuracy 下降到 0.5626±0.0022，macro-F1 下降到 0.5560±0.0075，ECE 升至 0.1783±0.0332，说明基础 Judge 分数是 pairwise 任务最不可替代的核心信号。"
    )
    add_paragraph(
        doc,
        "QLoRA 回放中，去除 Evidence Module 对 pairwise 指标几乎没有影响，但对 factuality 造成断崖式下降：Full factuality macro-F1 为 0.7405±0.0000，w/o Evidence Module 仅为 0.2629±0.0000，ECE 升至 0.6274±0.0000。这与冻结四模块实验一致，进一步强化了“证据增强主要支撑事实性可靠性”的结论。"
    )
    add_paragraph(
        doc,
        "w/o Tie Policy 的 pairwise accuracy 提升到 0.8288±0.0014，ECE 降至 0.0197±0.0058，但 Tie recall 降至 0.3359±0.0193。这表明 Tie 策略对边界样本有实质影响：如果只追求总体 accuracy，模型会倾向减少 Tie 判定；如果关注评价公平性和模糊样本识别，则需要保留或进一步优化 Tie policy。"
    )

    doc.add_heading("6. 综合讨论", level=1)
    add_bullets(
        doc,
        [
            "基础 Judge 分数是 pairwise 偏好判断的主干信号；缺少基础分数后，QLoRA 回放中的 accuracy 和 macro-F1 均明显下降。",
            "证据增强模块是 factuality 判断的主要贡献来源；无论冻结线还是 QLoRA 回放，去除证据后 factuality macro-F1 均显著降低。",
            "融合校准模块对 pairwise 的类别均衡和决策稳定性有贡献，但它不是单纯最小化 ECE 的模块；在正式表中，w/o Calibration 的 ECE 更低而 accuracy/macro-F1 更差，说明校准目标存在性能与置信度形态之间的权衡。",
            "偏差模块应被严格定位为 risk-control 和 review prioritization 组件。正式消融中去除偏差后点估计指标更高，因此论文中不应宣称偏差模块直接提升总体 accuracy。",
            "Tie policy 会影响总体 accuracy 与 Tie recall 的取舍。对于自动评估系统，Tie recall 不是附属指标，而是反映边界样本识别能力和评估谨慎性的关键指标。",
        ]
    )

    doc.add_heading("7. 结论", level=1)
    add_paragraph(
        doc,
        "本项目的四模块消融实验表明，BEA-Judge 的不同模块承担互补角色：基础 Judge 分数决定 pairwise 判别的主体性能，证据增强模块决定 factuality 可靠性，融合校准模块改善 pairwise 决策质量和类别均衡，偏差感知模块主要服务于风险审计和复核优先级。报告写作和论文投稿时，应将 evidence enhancement 与 calibration 作为可由消融直接支撑的性能贡献，将 bias awareness 表述为可信评估流程中的风险控制模块，并单独讨论 Tie policy 带来的 accuracy 与 Tie recall 权衡。"
    )

    doc.add_heading("8. 数据与可复现性来源", level=1)
    add_bullets(
        doc,
        [
            "正式主表：datasets/model_outputs/sci_tables_v2_20260521_110114/main_results_table.csv",
            "正式消融表：datasets/model_outputs/sci_tables_v2_20260521_110114/ablation_table.csv",
            "显著性检验：datasets/model_outputs/sci_tables_v2_20260521_110114/ablation_significance_table.csv",
            "证据特征组消融：datasets/model_outputs/sci_tables_v2_20260521_110114/evidence_feature_group_ablation_table.csv",
            "QLoRA 三种子消融：datasets/model_outputs/qlora_ablation_3seed_epoch1_1024_summary/ablation_3seed_summary.json",
            "方法摘要显示正式实验通过了 base-score 覆盖、修复后解析覆盖、无启发式正式结果、偏差覆盖和证据 profile 完整性等门禁。",
        ]
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_report())
