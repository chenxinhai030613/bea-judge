from __future__ import annotations

import argparse
import csv
import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    HAS_PYTHON_DOCX = True
except ImportError:  # pragma: no cover - exercised when optional dependency is absent.
    Document = None  # type: ignore[assignment]
    WD_TABLE_ALIGNMENT = None  # type: ignore[assignment]
    WD_CELL_VERTICAL_ALIGNMENT = None  # type: ignore[assignment]
    WD_ALIGN_PARAGRAPH = None  # type: ignore[assignment]
    OxmlElement = None  # type: ignore[assignment]
    qn = None  # type: ignore[assignment]
    Inches = None  # type: ignore[assignment]
    Pt = None  # type: ignore[assignment]
    RGBColor = None  # type: ignore[assignment]
    HAS_PYTHON_DOCX = False


ROOT = Path(__file__).resolve().parent.parent
DEFAULT_TITLE = "BEA-Judge-10K-v2 消融实验结果报告"
DEFAULT_OUTPUT = ROOT / "paper" / "bea_judge_ablation_report.docx"

CORE_TABLES = {
    "ablation": "ablation_table.csv",
    "significance": "ablation_significance_table.csv",
    "evidence_groups": "evidence_feature_group_ablation_table.csv",
    "bias_utility": "bias_risk_utility_table.csv",
    "calibration": "calibration_methods_table.csv",
    "risk_coverage": "risk_coverage_table.csv",
    "ragtruth": "ragtruth_results_table.csv",
}

COLUMN_LABELS = {
    "variant": "Variant",
    "head": "Task",
    "n": "N",
    "accuracy": "Acc.",
    "macro_f1": "Macro-F1",
    "ece": "ECE",
    "brier": "Brier",
    "tie_recall": "Tie recall",
    "review_rate": "Review rate",
    "paired_n": "Paired N",
    "delta_accuracy_full_minus_variant": "Delta Acc.",
    "delta_macro_f1_full_minus_variant": "Delta Macro-F1",
    "mcnemar_full_only_correct": "Full-only",
    "mcnemar_variant_only_correct": "Variant-only",
    "mcnemar_p": "McNemar p",
    "feature_group": "Feature group",
    "weighted_calibration": "Weighted cal.",
    "feature_count": "Feature count",
    "setting": "Setting",
    "method": "Method",
    "split": "Split",
    "mce": "MCE",
    "nll": "NLL",
    "coverage": "Coverage",
    "set_size_avg": "Set size",
    "review_count": "Review N",
    "error_capture_rate": "Error capture",
    "auto_accept_count": "Auto-accept N",
    "auto_accept_accuracy": "Auto-accept Acc.",
    "risk_threshold": "Risk threshold",
    "supported_to_unsupported": "Sup.->Unsup.",
    "unsupported_to_supported": "Unsup.->Sup.",
}

VALUE_LABELS = {
    "w/o Evidence Module": "w/o Evidence",
    "w/o Bias Module": "w/o Bias",
    "w/o Base Judge Scores": "w/o Base Scores",
    "w/o Review Threshold": "w/o Review",
    "Base + fusion calibration only": "Base+fusion only",
    "Raw M-Prometheus-3B only": "Raw M-Prometheus",
    "Text/metadata-only": "Text/metadata",
    "bias_as_decision_features": "bias decision features",
    "no_bias_decision_features": "no bias features",
    "bias_risk_only_review": "bias risk review only",
}


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")


def read_csv_rows(path: Path) -> List[Dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return [dict(row) for row in csv.DictReader(handle)]


def load_json(path: Path) -> Dict[str, Any]:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def first_row(rows: Sequence[Dict[str, str]], **filters: str) -> Optional[Dict[str, str]]:
    for row in rows:
        if all(str(row.get(key, "")) == str(value) for key, value in filters.items()):
            return row
    return None


def cell(row: Optional[Dict[str, str]], key: str, default: str = "-") -> str:
    if not row:
        return default
    value = row.get(key, "")
    return value if value not in {"", None} else default


def compact_float(value: str) -> str:
    try:
        return f"{float(value):.4f}"
    except (TypeError, ValueError):
        return value or "-"


def select_columns(rows: Sequence[Dict[str, str]], columns: Sequence[str]) -> List[List[str]]:
    header = [COLUMN_LABELS.get(column, column) for column in columns]
    body = [[VALUE_LABELS.get(row.get(column, ""), row.get(column, "")) for column in columns] for row in rows]
    return [header, *body] if body else [header, ["暂无数据", *["" for _ in header[1:]]]]


def limit_rows(rows: Sequence[Dict[str, str]], limit: int) -> List[Dict[str, str]]:
    return list(rows[:limit])


def set_cell_shading(cell_obj, fill: str) -> None:
    tc_pr = cell_obj._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell_obj, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc = cell_obj._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_width(cell_obj, width_dxa: int) -> None:
    tc_pr = cell_obj._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_grid(table, widths_dxa: Sequence[int]) -> None:
    tbl = table._tbl
    existing_grid = tbl.tblGrid
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)


def column_widths(data: Sequence[Sequence[str]], total_dxa: int = 9360) -> List[int]:
    if not data or not data[0]:
        return []
    columns = len(data[0])
    scores: List[float] = []
    for column_index in range(columns):
        texts = [str(row[column_index]) if column_index < len(row) else "" for row in data]
        header = texts[0].replace("_", " ")
        max_body = max((len(text) for text in texts[1:]), default=0)
        score = max(len(header) * 0.85, min(max_body, 34) * 0.55, 8)
        if column_index == 0:
            score *= 1.45
        scores.append(score)
    minimum = 620
    remaining = max(total_dxa - minimum * columns, columns)
    score_sum = sum(scores) or 1
    widths = [minimum + int(remaining * score / score_sum) for score in scores]
    widths[-1] += total_dxa - sum(widths)
    return widths


def apply_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title_page(document: Document, title: str, ablation_report: Dict[str, Any]) -> None:
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("四模块模型框架的模块级、证据、偏差、校准与复核消融分析").italic = True

    document.add_heading("1. 标题页与实验元信息", level=1)
    metadata = [
        ("生成时间", utc_now()),
        ("消融报告时间戳", str(ablation_report.get("created_at", "-"))),
        ("输入数据", str(ablation_report.get("input_dataset", "-"))),
        ("基础 Judge 输出", str(ablation_report.get("judge_output_path", "-"))),
        ("正式定位", "基于真实 M-Prometheus-3B 输出的偏差感知与证据增强校准框架"),
    ]
    add_kv_table(document, metadata)
    document.add_page_break()


def add_kv_table(document: Document, items: Sequence[tuple[str, str]]) -> None:
    table = document.add_table(rows=len(items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table)
    widths = [1900, 7460]
    set_table_grid(table, widths)
    for row_index, (key, value) in enumerate(items):
        cells = table.rows[row_index].cells
        cells[0].text = key
        cells[1].text = value
        set_cell_shading(cells[0], "F2F4F7")
        for column_index, cell_obj in enumerate(cells):
            cell_obj.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell_obj, widths[column_index])
            set_cell_margins(cell_obj)


def add_callout(document: Document, title: str, body: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table)
    cell_obj = table.cell(0, 0)
    set_cell_shading(cell_obj, "F4F6F9")
    set_cell_margins(cell_obj, top=120, bottom=120, start=160, end=160)
    paragraph = cell_obj.paragraphs[0]
    run = paragraph.add_run(f"{title}：")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F3A5F")
    paragraph.add_run(body)


def add_table(document: Document, caption: str, data: Sequence[Sequence[str]], font_size: float = 8.0) -> None:
    document.add_paragraph(caption, style="Heading 3")
    if not data:
        document.add_paragraph("暂无数据。")
        return
    table = document.add_table(rows=len(data), cols=len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table)
    widths = column_widths(data)
    set_table_grid(table, widths)
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row):
            cell_obj = table.cell(r_idx, c_idx)
            cell_obj.text = str(value)
            cell_obj.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if c_idx < len(widths):
                set_cell_width(cell_obj, widths[c_idx])
            set_cell_margins(cell_obj)
            if r_idx == 0:
                set_cell_shading(cell_obj, "F2F4F7")
            for paragraph in cell_obj.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "Calibri"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    if r_idx == 0:
                        run.bold = True
    document.add_paragraph()


def summarize_findings(tables: Dict[str, List[Dict[str, str]]]) -> List[str]:
    ablation = tables["ablation"]
    significance = tables["significance"]
    bias_utility = tables["bias_utility"]
    full_pair = first_row(ablation, variant="Full BEA-Judge", head="pairwise")
    full_fact = first_row(ablation, variant="Full BEA-Judge", head="factuality")
    no_evidence_fact = first_row(ablation, variant="w/o Evidence Module", head="factuality")
    no_cal_pair_sig = first_row(significance, variant="w/o Calibration", head="pairwise")
    no_evidence_sig = first_row(significance, variant="w/o Evidence Module", head="factuality")
    bias_full = first_row(bias_utility, setting="bias_as_decision_features")
    no_bias = first_row(bias_utility, setting="no_bias_decision_features")
    return [
        (
            "完整 BEA-Judge 在 pairwise test 上达到 "
            f"accuracy={cell(full_pair, 'accuracy')}、macro-F1={cell(full_pair, 'macro_f1')}、"
            f"Tie recall={cell(full_pair, 'tie_recall')}；在 factuality test 上达到 "
            f"accuracy={cell(full_fact, 'accuracy')}、macro-F1={cell(full_fact, 'macro_f1')}。"
        ),
        (
            "证据增强是事实性任务的主贡献来源：移除 evidence 后 factuality macro-F1 为 "
            f"{cell(no_evidence_fact, 'macro_f1')}，Full 相对该变体的 Δmacro-F1 为 "
            f"{cell(no_evidence_sig, 'delta_macro_f1_full_minus_variant')}，McNemar p={cell(no_evidence_sig, 'mcnemar_p')}。"
        ),
        (
            "校准模块主要支撑 pairwise 决策稳定性：Full 相对 w/o Calibration 的 Δmacro-F1 为 "
            f"{cell(no_cal_pair_sig, 'delta_macro_f1_full_minus_variant')}，McNemar p={cell(no_cal_pair_sig, 'mcnemar_p')}。"
        ),
        (
            "偏差模块不应被表述为单纯 accuracy booster。no-bias decision features 的 macro-F1 为 "
            f"{cell(no_bias, 'macro_f1')}，而 bias-as-decision-features 的 review capture rate 为 "
            f"{cell(bias_full, 'review_capture_rate')}。这支持将 bias module 定位为风险识别与复核优先级机制。"
        ),
    ]


def filter_risk_rows(rows: Sequence[Dict[str, str]]) -> List[Dict[str, str]]:
    selected: List[Dict[str, str]] = []
    wanted = {"0.2", "0.2000", "0.2004", "0.4", "0.3998", "0.4000", "0.499", "0.4990", "0.4995"}
    for row in rows:
        if row.get("review_rate") in wanted:
            selected.append(row)
    return selected or limit_rows(rows, 8)


def build_document(
    *,
    title: str,
    tables_dir: Path,
    ablation_report_path: Path,
    output_path: Path,
) -> Path:
    if not HAS_PYTHON_DOCX:
        raise ImportError("python-docx is required to build the ablation report")
    tables = {name: read_csv_rows(tables_dir / filename) for name, filename in CORE_TABLES.items()}
    ablation_report = load_json(ablation_report_path)

    document = Document()
    apply_styles(document)
    add_title_page(document, title, ablation_report)

    document.add_heading("2. 执行摘要", level=1)
    for finding in summarize_findings(tables):
        document.add_paragraph(finding)
    add_callout(
        document,
        "投稿叙事",
        "本文应强调证据增强与融合校准的可验证贡献，同时将偏差模块严格定位为风险控制和复核优先级，而不是泛化为整体准确率提升模块。",
    )
    document.add_page_break()

    document.add_heading("3. 模块级消融结果", level=1)
    add_table(
        document,
        "表1 模块级消融主结果",
        select_columns(
            tables["ablation"],
            ["variant", "head", "n", "accuracy", "macro_f1", "ece", "brier", "tie_recall", "review_rate"],
        ),
    )
    document.add_page_break()

    document.add_heading("4. 显著性检验", level=1)
    add_table(
        document,
        "表2 Full BEA-Judge 相对消融变体的配对显著性",
        select_columns(
            tables["significance"],
            [
                "variant",
                "head",
                "paired_n",
                "delta_accuracy_full_minus_variant",
                "delta_macro_f1_full_minus_variant",
                "mcnemar_full_only_correct",
                "mcnemar_variant_only_correct",
                "mcnemar_p",
            ],
        ),
        font_size=7.5,
    )

    document.add_heading("5. 证据增强模块细粒度消融", level=1)
    add_table(
        document,
        "表3 Factuality evidence feature group ablation",
        select_columns(
            tables["evidence_groups"],
            ["feature_group", "weighted_calibration", "feature_count", "accuracy", "macro_f1", "ece", "brier"],
        ),
    )

    document.add_page_break()
    document.add_heading("6. 偏差模块风险效用分析", level=1)
    add_table(
        document,
        "表4 Bias risk utility",
        select_columns(
            tables["bias_utility"],
            ["setting", "head", "n", "accuracy", "macro_f1", "ece", "review_rate", "review_capture_rate"],
        ),
    )

    document.add_heading("7. 校准方法对比", level=1)
    add_table(
        document,
        "表5 Calibration methods comparison",
        select_columns(
            tables["calibration"],
            ["method", "split", "accuracy", "ece", "mce", "brier", "nll", "coverage", "set_size_avg"],
        ),
        font_size=7.5,
    )

    document.add_page_break()
    document.add_heading("8. Risk-coverage 与 review capture", level=1)
    add_table(
        document,
        "表6 代表性 review budget 下的风险覆盖结果",
        select_columns(
            filter_risk_rows(tables["risk_coverage"]),
            [
                "head",
                "split",
                "review_rate",
                "review_count",
                "error_capture_rate",
                "auto_accept_count",
                "auto_accept_accuracy",
                "risk_threshold",
            ],
        ),
        font_size=7.5,
    )

    document.add_heading("9. RAGTruth 事实性压力测试", level=1)
    add_table(
        document,
        "表7 RAGTruth 错误分类与校准表现",
        select_columns(
            tables["ragtruth"],
            [
                "split",
                "n",
                "accuracy",
                "macro_f1",
                "ece",
                "brier",
                "review_rate",
                "supported_to_unsupported",
                "unsupported_to_supported",
            ],
        ),
    )

    document.add_heading("10. 结论与投稿叙事建议", level=1)
    recommendations = [
        "将 evidence module 写成 factuality reliability 的主要贡献，并用 w/o Evidence 的显著下降支撑。",
        "将 calibration module 写成概率可靠性、Tie policy 和复核阈值的稳定化机制。",
        "将 bias module 写成 risk-control 与 review prioritization，不宣称其必然提升总体 accuracy。",
        "将 RAGTruth 单独作为 response-level hallucination 压力测试，避免总体 factuality 指标掩盖其难度。",
        "主文报告模块级消融和显著性检验；附录保留校准方法、risk-coverage 和细粒度证据特征分析。",
    ]
    for item in recommendations:
        document.add_paragraph(item, style="List Bullet")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a DOCX ablation experiment report.")
    parser.add_argument("--tables-dir", type=Path, required=True)
    parser.add_argument("--ablation-report", type=Path, required=True)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--title", default=DEFAULT_TITLE)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    output = build_document(
        title=args.title,
        tables_dir=args.tables_dir,
        ablation_report_path=args.ablation_report,
        output_path=args.output,
    )
    print(json.dumps({"output": str(output.resolve())}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
