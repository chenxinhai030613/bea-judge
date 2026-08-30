from __future__ import annotations

import csv
import subprocess
from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent.parent
PAPER = ROOT / "paper_ieee_bea_judge"
FORMAL = ROOT / "datasets" / "model_outputs" / "sci_tables_v2_20260521_110114"
EXTENDED = ROOT / "datasets" / "model_outputs" / "sci_tables_extended_20260522"
MD = PAPER / "bea_judge_ieee_word_source.md"
DOCX = PAPER / "bea_judge_ieee.docx"
CSL = Path("D:/texlive/2026/texmf-dist/tex/latex/citation-style-language/styles/ieee.csl")


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def fnum(value: str, digits: int = 4) -> str:
    if not value:
        return "--"
    return f"{float(value):.{digits}f}"


def table_markdown(headers: list[str], rows: list[list[str]]) -> str:
    header = "| " + " | ".join(headers) + " |"
    sep = "| " + " | ".join("---" for _ in headers) + " |"
    body = ["| " + " | ".join(row) + " |" for row in rows]
    return "\n".join([header, sep, *body])


def normalize_markdown(text: str) -> str:
    lines = text.strip().splitlines()
    return "\n".join(line[8:] if line.startswith("        ") else line for line in lines) + "\n"


def sources_table() -> str:
    rows = read_csv(FORMAL / "source_provenance_table.csv")
    names = {
        "helpsteer2": "HelpSteer2",
        "oasst1": "OASST1",
        "offsetbias": "OffsetBias",
        "ragtruth": "RAGTruth",
        "rewardbench": "RewardBench",
    }
    return table_markdown(
        ["来源", "许可", "纳入样本", "可再分发", "训练纳入"],
        [
            [
                names.get(r["source"], r["source"]),
                r["license"],
                r["accepted_records"],
                "是" if r["redistribution_allowed"] == "True" else "否",
                "是" if r["admission_allowed"] == "True" else "否",
            ]
            for r in rows
        ],
    )


def main_table() -> str:
    rows = [r for r in read_csv(FORMAL / "main_results_table.csv") if r["split"] == "test"]
    out = []
    for r in rows:
        out.append(
            [
                "成对偏好" if r["head"] == "pairwise" else "事实性",
                r["n"],
                fnum(r["accuracy"]),
                fnum(r["macro_f1"]),
                fnum(r["ece"]),
                fnum(r["brier"]),
                fnum(r["tie_recall"]) if r["tie_recall"] else "--",
                fnum(r["review_rate"]),
            ]
        )
    return table_markdown(["任务头", "n", "Acc.", "Macro-F1", "ECE", "Brier", "Tie-R", "Review"], out)


def ablation_table() -> str:
    rows = [r for r in read_csv(FORMAL / "ablation_table.csv") if r["split"] == "test"]
    names = {
        "Full BEA-Judge": "完整模型",
        "w/o Bias Module": "去除偏差模块",
        "w/o Evidence Module": "去除证据模块",
        "w/o Calibration": "去除校准策略",
    }
    order = ["Full BEA-Judge", "w/o Bias Module", "w/o Evidence Module", "w/o Calibration"]
    out = []
    for variant in order:
        for head in ["pairwise", "factuality"]:
            r = next((x for x in rows if x["variant"] == variant and x["head"] == head), None)
            if not r:
                continue
            out.append(
                [
                    names[variant],
                    "成对偏好" if head == "pairwise" else "事实性",
                    fnum(r["accuracy"]),
                    fnum(r["macro_f1"]),
                    fnum(r["ece"]),
                    fnum(r["brier"]),
                    fnum(r["tie_recall"]) if r["tie_recall"] else "--",
                ]
            )
    return table_markdown(["变体", "任务头", "Acc.", "Macro-F1", "ECE", "Brier", "Tie-R"], out)


def significance_table() -> str:
    rows = read_csv(FORMAL / "ablation_significance_table.csv")
    keep = {
        ("w/o Evidence Module", "factuality"),
        ("w/o Calibration", "pairwise"),
        ("w/o Bias Module", "pairwise"),
    }
    names = {
        "w/o Evidence Module": "去除证据模块",
        "w/o Calibration": "去除校准策略",
        "w/o Bias Module": "去除偏差模块",
    }
    out = []
    for r in rows:
        if (r["variant"], r["head"]) not in keep:
            continue
        out.append(
            [
                names[r["variant"]],
                "成对偏好" if r["head"] == "pairwise" else "事实性",
                r["paired_n"],
                fnum(r["delta_accuracy_full_minus_variant"]),
                f"[{fnum(r['delta_accuracy_ci95_low'])}, {fnum(r['delta_accuracy_ci95_high'])}]",
                fnum(r["mcnemar_p"], 6),
            ]
        )
    return table_markdown(["变体", "任务头", "n", "Delta Acc.", "95% CI", "p"], out)


def calibration_table() -> str:
    rows = [r for r in read_csv(EXTENDED / "calibration_methods_table.csv") if r["split"] == "test"]
    names = {
        "temperature": "Temperature Scaling",
        "platt": "Platt Scaling",
        "isotonic": "Isotonic Regression",
        "vector_scaling": "Vector Scaling",
        "conformal": "Conformal Prediction",
    }
    out = []
    for r in rows:
        out.append(
            [
                names.get(r["method"], r["method"]),
                fnum(r["accuracy"]),
                fnum(r["ece"]),
                fnum(r["mce"]),
                fnum(r["brier"]),
                fnum(r["nll"]),
                fnum(r["coverage"]) if r["coverage"] else "--",
            ]
        )
    return table_markdown(["方法", "Acc.", "ECE", "MCE", "Brier", "NLL", "Coverage"], out)


def eq(body: str, number: int) -> str:
    return f"$$\n{body}\\qquad({number})\n$$"


def make_markdown() -> str:
    sources = sources_table()
    main_results = main_table()
    ablations = ablation_table()
    significance = significance_table()
    calibration = calibration_table()
    eq1 = eq("p_\\theta(y\\mid x_i,h)=\\operatorname{softmax}\\left(\\mathbf{W}_h\\mathbf{z}_i+\\mathbf{b}_h\\right)_y,\\quad y\\in\\mathcal{Y}_h", 1)
    eq2 = eq("z_{ij}=\\frac{\\phi_j(x_i)-\\mu_j}{\\sigma_j+\\epsilon}", 2)
    eq3 = eq("m_i=|\\Delta s_i|", 3)
    eq4 = eq("r_i^b=\\max\\{r_i^{pos},r_i^{len},r_i^{fmt},r_i^{rub},r_i^{src}\\}", 4)
    eq5 = eq("S(a,c,r)=\\max\\left(S(a,c\\oplus r),\\;0.65S(a,c)+0.35S(a,r)\\right)", 5)
    eq6 = eq("r_i^e=\\max_k g_k(x_i)", 6)
    eq7 = eq("\\mathcal{L}_h=-\\frac{1}{N_h}\\sum_{i=1}^{N_h}\\log p_\\theta(y_i\\mid x_i,h)+\\frac{\\lambda}{2}\\|\\mathbf{W}_h\\|_2^2", 7)
    eq8 = eq("J=\\operatorname{MacroF1}+0.25\\operatorname{Accuracy}-0.05\\operatorname{ECE}", 8)
    eq9 = eq("\\hat{p}_T(y\\mid x)=\\operatorname{softmax}\\left(\\frac{\\log p_\\theta(y\\mid x)}{T}\\right)", 9)
    eq10 = eq("\\operatorname{ECE}=\\sum_{b=1}^{B}\\frac{|I_b|}{n}\\left|\\operatorname{acc}(I_b)-\\operatorname{conf}(I_b)\\right|", 10)
    eq11 = eq("r_i=1-\\max_y \\hat{p}_T(y\\mid x_i)", 11)

    return normalize_markdown(dedent(
        f"""
        ---
        title: "BEA-Judge：面向生成式人工智能内容评估的偏差感知证据增强评判校准框架"
        author: "作者信息待补充"
        lang: zh-CN
        bibliography: references.bib
        csl: "D:/texlive/2026/texmf-dist/tex/latex/citation-style-language/styles/ieee.csl"
        ---

        **摘要—** 生成式人工智能内容评估正在从静态基准测试转向可解释、可校准且可复核的评判系统。现有 LLM-as-a-Judge 方法能够降低人工评估成本，但在成对偏好排序、事实性判断和偏差控制上仍面临三类问题：评判器对位置、长度和格式等表层因素敏感；事实性错误常以局部证据缺失、数值或实体错配等细粒度形式出现；单次模型输出缺乏面向生产评审的置信度校准与复核机制。本文提出 BEA-Judge，一个建立在真实 M-Prometheus-3B 输出之上的偏差感知证据增强评判校准框架。该框架不训练新的大语言模型，而是将基座评判器分数、文本差异、偏差风险、证据支持特征和数据源元信息输入任务特定 softmax 头，并在开发集上进行温度缩放、Tie 策略选择和风险阈值选择。我们构建 BEA-Judge-10K v2，共 10,200 条样本，覆盖开放问答、成对偏差评估和 RAG 事实性判断。冻结测试结果显示，成对偏好头达到 0.7512 accuracy、0.6730 macro-F1 和 0.0558 ECE；事实性头达到 0.7649 accuracy、0.7405 macro-F1 和 0.0377 ECE。消融实验表明，证据增强是事实性可靠性的主要贡献，去除证据模块使事实性 macro-F1 从 0.7405 降至 0.6542；校准策略显著提升成对任务的 macro-F1 与 Tie 召回。实验同时显示，偏差模块更适合作为风险识别与复核优先级机制，而非简单追求总体准确率提升。

        **关键词—** 生成式人工智能评估；LLM-as-a-Judge；校准；事实性验证；偏差感知；风险复核；RAG 幻觉检测

        # I. 引言

        大语言模型输出的开放性使传统基于精确答案匹配的评估范式难以覆盖真实应用中的质量、偏好和事实性要求。LLM-as-a-Judge 通过将强模型或专用评判模型作为评审器，为开放式对话和响应排序提供了可扩展方案 [@zheng2023judging; @kim2024prometheus; @kim2024prometheus2]。然而，近期研究也表明，评判器容易受到位置偏差、长度偏差、格式偏差和自信度错配的影响；这些偏差会降低评估结论的可解释性，尤其会影响需要人工复核或质量门禁的场景 [@zheng2023judging; @park2024offsetbias]。

        本文提出 BEA-Judge。与训练新的评判大模型不同，BEA-Judge 将 M-Prometheus-3B 的真实成对评判输出作为基座信号，再加入文本结构、偏差风险、证据支持和数据源特征，训练轻量级任务特定 softmax 头。随后，系统在开发集上选择温度缩放、数据集级温度策略、Tie 决策策略和低置信复核阈值。该设计将自动评判拆分为“基础判断、风险诊断、概率校准和复核控制”四个可审计环节。

        本文贡献如下：

        - 提出一种面向生成式内容评估的偏差感知证据增强校准框架，将基座 LLM 评判输出、偏差风险特征和证据事实性特征统一到任务特定概率模型中。
        - 构建并审计 BEA-Judge-10K v2，覆盖开放问答、成对偏差和 RAG 事实性三类任务，并明确训练、开发和测试划分及数据许可状态。
        - 在冻结测试集上报告主结果、消融实验、显著性检验、校准方法比较、风险覆盖曲线和证据错误子类型分析，给出严格可复现的实验门禁。
        - 明确方法边界：BEA-Judge 不是新训练的大语言模型，也不声称解决原子级事实核验。

        ![图 1. BEA-Judge 框架流程。系统以 BEA-Judge-10K v2 为输入，使用 M-Prometheus-3B 产生真实基座评判输出，再融合偏差风险与证据支持特征，最终通过任务特定校准头输出标签、置信度、风险分数和复核标记。](figures/fig1_pipeline.svg)

        # II. 相关工作

        MT-Bench 和 Chatbot Arena 推动了使用 LLM 近似人类偏好评估的研究 [@zheng2023judging]。Prometheus 系列进一步强调开放评判模型、评分准则和参考材料的重要性 [@kim2024prometheus; @kim2024prometheus2]。M-Prometheus 将评判能力扩展到多语言直接评分和成对比较场景 [@pombal2025mprometheus]，因此适合作为本文中文与英文混合实验的基座评判器。

        LLM 评判器已知存在位置、长度、格式、准则敏感性等偏差 [@zheng2023judging; @park2024offsetbias]。事实性方面，RAGTruth 提供了检索增强生成中幻觉和证据不一致的细粒度语料 [@niu2024ragtruth]。本文借鉴其证据敏感问题设置，但采用确定性证据特征和任务特定概率头，而非直接进行原子声明级判定。

        # III. 数据集构建

        BEA-Judge-10K v2 由 10,200 条样本组成，包含三类任务：开放问答 4,000 条、成对偏差评估 2,700 条、RAG 事实性评估 3,500 条。数据来源包括 HelpSteer2 [@wang2024helpsteer2]、OpenAssistant Conversations [@kopf2023openassistant]、OffsetBias [@park2024offsetbias]、RAGTruth [@niu2024ragtruth] 以及项目内中文专业标注数据。训练、开发和测试划分分别为 7,084、1,578 和 1,538 条；语言分布为英文 9,148 条、中文 1,052 条。

        表 1. 数据来源、许可与纳入状态。

        {sources}

        ![图 2. BEA-Judge-10K v2 的数据来源分布。](figures/fig2_dataset_distribution.svg)

        数据构建执行了样本数、字段类型、枚举值、重复 ID、重复内容和跨 split 泄漏检查。正式结果门禁要求样本数位于 9,500 到 10,200 区间，成对基座评判覆盖率为 6,946/6,946，启发式 fallback 行数为 0，未解析失败为 0。实验索引显示所有正式门禁均通过。

        # IV. 方法

        对样本 $x_i$，成对偏好头的标签空间为 $\\mathcal{{Y}}_p=\\{{A>B,B>A,Tie\\}}$，事实性头的有效标签空间为 $\\mathcal{{Y}}_f=\\{{supported,unsupported\\}}$。给定基座评判器输出、文本特征、偏差风险和证据特征，BEA-Judge 学习任务特定概率分布：

        {eq1}

        特征标准化仅使用训练集统计量：

        {eq2}

        成对头使用 M-Prometheus-3B 输出，包括回答 A、B 的分数、分差、绝对 margin、基座预测标签及顺序置换诊断。对基座分差 $\\Delta s=s_A-s_B$，基座 margin 记为：

        {eq3}

        偏差模块检测位置、长度、格式、rubric sensitivity 和数据源风险。每个风险项被裁剪到 $[0,1]$，总体偏差风险定义为最大风险：

        {eq4}

        事实性模块计算回答与上下文、参考答案之间的证据支持。令 $a$ 为回答，$c$ 为上下文，$r$ 为参考答案，则回答支持度定义为：

        {eq5}

        证据风险由数值缺口、日期缺口、实体缺口、否定错配、比较错配和低支持句子比例共同决定：

        {eq6}

        每个任务头使用交叉熵与 $L_2$ 正则训练：

        {eq7}

        学习率、batch size、$L_2$ 和 epoch 数在开发集上搜索，选择目标为：

        {eq8}

        训练后使用开发集选择温度 $T$：

        {eq9}

        ECE 采用等宽置信度分桶：

        {eq10}

        对任一输出，复核风险定义为：

        {eq11}

        若 $r_i\\ge\\tau_h$，则样本进入人工复核队列。阈值 $\\tau_h$ 在开发集上选择，目标是在可行时捕获至少 80% 的开发集错误。

        # V. 实验设置

        所有超参数、温度、Tie 策略和复核阈值仅使用训练集和开发集选择，测试集仅用于最终报告。成对头训练样本数为 4,806、开发样本数为 1,087、测试样本数为 1,053。事实性头训练样本数为 2,278、开发样本数为 491、测试样本数为 485。基座评判器为 Unbabel/M-Prometheus-3B，温度为 0.0，最大新 token 数为 256，正式实验不允许启发式 fallback。

        本文报告 Accuracy、Macro-F1、ECE、Brier score、Tie recall 和 review rate。消融实验包括去除偏差模块、去除证据模块和去除校准策略。扩展基线包括 Raw M-Prometheus-3B only、Text/metadata-only、Base + fusion calibration only、去除基座分数、去除 Tie 策略和去除复核阈值。校准对比包括 temperature scaling、Platt scaling、isotonic regression、vector scaling 与 split conformal prediction [@guo2017calibration; @platt1999probabilistic]。

        # VI. 实验结果

        表 2 与图 3 给出冻结测试集主结果。成对偏好头在 1,053 条测试样本上达到 0.7512 accuracy 和 0.6730 macro-F1，ECE 为 0.0558，Tie recall 为 0.5231。事实性头在 485 条测试样本上达到 0.7649 accuracy 和 0.7405 macro-F1，ECE 为 0.0377。

        表 2. 冻结测试集主结果。

        {main_results}

        ![图 3. 冻结测试集上成对偏好头与事实性头的主指标。](figures/fig3_main_results.svg)

        表 3 和图 4 展示关键模块消融。事实性头对证据模块高度敏感：去除证据后 accuracy 从 0.7649 降至 0.6928，macro-F1 从 0.7405 降至 0.6542。成对头对校准策略高度敏感：去除校准后 macro-F1 从 0.6730 降至 0.6402，Tie recall 从 0.5231 降至 0.3923。

        表 3. 模块消融实验。

        {ablations}

        ![图 4. 关键模块消融的 Macro-F1 对比。](figures/fig4_ablation.svg)

        表 4. 关键消融的配对显著性检验。Delta 表示完整模型减去对应消融变体。

        {significance}

        表 5 报告成对偏好头的扩展校准对比。Platt scaling 在测试集上获得最低 ECE 0.0194，isotonic regression 获得最高 accuracy 0.7692 和较低 ECE 0.0263，但 MCE 较高。本文正式结果采用温度缩放和任务策略组合，因为其参数更少、可解释性更高，并与开发集选择目标一致。

        表 5. 成对偏好头校准方法对比。

        {calibration}

        图 5 给出测试集风险覆盖曲线。成对头在复核约 49.95% 样本时捕获 85.11% 错误，自动接受部分 accuracy 达到 0.9260；事实性头在复核约 49.90% 样本时捕获 75.44% 错误，自动接受部分 accuracy 达到 0.8848。

        ![图 5. 风险阈值复核曲线。](figures/fig5_risk_coverage.svg)

        RAGTruth 测试子集包含 372 条事实性样本，accuracy 为 0.6962，macro-F1 为 0.6363，ECE 为 0.0360，review rate 为 0.6909。图 6 显示，numeric evidence gap、low support sentence ratio、low support anchor sentence ratio 等子类型具有较高错误率，同时复核捕获率通常超过 0.87。

        ![图 6. 事实性证据错误子类型诊断。](figures/fig6_evidence_subtypes.svg)

        分数据集结果显示，事实性头在 ARES-NQ 测试切片上达到 0.9912 accuracy 与 0.9911 macro-F1，在 RAGTruth 上为 0.6962 accuracy 与 0.6363 macro-F1，说明 RAGTruth 是当前事实性任务的主要困难来源。成对头在 OffsetBias 上达到 0.8690 accuracy 与 0.8690 macro-F1，在 HelpSteer2 上为 0.6678 accuracy 与 0.5373 macro-F1，在 MT-Bench 与 PandaLM 切片上表现较弱。

        # VII. 讨论

        事实性错误通常不是全局语义相似度不足，而是局部证据与回答之间存在实体、数值、日期、否定或比较关系错配。证据特征组消融显示，从 overlap-only 到加入 numeric/date/entity 特征，accuracy 由 0.7175 提升至 0.7361；加入 sentence/local-risk 后进一步提升至 0.7629；再加入加权校准后达到 0.7649 accuracy 和 0.7405 macro-F1。

        成对偏好任务中的 Tie 标签常对应边界样本。未校准模型在开发集的 Tie recall 为 0.3636，完整模型提升到 0.6558；测试集上未校准 Tie recall 为 0.3923，完整模型为 0.5231。虽然 Tie 策略可能牺牲一部分总体 accuracy，但能提升宏平均指标和边界样本可解释性。

        本文有三个明确边界。第一，BEA-Judge 不是新训练的大语言模型，无法替代基座评判器的语言理解能力。第二，证据模块是确定性特征抽取与统计校准，不等价于完整的原子声明级事实核验。第三，部分数据集切片样本量较小，分数据集结论应作为诊断而非泛化保证。

        # VIII. 结论

        本文提出 BEA-Judge，一个基于真实 M-Prometheus-3B 输出的偏差感知证据增强评判校准框架。该框架通过轻量 softmax 头、证据风险特征、偏差风险特征和开发集校准策略，在 BEA-Judge-10K v2 上实现了可复核的生成式内容评估。冻结测试结果显示，BEA-Judge 在成对偏好与事实性任务上均取得较稳定的 accuracy、macro-F1 和 ECE；消融实验确认，证据增强是事实性可靠性的主要来源，校准和 Tie 策略是成对偏好稳定性的关键。

        # 数据与代码可用性

        处理后的 BEA-Judge-10K v2、源数据 manifest、预处理脚本、校验报告、模型输出表和图表应在投稿前归档至可签发 DOI 的仓库。HelpSteer2、OASST1、OffsetBias 和 RAGTruth 的再分发许可已通过项目审计；RewardBench 因混合子集许可限制未纳入正式训练。正式结果使用的关键路径包括 `datasets/processed/bea_judge_cleaned_10000.json`、`datasets/model_outputs/bea_judge_20260521_110114/validation_report.json` 和 `datasets/model_outputs/sci_tables_v2_20260521_110114/`。

        # 致谢

        作者感谢公开数据集和开放评判模型社区。本文草稿使用人工智能工具辅助组织语言、生成 Word/LaTeX 和绘制矢量图；所有实验事实、结论边界和投稿版本仍需由作者逐项核验并承担责任。

        # 参考文献

        ::: {{#refs}}
        :::
        """
    ))


def write_markdown() -> None:
    PAPER.mkdir(parents=True, exist_ok=True)
    MD.write_text(make_markdown(), encoding="utf-8", newline="\n")


def build_docx() -> None:
    cmd = [
        "pandoc",
        str(MD.name),
        "--from=markdown+tex_math_dollars",
        "--to=docx",
        "--citeproc",
        f"--csl={CSL}",
        "--resource-path=.;figures",
        f"--output={DOCX.name}",
    ]
    subprocess.run(cmd, cwd=PAPER, check=True)


def style_docx() -> None:
    doc = Document(DOCX)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    for style_name in ["Normal", "Body Text", "First Paragraph"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style.font.size = Pt(10)

    for style_name in ["Title"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style.font.size = Pt(18)
            style.font.bold = True

    for style_name in ["Heading 1", "Heading 2"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style.font.size = Pt(12 if style_name == "Heading 1" else 11)
            style.font.bold = True

    for p in doc.paragraphs:
        if p.style.name in {"Title", "Author"}:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.text.startswith("**摘要"):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for table in doc.tables:
        for style_name in ("Table Grid", "TableGrid", "网格型"):
            try:
                table.style = style_name
                break
            except KeyError:
                continue
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)

    # Keep vector figures editable/sharp in Word while preventing oversized placement.
    for shape in doc.inline_shapes:
        if shape.width > Inches(6.5):
            ratio = Inches(6.5) / shape.width
            shape.width = Inches(6.5)
            shape.height = int(shape.height * ratio)

    doc.save(DOCX)


def main() -> None:
    write_markdown()
    build_docx()
    style_docx()
    print(DOCX)


if __name__ == "__main__":
    main()
